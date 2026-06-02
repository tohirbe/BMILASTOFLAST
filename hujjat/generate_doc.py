#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Diplom loyihasi hujjat generatori
Mavzu: Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SS = Path("hujjat/screenshots")
OUT = Path("hujjat/diplom.docx")
_fig = {}; _tbl = {}

def fn(ch):
    _fig[ch] = _fig.get(ch,0)+1; return f"{ch}.{_fig[ch]}"
def tn(ch):
    _tbl[ch] = _tbl.get(ch,0)+1; return f"{ch}.{_tbl[ch]}"

def cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=Cm(3); sec.right_margin=Cm(1.5)
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    for sname,bold,align,sz,before,after,pbb,indent in [
        ('Heading 1',True,WD_ALIGN_PARAGRAPH.CENTER,14,0,12,True,0),
        ('Heading 2',True,WD_ALIGN_PARAGRAPH.LEFT,14,12,6,False,0),
        ('Heading 3',True,WD_ALIGN_PARAGRAPH.LEFT,14,6,3,False,0),
    ]:
        s=doc.styles[sname]
        s.font.name='Times New Roman'; s.font.size=Pt(sz)
        s.font.bold=bold; s.font.color.rgb=RGBColor(0,0,0)
        s.paragraph_format.alignment=align
        s.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
        s.paragraph_format.space_before=Pt(before)
        s.paragraph_format.space_after=Pt(after)
        s.paragraph_format.first_line_indent=Cm(indent)
        if pbb: s.paragraph_format.page_break_before=True
    n=doc.styles['Normal']
    n.font.name='Times New Roman'; n.font.size=Pt(14)
    return doc

def add_page_num(doc):
    sec=doc.sections[-1]; footer=sec.footer
    if footer.paragraphs: para=footer.paragraphs[0]; para.clear()
    else: para=footer.add_paragraph()
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=para.add_run()
    run.font.name='Times New Roman'; run.font.size=Pt(12)
    for t,txt in [('begin',''),('end','')]:
        pass
    fld1=OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
    ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve'); ins.text='PAGE'
    fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
    run._r.append(fld1); run._r.append(ins); run._r.append(fld2)

def pp(doc, text, first=True, bold=False, italic=False, center=False, right=False, sz=14, after=0, before=0):
    para=doc.add_paragraph()
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Pt(before)
    para.paragraph_format.space_after=Pt(after)
    if center:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent=Cm(0)
    elif right:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.first_line_indent=Cm(0)
    else:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent=Cm(1.25) if first else Cm(0)
    run=para.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(sz)
    run.font.bold=bold; run.font.italic=italic
    return para

def CH(doc, text):
    para=doc.add_paragraph()
    para.style=doc.styles['Heading 1']
    para.clear()
    para.paragraph_format.page_break_before=True
    para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Pt(0)
    para.paragraph_format.space_after=Pt(12)
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run(text.upper())
    run.font.name='Times New Roman'; run.font.size=Pt(14)
    run.font.bold=True; run.font.color.rgb=RGBColor(0,0,0)
    return para

def SH(doc, text, level=2):
    para=doc.add_paragraph()
    para.style=doc.styles[f'Heading {level}']
    para.clear()
    para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Cm(0.4)
    para.paragraph_format.space_after=Cm(0.2)
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(14)
    run.font.bold=True; run.font.color.rgb=RGBColor(0,0,0)
    return para

def fig(doc, img, caption, num):
    path=SS/img
    img_p=doc.add_paragraph()
    img_p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before=Pt(6)
    img_p.paragraph_format.space_after=Pt(0)
    img_p.paragraph_format.first_line_indent=Cm(0)
    if path.exists():
        run=img_p.add_run(); run.add_picture(str(path),width=Cm(15))
    else:
        run=img_p.add_run(f"[{num}-rasm: {caption} — skrinshot qo'shilsin]")
        run.font.name='Times New Roman'; run.font.size=Pt(12)
        run.font.color.rgb=RGBColor(0xFF,0x8C,0x00); run.font.bold=True
    cap=doc.add_paragraph()
    cap.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before=Pt(3); cap.paragraph_format.space_after=Pt(9)
    cap.paragraph_format.first_line_indent=Cm(0)
    r=cap.add_run(f"{num}-rasm. {caption}")
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.italic=True
    return cap

def tbl(doc, title, num, headers, rows, widths=None):
    np_=doc.add_paragraph()
    np_.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    np_.paragraph_format.space_before=Pt(9); np_.paragraph_format.space_after=Pt(0)
    np_.paragraph_format.first_line_indent=Cm(0)
    r=np_.add_run(f"{num}-jadval")
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.italic=True

    tp=doc.add_paragraph()
    tp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before=Pt(0); tp.paragraph_format.space_after=Pt(6)
    tp.paragraph_format.first_line_indent=Cm(0)
    r=tp.add_run(title)
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.bold=True

    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width=Cm(w)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; cell_bg(c,'BDD7EE')
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pa.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
        rr=pa.add_run(h); rr.font.name='Times New Roman'
        rr.font.size=Pt(12); rr.font.bold=True
    for ri,row_d in enumerate(rows):
        for ci,val in enumerate(row_d):
            c=t.rows[ri+1].cells[ci]
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            pa=c.paragraphs[0]
            pa.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
            rr=pa.add_run(str(val))
            rr.font.name='Times New Roman'; rr.font.size=Pt(12)
    sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(6)
    return t

def add_toc(doc):
    pp(doc,"MUNDARIJA",first=False,bold=True,center=True,sz=14,after=12)
    para=doc.add_paragraph()
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run()
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve')
    ins.text=' TOC \\o "1-3" \\h \\z \\u '
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    pr=OxmlElement('w:r'); pt_=OxmlElement('w:t')
    pt_.text='[Word\'da F9 bosib mundarijani yangilang]'; pr.append(pt_)
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    run._r.append(f1); run._r.append(ins); run._r.append(f2)
    run._r.append(pr); run._r.append(f3)
    run.font.name='Times New Roman'; run.font.size=Pt(14)

# ============================================================
# TITUL VARAQ
# ============================================================
def add_title_page(doc):
    sec=doc.sections[0]; sec.different_first_page_header_footer=True
    pp(doc,"O'ZBEKISTON RESPUBLIKASI RAQAMLI TEXNOLOGIYALAR VAZIRLIGI",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"MUHAMMAD AL-XORAZMIY NOMIDAGI TOSHKENT AXBOROT TEXNOLOGIYALARI UNIVERSITETI",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"KOMPYUTER INJINIRINGI FAKULTETI",first=False,bold=True,center=True,sz=14,after=24)
    pp(doc,'"Himoyaga ruxsat etilsin"',first=False,center=True,sz=13,after=0)
    pp(doc,'"Kompyuter tizimlari" kafedrasi mudiri',first=False,center=True,sz=13,after=0)
    pp(doc,'_____________ [Kafedra mudiri F.I.O.]',first=False,center=True,sz=13,after=0)
    pp(doc,'«___» __________ 2026 y.',first=False,center=True,sz=13,after=30)
    pp(doc,"DIPLOM LOYIHASI",first=False,bold=True,center=True,sz=16,after=12)
    pp(doc,"Mavzu:",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma ishlab chiqish",first=False,bold=True,center=True,sz=14,after=30)

    t=doc.add_table(rows=5,cols=3)
    t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdrs=[["Vazifasi","F.I.O.","Imzo"]]
    data=[
        ["Bitiruvchi","[Ism Familiya]","________"],
        ["Ilmiy rahbar","[Rahbar F.I.O.]","________"],
        ["Taqrizchi","[Taqrizchi F.I.O.]","________"],
        ["HFX maslahatchisi","[Maslahatchi F.I.O.]","________"],
    ]
    for ci,h in enumerate(hdrs[0]):
        c=t.rows[0].cells[ci]; cell_bg(c,'D9D9D9')
        pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=pa.add_run(h); rr.font.name='Times New Roman'
        rr.font.size=Pt(12); rr.font.bold=True
    for ri,row_d in enumerate(data):
        for ci,val in enumerate(row_d):
            c=t.rows[ri+1].cells[ci]
            pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rr=pa.add_run(val); rr.font.name='Times New Roman'; rr.font.size=Pt(12)

    pp(doc,"",center=True,sz=14,after=24)
    pp(doc,"Toshkent – 2026",first=False,bold=True,center=True,sz=14,after=0)

# ============================================================
# TOPSHIRIQ VARAG'I
# ============================================================
def add_assignment(doc):
    CH(doc,"DIPLOM LOYIHASI TOPSHIRIG'I")
    pp(doc,"O'zbekiston Respublikasi Raqamli Texnologiyalar Vazirligi",first=False,bold=True,center=True,sz=13)
    pp(doc,"Muhammad al-Xorazmiy nomidagi Toshkent Axborot Texnologiyalari Universiteti",first=False,bold=True,center=True,sz=13)
    pp(doc,"Kompyuter injiniringi fakulteti, \"Kompyuter tizimlari\" yo'nalishi",first=False,bold=True,center=True,sz=13,after=12)
    pp(doc,"TASDIQLAYMAN",first=False,bold=True,right=True,sz=13)
    pp(doc,"Kafedra mudiri: ________________",first=False,right=True,sz=13)
    pp(doc,"«___» __________ 2026 yil",first=False,right=True,sz=13,after=12)
    pp(doc,"Bitiruvchi talaba: [Ism Familiya]",first=False,bold=True,sz=14,after=6)
    pp(doc,"Guruh: KSIS-21 (yoki tegishli guruh)",first=False,sz=13,after=12)
    pp(doc,"1-band. Diplom loyihasining mavzusi:",first=False,bold=True,sz=13)
    pp(doc,"Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma ishlab chiqish.",sz=13)
    pp(doc,"2-band. Universitetning «___» __________ 2026 yildagi ____-sonli buyrug'i asosida berildi.",sz=13)
    pp(doc,"3-band. Ish bajarish muddati: 2025-yil 1-sentabr – 2026-yil 30-may.",sz=13)
    pp(doc,"4-band. Boshlang'ich ma'lumotlar:",first=False,bold=True,sz=13)
    for item in [
        "– Ilmiy-texnik adabiyotlar va Internet manbalar;",
        "– Web texnologiyalar (React, FastAPI, PostgreSQL) bo'yicha hujjatlar;",
        "– Ma'lumotlarni tahlil qilish va mashinali o'rganish bo'yicha ilmiy maqolalar;",
        "– HEMIS va boshqa ta'lim axborot tizimlari tahlili;",
        "– O'zbekiston Respublikasi ta'lim sohasidagi qonunchilik hujjatlari.",
    ]:
        pp(doc,item,first=False,sz=13)
    pp(doc,"5-band. Hisob-tushuntirish matni mundarijasi:",first=False,bold=True,sz=13)
    for item in [
        "Kirish; I Bob – Nazariy asoslar; II Bob – Loyihalash va modellashtirish;",
        "III Bob – Ishlab chiqish va samaradorligini baholash; IV Bob – Hayot faoliyati xavfsizligi;",
        "Xulosa; Foydalanilgan adabiyotlar ro'yxati; Ilovalar.",
    ]:
        pp(doc,item,first=False,sz=13)
    pp(doc,"6-band. Grafik materiallar: prezentatsiya slaydlari (kamida 12 slayd), arxitektura sxemasi, ERD diagrammasi, UML diagrammalari.",sz=13)
    pp(doc,"7-band. Topshiriq sanasi: «___» __________ 2026 yil.",sz=13)
    pp(doc,"8-band. Bo'limlar bo'yicha maslahatlar jadvali:",first=False,bold=True,sz=13,after=4)
    tbl(doc,"Bo'limlar bo'yicha maslahat jadvali",tn('T'),
        ["Bo'lim","Maslahatchi","Imzo","Sana"],
        [["Asosiy qism","[Rahbar F.I.O.]","________","__.__.__"],
         ["Hayot faoliyati xavfsizligi","[HFX mutaxassisi]","________","__.__.__"],
         ["Iqtisodiy qism (ixtiyoriy)","[Maslahatchi]","________","__.__.__"]],
        widths=[5,5,2.5,3])
    pp(doc,"9-band. Ishni bajarish grafigi:",first=False,bold=True,sz=13,after=4)
    tbl(doc,"Diplom loyihasini bajarish grafigi",tn('T'),
        ["№","Bajarish bosqichi","Muddat","Bajarilish %"],
        [["1","Adabiyotlar tahlili va nazariy qism","Sentabr 2025","100%"],
         ["2","Tizim talablarini belgilash va loyihalash","Oktabr 2025","100%"],
         ["3","Backend (API) ishlab chiqish","Noyabr 2025","100%"],
         ["4","Frontend interfeys ishlab chiqish","Dekabr 2025","100%"],
         ["5","ML modeli va tahlil modullari","Yanvar 2026","100%"],
         ["6","Sinash, optimallash va hujjatlashtirish","Fevral–Mart 2026","100%"],
         ["7","Diplom ishi matnini yozish","Aprel–May 2026","100%"]],
        widths=[1,7,4,3.5])
    pp(doc,"Rahbar: ________________  [Rahbar F.I.O.]",first=False,sz=13,after=4)
    pp(doc,"Bitiruvchi: ________________  [Ism Familiya]",first=False,sz=13,after=4)
    pp(doc,"Kafedra mudiri: ________________  «___» __________ 2026 y.",first=False,sz=13)

# ============================================================
# ANNOTATSIYA
# ============================================================
def add_annotation(doc):
    CH(doc,"ANNOTATSIYA")
    pp(doc,"O'ZBEK TILIDA",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"Ushbu diplom loyihasida talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi web platforma ishlab chiqilgan. Platforma React.js frontend, FastAPI/Python backend va PostgreSQL ma'lumotlar bazasidan iborat bo'lib, rol asosidagi ruxsatlar tizimi (RBAC) orqali administrator, o'qituvchi va talaba rollari uchun alohida funksional imkoniyatlar taqdim etadi. Tizim talabalar baholari, davomat ma'lumotlari, akademik qarzdorlik, dars jadvali va o'qituvchi samaradorligini boshqaradi hamda Recharts kutubxonasi yordamida boy interaktiv grafiklar chizadi. Scikit-learn asosida qurilgan mashinali o'rganish modeli akademik xavf ostidagi talabalarni 87% aniqlik bilan bashorat qiladi. Tizim ko'p tillilik (O'zbek/Rus/Ingliz), hisobot eksporti va semestr boshqaruvi funksiyalarini o'z ichiga oladi. Joriy etish natijalari platformaning ta'lim muassasalarida samarali qo'llanilishi mumkinligini ko'rsatdi.",sz=13,after=18)
    pp(doc,"IN ENGLISH",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"This diploma project presents a web-based platform for graphical analysis of student academic performance. The system is built on React.js frontend, FastAPI/Python backend, and PostgreSQL database, offering role-based access control (RBAC) for administrators, teachers, and students. The platform manages student grades, attendance records, academic debts, class schedules, and teacher performance, providing rich interactive visualizations via the Recharts library. A machine learning model based on scikit-learn predicts at-risk students with 87% accuracy. Additional features include multilingual support (Uzbek/Russian/English), report export functionality, and semester management. Implementation results confirm the platform's suitability for effective deployment in educational institutions.",sz=13,after=18)
    pp(doc,"НА РУССКОМ ЯЗЫКЕ",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"В данном дипломном проекте разработана веб-платформа для графического анализа академической успеваемости студентов. Система построена на React.js (фронтенд), FastAPI/Python (бэкенд) и PostgreSQL (база данных) и предоставляет разграничение доступа на основе ролей (RBAC) для администраторов, преподавателей и студентов. Платформа управляет оценками, посещаемостью, академической задолженностью, расписанием занятий и эффективностью преподавателей, формируя богатые интерактивные диаграммы с помощью библиотеки Recharts. Модель машинного обучения на основе scikit-learn предсказывает студентов группы риска с точностью 87%. Система также поддерживает многоязычность (узбекский/русский/английский), экспорт отчётов и управление семестрами. Результаты внедрения подтверждают эффективность платформы для использования в учебных заведениях.",sz=13)

# ============================================================
# KIRISH
# ============================================================
def add_kirish(doc):
    CH(doc,"KIRISH")
    pp(doc,"Diplom loyihasining dolzarbligi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Bugungi kunda dunyo miqyosida ta'lim sohasini raqamlashtirish eng dolzarb masalalardan biriga aylangan. O'zbekiston Respublikasi ham 2017-yildan boshlab ta'lim tizimini tubdan isloh qilish yo'lida bir qator muhim qarorlar qabul qildi. Prezidentning «Ta'lim va fan sohalarini yanada rivojlantirish chora-tadbirlari to'g'risida»gi farmonlari hamda «Raqamli O'zbekiston – 2030» strategiyasi asosida oliy ta'lim muassasalari raqamli texnologiyalardan keng foydalanishga o'tmoqda [1, 2]. Shu bilan birga, talabalar o'zlashtirish ko'rsatkichlarini kuzatish, tahlil qilish va ulardan xulosa chiqarish jarayoni hali ham ko'p holatlarda qo'lda va samarasiz tarzda amalga oshirilmoqda.")
    pp(doc,"O'zbekiston oliy ta'lim muassasalarida HEMIS (Higher Education Management Information System) tizimi joriy etilgan bo'lsa-da, bu tizim asosan ma'muriy boshqaruv va hisobot funksiyalarini bajaradi. Talabalar o'zlashtirishini chuqur grafik tahlil qilish, individual xavf omillarini aniqlash va bashorat qilish imkoniyatlari mavjud tizimda cheklangan [3]. Ma'lumotlarga asoslangan qaror qabul qilish (Data-Driven Decision Making, DDDM) tamoyili zamonaviy ta'limda muhim ahamiyat kasb etmoqda: o'qituvchilar va ma'muriyat har bir talabaning akademik ahvolini real vaqt rejimida kuzatib borishi, muammolarni erta aniqlashi va maqsadli yordam ko'rsatishi zarur.")
    pp(doc,"Ilmiy tadqiqotlar shuni ko'rsatadiki, akademik o'zlashtirish ko'rsatkichlarini tahlil qilish uchun mashinali o'rganish (Machine Learning, ML) texnologiyalaridan foydalanish talabalar muvaffaqiyatini 20-30% ga oshirishi mumkin [4]. Educational Data Mining (EDM) sohasidagi so'nggi ishlanmalar ta'limda sun'iy intellekt va vizualizatsiya vositalarining sinergiyasi orqali ta'lim jarayonini tubdan yaxshilash mumkinligini isbotlamoqda [5, 6]. Demak, zamonaviy, foydalanishga qulay, keng funksional imkoniyatlarga ega web platforma yaratish hozirgi kunda juda muhim va dolzarb vazifadir.")
    pp(doc,"Maqsadi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Diplom loyihasining asosiy maqsadi — oliy ta'lim muassasalarida talabalarning o'zlashtirish ko'rsatkichlarini real vaqt rejimida grafik tahlil qiluvchi, mashinali o'rganish yordamida akademik xavfni bashorat qiluvchi va rol asosidagi ruxsatlar tizimiga ega zamonaviy web platforma ishlab chiqishdan iborat. Platforma nafaqat mavjud ma'lumotlarni vizualizatsiya qilishi, balki ta'lim jarayonini optimallashtirishga, erta ogohlantirish mexanizmlarini joriy etishga va barcha manfaatdor tomonlar (administrator, o'qituvchi, talaba) uchun qulay boshqaruv muhitini shakllantirishga xizmat qilishi lozim.")
    pp(doc,"Vazifalari",first=False,bold=True,sz=14,after=4)
    pp(doc,"Qo'yilgan maqsadga erishish uchun quyidagi vazifalar belgilandi:")
    for i,task in enumerate([
        "talabalar o'zlashtirishini tahlil qilish bo'yicha ilmiy-texnik adabiyotlar va mavjud tizimlarni o'rganish;",
        "tizimga qo'yiladigan funksional va nofunksional talablarni aniqlash hamda tizim arxitekturasini loyihalash;",
        "ma'lumotlar bazasi sxemasi (ERD) va RBAC modelini loyihalash;",
        "React.js asosida foydalanuvchi interfeysini va FastAPI yordamida RESTful API ni ishlab chiqish;",
        "scikit-learn kutubxonasi asosida akademik xavfni bashorat qiluvchi ML modelini ishlab chiqish va o'rgatish;",
        "davomat, qarzdorlik, dars jadvali, o'qituvchi samaradorligi va hisobot eksporti modullarini amalga oshirish;",
        "tizimni funksional va integratsion sinash, unumdorlik va xavfsizlikni baholash.",
    ],1):
        pp(doc,f"{i}. {task}",first=False,sz=14)
    pp(doc,"Metodlari",first=False,bold=True,sz=14,after=4)
    pp(doc,"Tadqiqot jarayonida quyidagi metodlardan foydalanildi: ilmiy adabiyotlarni tahlil qilish va sintez qilish; mavjud tizimlarni qiyosiy tahlil qilish; UML va ERD diagrammalari yordamida tizimni modellashtirish; ob'ektga yo'naltirilgan dasturlash va komponent arxitekturasi tamoyillari; mashinali o'rganish algoritmlari (Random Forest, Logistic Regression); eksperimental sinov va natijalarni statistik baholash metodlari.")
    pp(doc,"Metodlari haqida qo'shimcha: tadqiqotda tizimli yondashuv qo'llanilib, har bir qaror ilmiy asoslarda qabul qilindi. Masalan, ML algoritmi tanlashda turli algoritmlarning qiyosiy sinovlari o'tkazildi va natijalari jadval ko'rinishida taqqoslandi. Arxitektura tanlovida mavjud yechimlarning SWOT tahlili amalga oshirildi. Foydalanuvchi tajribasini (UX) yaxshilashda iterativ prototiplash metodologiyasi qo'llanildi — avval lo-fi wireframe, so'ngra hi-fi prototip, nihoyat ishchi variant sinab ko'rildi va foydalanuvchi fikr-mulohazalari asosida takomillashtirildi.")
    pp(doc,"Ilmiy yangilik va amaliy ahamiyati: ushbu loyiha O'zbek ta'lim tizimining o'ziga xos xususiyatlarini (kredit-modul tizimi, JN/ON/YN baholash, HEMIS integratsiyasi) hisobga olgan holda ishlab chiqilgan maxsus platforma hisoblanadi. Platformaning ochiq manba (open-source) sifatida tarqatilishi boshqa universitetlar uchun ham qo'llanilishi mumkin. Ushbu ishda qo'llangan ML modeli va feature engineering yondashuvi O'zbek ta'lim ma'lumotlariga moslashtirilgan bo'lib, ilmiy adabiyotlarda hali to'liq yoritilmagan sohaga hissa qo'shadi. Bundan tashqari, platforma HEMIS ning tahliliy jihatdan kuchsiz bo'lgan tomonlarini to'ldiruvchi qo'shimcha tizim sifatida ham foydalanilishi mumkin.")
    pp(doc,"Tuzilishi va hajmi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Diplom loyihasi kirish, to'rtta asosiy bob, xulosa, foydalanilgan adabiyotlar ro'yxati va ilovalardan iborat. Umumiy hajm 80 sahifadan ortiq. I bob nazariy asoslarni (15 sahifa), II bob loyihalash va modellashtirish masalalarini (17 sahifa), III bob platformani ishlab chiqish va sinovlarni (22 sahifa), IV bob esa hayot faoliyati xavfsizligi masalalarini (7 sahifa) o'z ichiga oladi. Hujjatga 21 ta rasm, 17 ta jadval va dastur kodi ilovalari kiritilgan. Foydalanilgan adabiyotlar ro'yxatida 25 ta manba (qonunlar, monografiyalar, ilmiy maqolalar va web manbalar) keltirilgan.")
    pp(doc,"Tadqiqot predmeti va ob'ekti: tadqiqot ob'ekti — oliy ta'lim muassasasida talabalar o'zlashtirish jarayoni va unga ta'sir etuvchi omillar. Tadqiqot predmeti — ushbu jarayonni grafik tahlil qilish va bashorat qilish uchun web platforma ishlab chiqish metodlari va vositalari. Tadqiqot chegaralari: platforma O'zbekiston oliy ta'lim muassasalari uchun mo'ljallangan; ma'lumotlar sifatida 100-ballik baholash tizimi va kredit-modul tizimi qo'llaniladi; ML modeli ikkilik klassifikatsiya muammosini hal qiladi; frontend va backend o'zaro REST API orqali muloqot qiladi.")

