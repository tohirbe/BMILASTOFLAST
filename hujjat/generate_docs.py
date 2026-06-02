#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Platforma foydalanuvchi qo'llanmasi — screenshot + tavsif"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SS  = Path("hujjat/role_screenshots")
OUT = Path("hujjat/dokumentatsiya.docx")

# ── Rang konstantalari ──
RED    = RGBColor(0xDC, 0x26, 0x26)   # admin
BLUE   = RGBColor(0x25, 0x63, 0xEB)   # dekanat
GREEN  = RGBColor(0x05, 0x96, 0x69)   # o'qituvchi
PURPLE = RGBColor(0x70, 0x3B, 0xEE)   # talaba
GRAY   = RGBColor(0x47, 0x55, 0x69)   # API

# ── SAHIFA MA'LUMOTLARI ──
PAGES = {
  "login": {
    "title": "Tizimga kirish (Login)",
    "folder": "login",
    "file": "login.png",
    "role_color": GRAY,
    "desc": "Platforma ishga tushganda foydalanuvchi ushbu sahifani ko'radi. Username va parol kiritilgach, tizim avtomatik ravishda foydalanuvchi roliga mos dashboardga yo'naltiradi.",
    "funcs": [
      "Username va parol kiritish maydonlari",
      "Parolni ko'rsatish/yashirish (ko'z ikonasi)",
      "«Kirish» tugmasi — JWT token olish",
      "Xato holatda qizil xabar ko'rsatiladi",
      "Muvaffaqiyatli kirishda rol asosida sahifaga yo'naltirish",
    ]
  },
  "api_swagger": {
    "title": "API Hujjatlar — Swagger UI",
    "folder": "api",
    "file": "swagger_docs.png",
    "role_color": GRAY,
    "desc": "Backend FastAPI tomonidan avtomatik yaratilgan interaktiv API hujjatlari. Har bir endpoint ni to'g'ridan-to'g'ri shu sahifadan sinab ko'rish mumkin.",
    "funcs": [
      "Barcha API endpoint larni ko'rish (GET / POST / PUT / DELETE)",
      "Har bir endpoint ni «Try it out» orqali sinash",
      "Request/Response sxemalarini ko'rish",
      "JWT token kiritib autentifikatsiya qilish",
      "Endpoint larni toifa bo'yicha ko'rish (auth, grades, attendance...)",
    ]
  },

  # ─────────────────── ADMIN ───────────────────
  "admin_dashboard": {
    "title": "ADMIN — Dashboard",
    "folder": "admin", "file": "01_dashboard.png",
    "role_color": RED,
    "desc": "Adminning asosiy boshqaruv paneli. Tizim bo'yicha umumiy raqamli ko'rsatkichlar va grafiklar real vaqtda yangilanib turadi.",
    "funcs": [
      "Jami talabalar, faol foydalanuvchilar, xavf ostidagilar soni (metrika kartochkalari)",
      "Tizim bo'yicha o'rtacha GPA ko'rsatkichi",
      "So'nggi 6 oy o'zlashtirish tendensiyasi (chiziqli grafik)",
      "Guruhlar bo'yicha o'rtacha ball (ustunli diagramma)",
      "Baholash holati taqsimoti (doira diagrammasi: a'lo / yaxshi / qoniqarli / qoniqarsiz)",
    ]
  },
  "admin_talabalar": {
    "title": "ADMIN — Talabalar ro'yxati",
    "folder": "admin", "file": "02_talabalar.png",
    "role_color": RED,
    "desc": "Tizimda ro'yxatdan o'tgan barcha talabalar jadvali. Qidiruv, filtr va to'liq CRUD imkoniyatlari mavjud.",
    "funcs": [
      "Ism / talaba ID bo'yicha qidirish",
      "Guruh, kurs, fakultet bo'yicha filtrash",
      "Yangi talaba qo'shish (modal forma)",
      "Talaba ma'lumotlarini tahrirlash",
      "Talabani o'chirish (tasdiqlash oynasi bilan)",
      "Talaba ismiga bosib batafsil profil sahifasiga o'tish",
      "Sahifalash (pagination)",
    ]
  },
  "admin_talaba_profil": {
    "title": "ADMIN — Talaba profili (batafsil)",
    "folder": "admin", "file": "03_talaba_profili.png",
    "role_color": RED,
    "desc": "Bitta talabaning to'liq akademik portfeli. Barcha semestrlar, fanlar, baholar va davomat ma'lumotlari bir joyda.",
    "funcs": [
      "Shaxsiy ma'lumotlar: ism, guruh, talaba ID",
      "GPA dinamikasi — semestrlar bo'yicha chiziqli grafik",
      "Barcha fanlar bo'yicha JN / ON / YN / Jami ballar jadvali",
      "Davomat foizi va darsga qatnashish statistikasi",
      "Aktiv qarzdorliklar ro'yxati",
      "ML xavf darajasi ko'rsatkichi (foiz va rang indikatori)",
    ]
  },
  "admin_fanlar": {
    "title": "ADMIN — Fanlar",
    "folder": "admin", "file": "04_fanlar.png",
    "role_color": RED,
    "desc": "O'quv fanlarini boshqarish sahifasi. Har bir fan kredit soni, semestr va biriktirilgan o'qituvchilar bilan ko'rsatiladi.",
    "funcs": [
      "Fanlar ro'yxatini ko'rish (nomi, kodi, kreditlar, semestr)",
      "Yangi fan qo'shish",
      "Fan ma'lumotlarini tahrirlash",
      "Fanni o'chirish",
      "Fanga o'qituvchi va guruh biriktirish",
      "Fan bo'yicha o'rtacha ball statistikasi",
    ]
  },
  "admin_guruhlar": {
    "title": "ADMIN — Guruhlar",
    "folder": "admin", "file": "05_guruhlar.png",
    "role_color": RED,
    "desc": "O'quv guruhlarini boshqarish. Har bir guruh talabalar soni, kurs va fakultet bilan ko'rsatiladi.",
    "funcs": [
      "Guruhlar ro'yxatini ko'rish",
      "Yangi guruh yaratish",
      "Guruh nomini / kursini tahrirlash",
      "Guruhni o'chirish",
      "Guruh bo'yicha o'rtacha GPA ko'rsatkichi",
      "Guruh tarkibiga kirish (talabalar ro'yxati)",
    ]
  },
  "admin_baholar": {
    "title": "ADMIN — Baho kiritish",
    "folder": "admin", "file": "06_baholar.png",
    "role_color": RED,
    "desc": "Guruh va fan tanlangandan so'ng talabalar ro'yxati chiqadi. Har bir talaba uchun JN, ON, YN ballarini kiritish yoki tahrirlash mumkin.",
    "funcs": [
      "Guruh va fan bo'yicha filtrash",
      "Semestr tanlash",
      "JN (0–40), ON (0–20), YN (0–40) ballarini kiritish",
      "Jami ball avtomatik hisoblanadi (JN+ON+YN)",
      "O'tdi / Qoniqarsiz holati rang bilan belgilanadi",
      "Baholash oynasi yopiq bo'lsa kiritish bloklanadi",
      "Ommaviy saqlash (bir tugma bilan barcha baholar)",
    ]
  },
  "admin_grade_windows": {
    "title": "ADMIN — Baholash oynalari",
    "folder": "admin", "file": "07_baholash_oynalari.png",
    "role_color": RED,
    "desc": "Baholash muddatlarini boshqarish. JN, ON va YN uchun alohida oynalar ochiladi va yopiladi. Oyna yopiq bo'lsa o'qituvchilar baho kira olmaydi.",
    "funcs": [
      "Joriy semestr baholash oynalarini ko'rish",
      "JN / ON / YN uchun boshlanish va tugash sanasini belgilash",
      "Oynani ochish / yopish (toggle)",
      "Oyna holati: Ochiq (yashil) / Yopiq (qizil) ko'rsatiladi",
      "Semestr bo'yicha filtrash",
    ]
  },
  "admin_davomat": {
    "title": "ADMIN — Davomat",
    "folder": "admin", "file": "08_davomat.png",
    "role_color": RED,
    "desc": "Barcha guruhlar va fanlar bo'yicha davomat ma'lumotlari. Admin ko'rish va tahrirlash huquqiga ega.",
    "funcs": [
      "Guruh, fan, sana bo'yicha davomat ko'rish",
      "Talabalar ro'yxatida: Keldi / Kech keldi / Kelmadi holati",
      "Davomat foizi statistikasi — chiziqli grafik",
      "60% dan past davomatli talabalar avtomatik ajratib ko'rsatiladi",
      "Davomat yozuvlarini tahrirlash",
      "Sana oralig'i bo'yicha filtrash",
    ]
  },
  "admin_qarzdorlik": {
    "title": "ADMIN — Qarzdorliklar",
    "folder": "admin", "file": "09_qarzdorliklar.png",
    "role_color": RED,
    "desc": "Barcha talabalarning akademik qarzdorliklari ro'yxati. Holat: Kutilmoqda / Hal qilindi / Bekor qilindi.",
    "funcs": [
      "Barcha qarzdorliklarni ko'rish (talaba, fan, ball, sana)",
      "Holat bo'yicha filtrash (pending / resolved / cancelled)",
      "Guruh va semestr bo'yicha filtrash",
      "Qarzdorlik holatini o'zgartirish",
      "Qayta imtihon sanasini belgilash",
      "Eng ko'p qarzdorlik uchratilgan fanlar statistikasi",
    ]
  },
  "admin_jadval": {
    "title": "ADMIN — Dars jadvali",
    "folder": "admin", "file": "10_dars_jadvali.png",
    "role_color": RED,
    "desc": "Haftalik dars jadvalini boshqarish. Barcha guruhlar va o'qituvchilar uchun yagona joyda.",
    "funcs": [
      "Haftalik jadval (Du—Yak) ko'rinishida ko'rish",
      "Yangi dars qo'shish: fan, guruh, o'qituvchi, xona, vaqt",
      "To'qnashuv tekshiruvi (xona/o'qituvchi bir vaqtda band bo'lsa ogohlantirish)",
      "Dars ma'lumotlarini tahrirlash va o'chirish",
      "Guruh bo'yicha filtrash",
      "PDF formatida eksport qilish",
    ]
  },
  "admin_teacher_perf": {
    "title": "ADMIN — O'qituvchi samaradorligi",
    "folder": "admin", "file": "11_oqituvchi_samaradorligi.png",
    "role_color": RED,
    "desc": "Barcha o'qituvchilar bo'yicha samaradorlik ko'rsatkichlari. Baholarni o'z vaqtida kiritish, o'rtacha GPA va boshqa metrikalar.",
    "funcs": [
      "O'qituvchilar reytingi jadvali",
      "O'rtacha GPA ko'rsatkichi (o'quvchilar natijalari asosida)",
      "Davomat belgilash foizi (qanchalik to'liq belgilagan)",
      "Baholar kiritish o'z vaqtiligi (%)",
      "Xavf ostidagi talabalar bilan ishlash ko'rsatkichi",
      "Semestr bo'yicha filtrash",
    ]
  },
  "admin_risk": {
    "title": "ADMIN — Xavf tahlili (ML)",
    "folder": "admin", "file": "12_xavf_tahlili.png",
    "role_color": RED,
    "desc": "Random Forest ML modeli asosida akademik xavf ostidagi talabalar aniqlanadi. Har bir talaba uchun xavf ehtimoli (%) ko'rsatiladi.",
    "funcs": [
      "Xavf ostidagi talabalar ro'yxati (ehtimollik % bilan)",
      "Xavf darajasi: Yuqori (qizil) / O'rta (sariq) / Past (yashil)",
      "Asosiy xavf omillari: davomat foizi, oldingi GPA, JN ballari",
      "Talaba profiliga o'tish (bir bosish)",
      "Semestr va guruh bo'yicha filtrash",
      "Tavsiyalar: qaysi talabaga qanday yordam kerak",
    ]
  },
  "admin_hisobotlar": {
    "title": "ADMIN — Hisobotlar",
    "folder": "admin", "file": "13_hisobotlar.png",
    "role_color": RED,
    "desc": "Ma'lumotlarni PDF va Excel formatida yuklab olish. Guruh, semestr yoki individual talaba bo'yicha hisobotlar.",
    "funcs": [
      "Guruh + semestr tanlash",
      "Hisobot turi: Baholar / Davomat / Qarzdorlik / Umumiy",
      "PDF formatida professional hujjat sifatida eksport",
      "Excel (XLSX) formatida ma'lumotlar jadvali eksport",
      "Eksport tugmasi bosilganda yuklab olish boshlanadi",
    ]
  },
  "admin_upload": {
    "title": "ADMIN — Ma'lumot yuklash",
    "folder": "admin", "file": "14_malumot_yuklash.png",
    "role_color": RED,
    "desc": "Excel yoki CSV fayl orqali ommaviy ma'lumot import qilish. Talabalar, baholar yoki davomat ma'lumotlari mass yuklash uchun.",
    "funcs": [
      "Fayl turini tanlash: Talabalar / Baholar / Davomat",
      "Drag & Drop yoki fayl tanlash tugmasi",
      "Namuna shablon faylni yuklab olish",
      "Yuklashdan oldin ma'lumotlar preview ko'rinishi",
      "Xatolar aniqlansa qator bo'yicha xabar ko'rsatiladi",
      "Muvaffaqiyatli yuklangandan so'ng statistika ko'rsatiladi",
    ]
  },
  "admin_users": {
    "title": "ADMIN — Foydalanuvchilar",
    "folder": "admin", "file": "15_foydalanuvchilar.png",
    "role_color": RED,
    "desc": "Tizim foydalanuvchilarini boshqarish. Yangi akkaunt yaratish, rol belgilash, faollashtirish/bloklash.",
    "funcs": [
      "Barcha foydalanuvchilar ro'yxati (ism, login, rol, holat)",
      "Yangi foydalanuvchi qo'shish (login, parol, rol, bog'liq talaba/o'qituvchi)",
      "Foydalanuvchi rolini o'zgartirish",
      "Akkauntni faollashtirish yoki bloklash (is_active toggle)",
      "Parolni tiklash",
      "Qidirish va rol bo'yicha filtrash",
    ]
  },
  "admin_settings": {
    "title": "ADMIN — Sozlamalar",
    "folder": "admin", "file": "16_sozlamalar.png",
    "role_color": RED,
    "desc": "Foydalanuvchi interfeys sozlamalari. Til, rang rejimi va qo'shimcha ko'rinish parametrlari.",
    "funcs": [
      "Interfeys tili: O'zbek / Русский / English",
      "Rang rejimi: Yorug' (Light) / Qorong'u (Dark)",
      "Compact mode: zichroq interfeys uchun",
      "Sozlamalar localStorage da saqlanadi",
    ]
  },

  # ─────────────────── DEKANAT ───────────────────
  "dek_dashboard": {
    "title": "DEKANAT — Dashboard",
    "folder": "dekanat", "file": "01_dashboard.png",
    "role_color": BLUE,
    "desc": "Dekanat uchun umumiy ko'rinish. Admin dashboardi bilan o'xshash, lekin foydalanuvchi boshqaruvi va ma'lumot yuklash yo'q.",
    "funcs": [
      "Fakultet/bo'lim miqyosidagi statistika",
      "O'rtacha GPA, xavf ostidagilar soni",
      "Guruhlar va fanlar bo'yicha tendensiya grafiklar",
      "Tezkor havolalar: hisobotlar, xavf tahlili",
    ]
  },
  "dek_students": {
    "title": "DEKANAT — Talabalar",
    "folder": "dekanat", "file": "02_talabalar.png",
    "role_color": BLUE,
    "desc": "Barcha talabalarni ko'rish va filtrash. Dekanat yangi talaba qo'sha oladi lekin o'chirish huquqi yo'q.",
    "funcs": [
      "Barcha talabalarni ko'rish va qidirish",
      "Guruh va kurs bo'yicha filtrash",
      "Talaba profilini ko'rish (baholar, davomat)",
      "Talaba ma'lumotlarini ko'rish (tahrirlash cheklangan)",
    ]
  },
  "dek_fanlar": {
    "title": "DEKANAT — Fanlar",
    "folder": "dekanat", "file": "03_fanlar.png",
    "role_color": BLUE,
    "desc": "Fanlar ro'yxati va statistikasi. Dekanat yangi fan qo'sha oladi.",
    "funcs": [
      "Fanlar ro'yxatini ko'rish",
      "Fan bo'yicha o'rtacha ball statistikasi",
      "Yangi fan qo'shish",
      "Fan ma'lumotlarini tahrirlash",
    ]
  },
  "dek_guruhlar": {
    "title": "DEKANAT — Guruhlar",
    "folder": "dekanat", "file": "04_guruhlar.png",
    "role_color": BLUE,
    "desc": "Guruhlar ro'yxati va ularning umumiy ko'rsatkichlari.",
    "funcs": [
      "Barcha guruhlarni ko'rish",
      "Guruh bo'yicha o'rtacha GPA",
      "Guruh tarkibiga kirish",
      "Yangi guruh yaratish",
    ]
  },
  "dek_davomat": {
    "title": "DEKANAT — Davomat",
    "folder": "dekanat", "file": "05_davomat.png",
    "role_color": BLUE,
    "desc": "Barcha guruhlar davomat statistikasi. Dekanat davomat ko'ra oladi, lekin o'zgartira olmaydi.",
    "funcs": [
      "Guruh va fan bo'yicha davomat statistikasini ko'rish",
      "Past davomatli talabalar ro'yxati",
      "Davomat foizi tendensiyasi grafigi",
      "Sana oralig'i bo'yicha filtrash",
    ]
  },
  "dek_qarzdorlik": {
    "title": "DEKANAT — Qarzdorliklar",
    "folder": "dekanat", "file": "06_qarzdorliklar.png",
    "role_color": BLUE,
    "desc": "Barcha guruhlar bo'yicha qarzdorliklar holati. Dekanat holatni o'zgartira oladi.",
    "funcs": [
      "Barcha qarzdorliklarni ko'rish va boshqarish",
      "Holat o'zgartirish: Hal qilindi / Bekor qilindi",
      "Qayta imtihon sanasini belgilash",
      "Guruh va semestr bo'yicha filtrash",
    ]
  },
  "dek_jadval": {
    "title": "DEKANAT — Dars jadvali",
    "folder": "dekanat", "file": "07_dars_jadvali.png",
    "role_color": BLUE,
    "desc": "Barcha guruhlar uchun dars jadvali. Dekanat jadvalga kirish va o'zgartirish huquqiga ega.",
    "funcs": [
      "Barcha guruhlar jadvali",
      "Guruh bo'yicha filtrash",
      "Jadval to'qnashuvlarini ko'rish",
      "Yangi dars qo'shish / tahrirlash",
    ]
  },
  "dek_teacher_perf": {
    "title": "DEKANAT — O'qituvchilar samaradorligi",
    "folder": "dekanat", "file": "08_oqituvchi_samaradorligi.png",
    "role_color": BLUE,
    "desc": "Adminnikidan bir xil — barcha o'qituvchilar reytingi va samaradorlik ko'rsatkichlari.",
    "funcs": [
      "O'qituvchilar reytingi (GPA, davomat, o'z vaqtlilik)",
      "Semestr bo'yicha tahlil",
      "Eng samarali / eng kam samarali o'qituvchilar",
    ]
  },
  "dek_risk": {
    "title": "DEKANAT — Xavf tahlili",
    "folder": "dekanat", "file": "09_xavf_tahlili.png",
    "role_color": BLUE,
    "desc": "ML asosida xavf ostidagi talabalar. Dekanat bu sahifaga to'liq kirish huquqiga ega.",
    "funcs": [
      "Xavf ostidagi talabalar ro'yxati va ehtimollik %",
      "Guruh bo'yicha filtrash",
      "Talaba profiliga o'tish",
    ]
  },
  "dek_hisobotlar": {
    "title": "DEKANAT — Hisobotlar",
    "folder": "dekanat", "file": "10_hisobotlar.png",
    "role_color": BLUE,
    "desc": "Adminnikidan bir xil — PDF va Excel eksport imkoniyatlari.",
    "funcs": [
      "Guruh + semestr tanlash",
      "PDF / Excel eksport",
      "Barcha hisobot turlari (baholar, davomat, qarzdorlik)",
    ]
  },
  "dek_grade_windows": {
    "title": "DEKANAT — Baholash oynalari",
    "folder": "dekanat", "file": "11_baholash_oynalari.png",
    "role_color": BLUE,
    "desc": "JN/ON/YN baholash oynalarini boshqarish. Dekanat ochish/yopish huquqiga ega.",
    "funcs": [
      "Oynalar holati: Ochiq / Yopiq",
      "Oyna ochish / yopish",
      "Sana chegaralarini belgilash",
    ]
  },
  "dek_settings": {
    "title": "DEKANAT — Sozlamalar",
    "folder": "dekanat", "file": "12_sozlamalar.png",
    "role_color": BLUE,
    "desc": "Til va tema sozlamalari — barcha rollarda bir xil.",
    "funcs": ["Til tanlash (UZ/RU/EN)", "Rang rejimi (Light/Dark)", "Compact mode"]
  },

  # ─────────────────── O'QITUVCHI ───────────────────
  "oq_dashboard": {
    "title": "O'QITUVCHI — Dashboard",
    "folder": "oqituvchi", "file": "01_dashboard.png",
    "role_color": GREEN,
    "desc": "O'qituvchi faqat o'ziga biriktirilgan guruhlar statistikasini ko'radi. Boshqa guruhlar ko'rinmaydi.",
    "funcs": [
      "O'z guruhlari bo'yicha o'rtacha GPA",
      "O'z darslarida davomat foizi",
      "Xavf ostidagi talabalar (faqat o'z guruhlari)",
      "Eng kam ball olganlar ro'yxati",
      "Haftalik dars jadvali preview",
    ]
  },
  "oq_talabalar": {
    "title": "O'QITUVCHI — O'z talabalari",
    "folder": "oqituvchi", "file": "02_talabalar.png",
    "role_color": GREEN,
    "desc": "O'qituvchi faqat o'ziga biriktirilgan guruhlar talabalarini ko'ra oladi. Boshqa guruhlar ko'rinmaydi.",
    "funcs": [
      "O'z guruhlari talabalar ro'yxati",
      "Talaba bo'yicha qidirish",
      "Talaba profilini ko'rish (batafsil baholar va davomat)",
      "Boshqa guruh talabalariga kirish bloklangan",
    ]
  },
  "oq_baholar": {
    "title": "O'QITUVCHI — Baho kiritish",
    "folder": "oqituvchi", "file": "03_baho_kiritish.png",
    "role_color": GREEN,
    "desc": "O'qituvchi faqat o'z fanlariga baho kirita oladi. Boshqa fanlar ro'yxatda ko'rinmaydi.",
    "funcs": [
      "Faqat o'z fanlari bo'yicha guruh tanlash",
      "JN (0–40), ON (0–20), YN (0–40) kiritish",
      "Jami ball avtomatik hisoblanadi",
      "Baholash oynasi yopiq bo'lsa kiritish bloklanadi",
      "Oldindan kiritilgan baholarni tahrirlash",
    ]
  },
  "oq_davomat": {
    "title": "O'QITUVCHI — Davomat belgilash",
    "folder": "oqituvchi", "file": "04_davomat.png",
    "role_color": GREEN,
    "desc": "Dars boshida o'qituvchi har bir talaba uchun davomat belgilaydi. Faqat o'z guruh va fanlari uchun.",
    "funcs": [
      "Guruh va fan tanlash (faqat o'ziniki)",
      "Sana tanlash (standart: bugun)",
      "Har talaba uchun: Keldi / Kech keldi / Kelmadi",
      "«Hammasi keldi» ommaviy belgilash tugmasi",
      "Davomat statistikasi grafigi (semestr bo'yicha)",
      "Bir saqlash tugmasi bilan barchasi bazaga yoziladi",
    ]
  },
  "oq_qarzdorlik": {
    "title": "O'QITUVCHI — Qarzdorliklar",
    "folder": "oqituvchi", "file": "05_qarzdorliklar.png",
    "role_color": GREEN,
    "desc": "O'z guruhlaridagi talabalarning qarzdorliklari. O'qituvchi holatni o'zgartira oladi.",
    "funcs": [
      "O'z guruhi bo'yicha qarzdorliklar ro'yxati",
      "Qayta imtihon natijasini kiritish",
      "Holat o'zgartirish: Hal qilindi / Bekor qilindi",
      "Talaba profiliga o'tish",
    ]
  },
  "oq_jadval": {
    "title": "O'QITUVCHI — Dars jadvali",
    "folder": "oqituvchi", "file": "06_dars_jadvali.png",
    "role_color": GREEN,
    "desc": "Faqat o'z dars jadvali ko'rsatiladi. O'qituvchi jadvalni o'zgartira olmaydi (faqat ko'radi).",
    "funcs": [
      "Haftalik jadval — faqat o'z darslari",
      "Dars vaqti, xona, guruh ma'lumotlari",
      "Hafta bo'yicha navigatsiya",
      "Jadval PDF eksport (ixtiyoriy)",
    ]
  },
  "oq_samaradorlik": {
    "title": "O'QITUVCHI — Mening samaradorligim",
    "folder": "oqituvchi", "file": "07_mening_samaradorligim.png",
    "role_color": GREEN,
    "desc": "O'qituvchi o'z ko'rsatkichlarini ko'radi — baholar kiritish timeliness, talabalar o'rtacha GPA, davomat.",
    "funcs": [
      "O'z guruhlari bo'yicha o'rtacha GPA",
      "Baholar kiritish o'z vaqtiligi (%)",
      "Davomat belgilash to'liqligi (%)",
      "Xavf ostidagi talabalar tendensiyasi",
      "Semestr bo'yicha taqqoslama grafik",
    ]
  },
  "oq_settings": {
    "title": "O'QITUVCHI — Sozlamalar",
    "folder": "oqituvchi", "file": "08_sozlamalar.png",
    "role_color": GREEN,
    "desc": "Til va tema sozlamalari.",
    "funcs": ["Til tanlash (UZ/RU/EN)", "Rang rejimi (Light/Dark)", "Compact mode"]
  },

  # ─────────────────── TALABA ───────────────────
  "tal_dashboard": {
    "title": "TALABA — Dashboard",
    "folder": "talaba", "file": "01_dashboard.png",
    "role_color": PURPLE,
    "desc": "Talabaning shaxsiy boshqaruv paneli. Faqat o'z ma'lumotlari ko'rsatiladi — boshqa hech kim ma'lumoti ko'rinmaydi.",
    "funcs": [
      "Joriy semestr baholarim (karta ko'rinishida)",
      "Davomat foizim (doira diagrammasi)",
      "O'rtacha GPA dinamikasi (chiziqli grafik)",
      "Aktiv qarzdorliklarim (agar mavjud bo'lsa)",
      "Bugungi dars jadvalim",
    ]
  },
  "tal_profil": {
    "title": "TALABA — Profilim",
    "folder": "talaba", "file": "02_profilim.png",
    "role_color": PURPLE,
    "desc": "Shaxsiy akademik portfolio. Barcha semestrlar bo'yicha baholar va GPA dinamikasi.",
    "funcs": [
      "Shaxsiy ma'lumotlar: ism, guruh, talaba ID",
      "Kümülatif GPA (barcha semestrlar)",
      "Semestrlar bo'yicha GPA grafigi",
      "Barcha fanlar baholari jadvali (JN/ON/YN/Jami)",
      "O'tgan va qoniqarsiz baholar rang bilan ajratilgan",
    ]
  },
  "tal_davomat": {
    "title": "TALABA — Davomatim",
    "folder": "talaba", "file": "03_davomatim.png",
    "role_color": PURPLE,
    "desc": "Talaba o'z dars davomati statistikasini ko'radi. Faqat o'qish — o'zgartirish mumkin emas.",
    "funcs": [
      "Umumiy davomat foizi (doira diagrammasi)",
      "Fan bo'yicha davomat foizlari",
      "Keldi / Kech keldi / Kelmadi kunlar ro'yxati",
      "Semestr bo'yicha davomat tendensiyasi grafigi",
      "60% dan past bo'lsa ogohlantirish xabari",
    ]
  },
  "tal_qarzdorlik": {
    "title": "TALABA — Qarzdorliklarim",
    "folder": "talaba", "file": "04_qarzdorliklarim.png",
    "role_color": PURPLE,
    "desc": "Faqat o'z qarzdorliklari ko'rsatiladi. Boshqa talabalar ma'lumotlari ko'rinmaydi.",
    "funcs": [
      "Aktiv qarzdorliklarim: fan nomi, ball, muddat",
      "Holat: Kutilmoqda / Hal qilindi / Bekor qilindi",
      "Muddati yaqinlashganda ogohlantirish",
      "Qarzdorlik tarixi (o'tgan semestrlar)",
    ]
  },
  "tal_jadval": {
    "title": "TALABA — Dars jadvali",
    "folder": "talaba", "file": "05_dars_jadvali.png",
    "role_color": PURPLE,
    "desc": "O'z guruhining haftalik dars jadvali. Faqat ko'rish — o'zgartirish mumkin emas.",
    "funcs": [
      "Haftalik jadval: kun, vaqt, fan, xona",
      "Bugungi darslar ajratib ko'rsatiladi",
      "O'qituvchi ismi har dars yonida",
      "Hafta bo'yicha navigatsiya",
    ]
  },
  "tal_settings": {
    "title": "TALABA — Sozlamalar",
    "folder": "talaba", "file": "06_sozlamalar.png",
    "role_color": PURPLE,
    "desc": "Til va tema sozlamalari — faqat o'z interfeysi uchun.",
    "funcs": ["Til tanlash (UZ/RU/EN)", "Rang rejimi (Light/Dark)", "Compact mode"]
  },
}

ORDER = [
  "login","api_swagger",
  "admin_dashboard","admin_talabalar","admin_talaba_profil","admin_fanlar",
  "admin_guruhlar","admin_baholar","admin_grade_windows","admin_davomat",
  "admin_qarzdorlik","admin_jadval","admin_teacher_perf","admin_risk",
  "admin_hisobotlar","admin_upload","admin_users","admin_settings",
  "dek_dashboard","dek_students","dek_fanlar","dek_guruhlar","dek_davomat",
  "dek_qarzdorlik","dek_jadval","dek_teacher_perf","dek_risk",
  "dek_hisobotlar","dek_grade_windows","dek_settings",
  "oq_dashboard","oq_talabalar","oq_baholar","oq_davomat","oq_qarzdorlik",
  "oq_jadval","oq_samaradorlik","oq_settings",
  "tal_dashboard","tal_profil","tal_davomat","tal_qarzdorlik",
  "tal_jadval","tal_settings",
]

ROLE_BANNERS = {
  "admin":    ("🔴 ADMINISTRATOR", RED,    "admin",    "admin123"),
  "dekanat":  ("🟡 DEKANAT",       BLUE,   "dekanat",  "dekan123"),
  "oqituvchi":("🟢 O'QITUVCHI",   GREEN,  "oqituvchi","teacher123"),
  "talaba":   ("🔵 TALABA",        PURPLE, "talaba",   "student123"),
}

def cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=Cm(2); sec.right_margin=Cm(1.5)
        sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def pp(doc, text, bold=False, italic=False, center=False, sz=11,
       color=None, after=0, before=0, first=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if center:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if first:
        para.paragraph_format.first_line_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = 'Segoe UI'; run.font.size = Pt(sz)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
    return para

def add_page_num(doc):
    sec = doc.sections[-1]; footer = sec.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.name = 'Segoe UI'; run.font.size = Pt(10)
    for tag, txt in [('begin',''),('end','')]: pass
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve'); ins.text='PAGE'
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'end')
    run._r.append(f1); run._r.append(ins); run._r.append(f2)

def role_banner(doc, label, color):
    """Rol bo'lim sarlavhasi — rangli blok"""
    para = doc.add_paragraph()
    para.paragraph_format.page_break_before = True
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(16)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"  {label}  ")
    run.font.name = 'Segoe UI'; run.font.size = Pt(22)
    run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    # Background via shading on paragraph
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_c = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_c); pPr.append(shd)

def add_page_entry(doc, key, data, idx):
    """Bitta sahifa: screenshot + tavsif jadval ko'rinishida"""
    img_path = SS / data['folder'] / data['file']
    color = data['role_color']
    hex_c = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

    # ── Sarlavha ──
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(f"  {idx}. {data['title']}")
    run.font.name = 'Segoe UI'; run.font.size = Pt(13)
    run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_c); pPr.append(shd)

    # ── Jadval: [Screenshot] | [Tavsif] ──
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(9.5)
    tbl.columns[1].width = Cm(7.5)

    # Sol: Screenshot
    left_cell = tbl.rows[0].cells[0]
    left_cell.width = Cm(9.5)
    img_para = left_cell.paragraphs[0]
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(4)
    img_para.paragraph_format.space_after  = Pt(4)
    if img_path.exists():
        img_para.add_run().add_picture(str(img_path), width=Cm(9))
    else:
        run = img_para.add_run(f"[{data['file']} — topilmadi]")
        run.font.color.rgb = RGBColor(0xFF,0x8C,0x00)
        run.font.bold = True

    # O'ng: Tavsif
    right_cell = tbl.rows[0].cells[1]
    right_cell.width = Cm(7.5)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # URL
    url_para = right_cell.add_paragraph()
    url_para.paragraph_format.space_before = Pt(6)
    url_para.paragraph_format.space_after  = Pt(4)
    url_run = url_para.add_run(f"localhost:5176{_get_path(data['folder'], data['file'])}")
    url_run.font.name = 'Courier New'; url_run.font.size = Pt(9)
    url_run.font.color.rgb = RGBColor(0x64,0x74,0x8B)
    url_run.font.italic = True

    # Tavsif matni
    desc_para = right_cell.add_paragraph()
    desc_para.paragraph_format.space_before = Pt(2)
    desc_para.paragraph_format.space_after  = Pt(8)
    desc_run = desc_para.add_run(data['desc'])
    desc_run.font.name = 'Segoe UI'; desc_run.font.size = Pt(10)

    # Funksiyalar sarlavhasi
    fh = right_cell.add_paragraph()
    fh.paragraph_format.space_before = Pt(0)
    fh.paragraph_format.space_after  = Pt(3)
    fr = fh.add_run("Funksiyalar:")
    fr.font.name = 'Segoe UI'; fr.font.size = Pt(10)
    fr.font.bold = True; fr.font.color.rgb = color

    # Funksiyalar ro'yxati
    for func in data['funcs']:
        fp = right_cell.add_paragraph()
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after  = Pt(2)
        fp.paragraph_format.left_indent  = Cm(0.3)
        fr2 = fp.add_run(f"• {func}")
        fr2.font.name = 'Segoe UI'; fr2.font.size = Pt(10)

    # Fayl nomi (pastda)
    fn_para = right_cell.add_paragraph()
    fn_para.paragraph_format.space_before = Pt(6)
    fn_para.paragraph_format.space_after  = Pt(4)
    fn_run = fn_para.add_run(f"📁 {data['folder']}/{data['file']}")
    fn_run.font.name = 'Segoe UI'; fn_run.font.size = Pt(9)
    fn_run.font.color.rgb = RGBColor(0x94,0xA3,0xB8)
    fn_run.font.italic = True

    # Bo'shliq
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

PATH_MAP = {
    "login":     {("login","login.png"): "/login"},
    "api":       {("api","swagger_docs.png"): "localhost:8000/docs"},
    "admin":     {
        "01_dashboard.png":"/", "02_talabalar.png":"/students",
        "03_talaba_profili.png":"/students/1", "04_fanlar.png":"/subjects",
        "05_guruhlar.png":"/groups", "06_baholar.png":"/grades",
        "07_baholash_oynalari.png":"/grade-windows", "08_davomat.png":"/attendance",
        "09_qarzdorliklar.png":"/debts", "10_dars_jadvali.png":"/schedule",
        "11_oqituvchi_samaradorligi.png":"/teacher-performance",
        "12_xavf_tahlili.png":"/risk", "13_hisobotlar.png":"/reports",
        "14_malumot_yuklash.png":"/upload", "15_foydalanuvchilar.png":"/users",
        "16_sozlamalar.png":"/settings",
    },
    "dekanat": {
        "01_dashboard.png":"/","02_talabalar.png":"/students",
        "03_fanlar.png":"/subjects","04_guruhlar.png":"/groups",
        "05_davomat.png":"/attendance","06_qarzdorliklar.png":"/debts",
        "07_dars_jadvali.png":"/schedule",
        "08_oqituvchi_samaradorligi.png":"/teacher-performance",
        "09_xavf_tahlili.png":"/risk","10_hisobotlar.png":"/reports",
        "11_baholash_oynalari.png":"/grade-windows","12_sozlamalar.png":"/settings",
    },
    "oqituvchi": {
        "01_dashboard.png":"/","02_talabalar.png":"/students",
        "03_baho_kiritish.png":"/grades","04_davomat.png":"/attendance",
        "05_qarzdorliklar.png":"/debts","06_dars_jadvali.png":"/schedule",
        "07_mening_samaradorligim.png":"/teacher-performance",
        "08_sozlamalar.png":"/settings",
    },
    "talaba": {
        "01_dashboard.png":"/","02_profilim.png":"/profile",
        "03_davomatim.png":"/attendance","04_qarzdorliklarim.png":"/debts",
        "05_dars_jadvali.png":"/schedule","06_sozlamalar.png":"/settings",
    },
}

def _get_path(folder, file):
    m = PATH_MAP.get(folder, {})
    return m.get(file, "/")

def main():
    doc = new_doc()
    add_page_num(doc)

    # ── MUQOVA ──
    pp(doc,"BMI TALABALAR O'ZLASHTIRISH TIZIMI",bold=True,center=True,sz=20,color=RGBColor(0x1E,0x40,0xAF),after=8)
    pp(doc,"PLATFORMA FOYDALANUVCHI QO'LLANMASI",bold=True,center=True,sz=15,after=4)
    pp(doc,"Har bir sahifa tavsifi va funksiyalar ro'yxati",italic=True,center=True,sz=12,
       color=RGBColor(0x64,0x74,0x8B),after=6)

    # Kirish jadvali
    intro = doc.add_table(rows=5, cols=2)
    intro.style = 'Table Grid'
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    roles_info = [
        ("🔴 ADMIN","admin / admin123","15 sahifa — to'liq boshqaruv"),
        ("🟡 DEKANAT","dekanat / dekan123","12 sahifa — kuzatuv va hisobot"),
        ("🟢 O'QITUVCHI","oqituvchi / teacher123","8 sahifa — o'z guruhi"),
        ("🔵 TALABA","talaba / student123","6 sahifa — faqat o'z ma'lumotlari"),
    ]
    # header
    for ci,(h,_,__) in enumerate([("Rol","Login","Sahifalar soni")]):
        pass
    hdr = intro.rows[0]
    for ci, h in enumerate(["Rol","Login / Parol","Sahifalar"]):
        c = hdr.cells[ci] if ci < 2 else hdr.cells[1]
    # header row
    intro.rows[0].cells[0].paragraphs[0].add_run("Rol").font.bold=True
    intro.rows[0].cells[1].paragraphs[0].add_run("Login / Parol  |  Sahifalar soni").font.bold=True
    for ri,(role,cred,pages) in enumerate(roles_info,1):
        intro.rows[ri].cells[0].paragraphs[0].add_run(role)
        intro.rows[ri].cells[1].paragraphs[0].add_run(f"{cred}  |  {pages}")
    pp(doc,"",after=6)
    pp(doc,f"Jami sahifalar: {len(ORDER)} ta  |  Qo'llanma tuzilgan: 2026",
       italic=True,center=True,sz=10,color=RGBColor(0x94,0xA3,0xB8),after=16)

    # ── ASOSIY KONTENT ──
    prev_folder = None
    idx = 0
    for key in ORDER:
        data = PAGES[key]
        folder = data['folder']

        # Rol bo'lim sarlavhasi (birinchi marta)
        if folder != prev_folder:
            if folder == 'login':
                role_banner(doc, "UMUMIY SAHIFALAR", GRAY)
            elif folder == 'admin':
                role_banner(doc, "🔴  ADMINISTRATOR  (admin / admin123)", RED)
            elif folder == 'api':
                pass  # swagger API — o'sha blok ichida
            elif folder == 'dekanat':
                role_banner(doc, "🟡  DEKANAT  (dekanat / dekan123)", BLUE)
            elif folder == 'oqituvchi':
                role_banner(doc, "🟢  O'QITUVCHI  (oqituvchi / teacher123)", GREEN)
            elif folder == 'talaba':
                role_banner(doc, "🔵  TALABA  (talaba / student123)", PURPLE)
            prev_folder = folder

        idx += 1
        add_page_entry(doc, key, data, idx)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK! Dokumentatsiya saqlandi: {OUT}")
    print(f"   Sahifalar tavsifi: {len(ORDER)} ta")

if __name__ == "__main__":
    main()
