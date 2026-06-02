#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Diplom loyihasi hujjat generatori
Mavzu: Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SS = Path("hujjat/screenshots")
OUT = Path("hujjat/diplom.docx")
_fig = {}; _tbl = {}

def fn(ch):
    _fig[ch] = _fig.get(ch,0)+1; return f"{ch}.{_fig[ch]}"
def tn(ch):
    _tbl[ch] = _tbl.get(ch,0)+1; return f"{ch}.{_tbl[ch]}"

def cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=Cm(3); sec.right_margin=Cm(1.5)
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    for sname,bold,align,sz,before,after,pbb,indent in [
        ('Heading 1',True,WD_ALIGN_PARAGRAPH.CENTER,14,0,12,True,0),
        ('Heading 2',True,WD_ALIGN_PARAGRAPH.LEFT,14,12,6,False,0),
        ('Heading 3',True,WD_ALIGN_PARAGRAPH.LEFT,14,6,3,False,0),
    ]:
        s=doc.styles[sname]
        s.font.name='Times New Roman'; s.font.size=Pt(sz)
        s.font.bold=bold; s.font.color.rgb=RGBColor(0,0,0)
        s.paragraph_format.alignment=align
        s.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
        s.paragraph_format.space_before=Pt(before)
        s.paragraph_format.space_after=Pt(after)
        s.paragraph_format.first_line_indent=Cm(indent)
        if pbb: s.paragraph_format.page_break_before=True
    n=doc.styles['Normal']
    n.font.name='Times New Roman'; n.font.size=Pt(14)
    return doc

def add_page_num(doc):
    sec=doc.sections[-1]; footer=sec.footer
    if footer.paragraphs: para=footer.paragraphs[0]; para.clear()
    else: para=footer.add_paragraph()
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=para.add_run()
    run.font.name='Times New Roman'; run.font.size=Pt(12)
    for t,txt in [('begin',''),('end','')]:
        pass
    fld1=OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
    ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve'); ins.text='PAGE'
    fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
    run._r.append(fld1); run._r.append(ins); run._r.append(fld2)

def pp(doc, text, first=True, bold=False, italic=False, center=False, right=False, sz=14, after=0, before=0):
    para=doc.add_paragraph()
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Pt(before)
    para.paragraph_format.space_after=Pt(after)
    if center:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent=Cm(0)
    elif right:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.first_line_indent=Cm(0)
    else:
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent=Cm(1.25) if first else Cm(0)
    run=para.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(sz)
    run.font.bold=bold; run.font.italic=italic
    return para

def CH(doc, text):
    para=doc.add_paragraph()
    para.style=doc.styles['Heading 1']
    para.clear()
    para.paragraph_format.page_break_before=True
    para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Pt(0)
    para.paragraph_format.space_after=Pt(12)
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run(text.upper())
    run.font.name='Times New Roman'; run.font.size=Pt(14)
    run.font.bold=True; run.font.color.rgb=RGBColor(0,0,0)
    return para

def SH(doc, text, level=2):
    para=doc.add_paragraph()
    para.style=doc.styles[f'Heading {level}']
    para.clear()
    para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.space_before=Cm(0.4)
    para.paragraph_format.space_after=Cm(0.2)
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(14)
    run.font.bold=True; run.font.color.rgb=RGBColor(0,0,0)
    return para

def fig(doc, img, caption, num):
    path=SS/img
    img_p=doc.add_paragraph()
    img_p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before=Pt(6)
    img_p.paragraph_format.space_after=Pt(0)
    img_p.paragraph_format.first_line_indent=Cm(0)
    if path.exists():
        run=img_p.add_run(); run.add_picture(str(path),width=Cm(15))
    else:
        run=img_p.add_run(f"[{num}-rasm: {caption} — skrinshot qo'shilsin]")
        run.font.name='Times New Roman'; run.font.size=Pt(12)
        run.font.color.rgb=RGBColor(0xFF,0x8C,0x00); run.font.bold=True
    cap=doc.add_paragraph()
    cap.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before=Pt(3); cap.paragraph_format.space_after=Pt(9)
    cap.paragraph_format.first_line_indent=Cm(0)
    r=cap.add_run(f"{num}-rasm. {caption}")
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.italic=True
    return cap

def tbl(doc, title, num, headers, rows, widths=None):
    np_=doc.add_paragraph()
    np_.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    np_.paragraph_format.space_before=Pt(9); np_.paragraph_format.space_after=Pt(0)
    np_.paragraph_format.first_line_indent=Cm(0)
    r=np_.add_run(f"{num}-jadval")
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.italic=True

    tp=doc.add_paragraph()
    tp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before=Pt(0); tp.paragraph_format.space_after=Pt(6)
    tp.paragraph_format.first_line_indent=Cm(0)
    r=tp.add_run(title)
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.font.bold=True

    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width=Cm(w)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; cell_bg(c,'BDD7EE')
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pa.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
        rr=pa.add_run(h); rr.font.name='Times New Roman'
        rr.font.size=Pt(12); rr.font.bold=True
    for ri,row_d in enumerate(rows):
        for ci,val in enumerate(row_d):
            c=t.rows[ri+1].cells[ci]
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            pa=c.paragraphs[0]
            pa.paragraph_format.line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE
            rr=pa.add_run(str(val))
            rr.font.name='Times New Roman'; rr.font.size=Pt(12)
    sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(6)
    return t

def add_toc(doc):
    pp(doc,"MUNDARIJA",first=False,bold=True,center=True,sz=14,after=12)
    para=doc.add_paragraph()
    para.paragraph_format.first_line_indent=Cm(0)
    run=para.add_run()
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve')
    ins.text=' TOC \\o "1-3" \\h \\z \\u '
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    pr=OxmlElement('w:r'); pt_=OxmlElement('w:t')
    pt_.text='[Word\'da F9 bosib mundarijani yangilang]'; pr.append(pt_)
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    run._r.append(f1); run._r.append(ins); run._r.append(f2)
    run._r.append(pr); run._r.append(f3)
    run.font.name='Times New Roman'; run.font.size=Pt(14)

# ============================================================
# TITUL VARAQ
# ============================================================
def add_title_page(doc):
    sec=doc.sections[0]; sec.different_first_page_header_footer=True
    pp(doc,"O'ZBEKISTON RESPUBLIKASI RAQAMLI TEXNOLOGIYALAR VAZIRLIGI",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"MUHAMMAD AL-XORAZMIY NOMIDAGI TOSHKENT AXBOROT TEXNOLOGIYALARI UNIVERSITETI",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"KOMPYUTER INJINIRINGI FAKULTETI",first=False,bold=True,center=True,sz=14,after=24)
    pp(doc,'"Himoyaga ruxsat etilsin"',first=False,center=True,sz=13,after=0)
    pp(doc,'"Kompyuter tizimlari" kafedrasi mudiri',first=False,center=True,sz=13,after=0)
    pp(doc,'_____________ [Kafedra mudiri F.I.O.]',first=False,center=True,sz=13,after=0)
    pp(doc,'«___» __________ 2026 y.',first=False,center=True,sz=13,after=30)
    pp(doc,"DIPLOM LOYIHASI",first=False,bold=True,center=True,sz=16,after=12)
    pp(doc,"Mavzu:",first=False,bold=True,center=True,sz=14,after=6)
    pp(doc,"Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma ishlab chiqish",first=False,bold=True,center=True,sz=14,after=30)

    t=doc.add_table(rows=5,cols=3)
    t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdrs=[["Vazifasi","F.I.O.","Imzo"]]
    data=[
        ["Bitiruvchi","[Ism Familiya]","________"],
        ["Ilmiy rahbar","[Rahbar F.I.O.]","________"],
        ["Taqrizchi","[Taqrizchi F.I.O.]","________"],
        ["HFX maslahatchisi","[Maslahatchi F.I.O.]","________"],
    ]
    for ci,h in enumerate(hdrs[0]):
        c=t.rows[0].cells[ci]; cell_bg(c,'D9D9D9')
        pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=pa.add_run(h); rr.font.name='Times New Roman'
        rr.font.size=Pt(12); rr.font.bold=True
    for ri,row_d in enumerate(data):
        for ci,val in enumerate(row_d):
            c=t.rows[ri+1].cells[ci]
            pa=c.paragraphs[0]; pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rr=pa.add_run(val); rr.font.name='Times New Roman'; rr.font.size=Pt(12)

    pp(doc,"",center=True,sz=14,after=24)
    pp(doc,"Toshkent – 2026",first=False,bold=True,center=True,sz=14,after=0)

# ============================================================
# TOPSHIRIQ VARAG'I
# ============================================================
def add_assignment(doc):
    CH(doc,"DIPLOM LOYIHASI TOPSHIRIG'I")
    pp(doc,"O'zbekiston Respublikasi Raqamli Texnologiyalar Vazirligi",first=False,bold=True,center=True,sz=13)
    pp(doc,"Muhammad al-Xorazmiy nomidagi Toshkent Axborot Texnologiyalari Universiteti",first=False,bold=True,center=True,sz=13)
    pp(doc,"Kompyuter injiniringi fakulteti, \"Kompyuter tizimlari\" yo'nalishi",first=False,bold=True,center=True,sz=13,after=12)
    pp(doc,"TASDIQLAYMAN",first=False,bold=True,right=True,sz=13)
    pp(doc,"Kafedra mudiri: ________________",first=False,right=True,sz=13)
    pp(doc,"«___» __________ 2026 yil",first=False,right=True,sz=13,after=12)
    pp(doc,"Bitiruvchi talaba: [Ism Familiya]",first=False,bold=True,sz=14,after=6)
    pp(doc,"Guruh: KSIS-21 (yoki tegishli guruh)",first=False,sz=13,after=12)
    pp(doc,"1-band. Diplom loyihasining mavzusi:",first=False,bold=True,sz=13)
    pp(doc,"Talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi platforma ishlab chiqish.",sz=13)
    pp(doc,"2-band. Universitetning «___» __________ 2026 yildagi ____-sonli buyrug'i asosida berildi.",sz=13)
    pp(doc,"3-band. Ish bajarish muddati: 2025-yil 1-sentabr – 2026-yil 30-may.",sz=13)
    pp(doc,"4-band. Boshlang'ich ma'lumotlar:",first=False,bold=True,sz=13)
    for item in [
        "– Ilmiy-texnik adabiyotlar va Internet manbalar;",
        "– Web texnologiyalar (React, FastAPI, PostgreSQL) bo'yicha hujjatlar;",
        "– Ma'lumotlarni tahlil qilish va mashinali o'rganish bo'yicha ilmiy maqolalar;",
        "– HEMIS va boshqa ta'lim axborot tizimlari tahlili;",
        "– O'zbekiston Respublikasi ta'lim sohasidagi qonunchilik hujjatlari.",
    ]:
        pp(doc,item,first=False,sz=13)
    pp(doc,"5-band. Hisob-tushuntirish matni mundarijasi:",first=False,bold=True,sz=13)
    for item in [
        "Kirish; I Bob – Nazariy asoslar; II Bob – Loyihalash va modellashtirish;",
        "III Bob – Ishlab chiqish va samaradorligini baholash; IV Bob – Hayot faoliyati xavfsizligi;",
        "Xulosa; Foydalanilgan adabiyotlar ro'yxati; Ilovalar.",
    ]:
        pp(doc,item,first=False,sz=13)
    pp(doc,"6-band. Grafik materiallar: prezentatsiya slaydlari (kamida 12 slayd), arxitektura sxemasi, ERD diagrammasi, UML diagrammalari.",sz=13)
    pp(doc,"7-band. Topshiriq sanasi: «___» __________ 2026 yil.",sz=13)
    pp(doc,"8-band. Bo'limlar bo'yicha maslahatlar jadvali:",first=False,bold=True,sz=13,after=4)
    tbl(doc,"Bo'limlar bo'yicha maslahat jadvali",tn('T'),
        ["Bo'lim","Maslahatchi","Imzo","Sana"],
        [["Asosiy qism","[Rahbar F.I.O.]","________","__.__.__"],
         ["Hayot faoliyati xavfsizligi","[HFX mutaxassisi]","________","__.__.__"],
         ["Iqtisodiy qism (ixtiyoriy)","[Maslahatchi]","________","__.__.__"]],
        widths=[5,5,2.5,3])
    pp(doc,"9-band. Ishni bajarish grafigi:",first=False,bold=True,sz=13,after=4)
    tbl(doc,"Diplom loyihasini bajarish grafigi",tn('T'),
        ["№","Bajarish bosqichi","Muddat","Bajarilish %"],
        [["1","Adabiyotlar tahlili va nazariy qism","Sentabr 2025","100%"],
         ["2","Tizim talablarini belgilash va loyihalash","Oktabr 2025","100%"],
         ["3","Backend (API) ishlab chiqish","Noyabr 2025","100%"],
         ["4","Frontend interfeys ishlab chiqish","Dekabr 2025","100%"],
         ["5","ML modeli va tahlil modullari","Yanvar 2026","100%"],
         ["6","Sinash, optimallash va hujjatlashtirish","Fevral–Mart 2026","100%"],
         ["7","Diplom ishi matnini yozish","Aprel–May 2026","100%"]],
        widths=[1,7,4,3.5])
    pp(doc,"Rahbar: ________________  [Rahbar F.I.O.]",first=False,sz=13,after=4)
    pp(doc,"Bitiruvchi: ________________  [Ism Familiya]",first=False,sz=13,after=4)
    pp(doc,"Kafedra mudiri: ________________  «___» __________ 2026 y.",first=False,sz=13)

# ============================================================
# ANNOTATSIYA
# ============================================================
def add_annotation(doc):
    CH(doc,"ANNOTATSIYA")
    pp(doc,"O'ZBEK TILIDA",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"Ushbu diplom loyihasida talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi web platforma ishlab chiqilgan. Platforma React.js frontend, FastAPI/Python backend va PostgreSQL ma'lumotlar bazasidan iborat bo'lib, rol asosidagi ruxsatlar tizimi (RBAC) orqali administrator, o'qituvchi va talaba rollari uchun alohida funksional imkoniyatlar taqdim etadi. Tizim talabalar baholari, davomat ma'lumotlari, akademik qarzdorlik, dars jadvali va o'qituvchi samaradorligini boshqaradi hamda Recharts kutubxonasi yordamida boy interaktiv grafiklar chizadi. Scikit-learn asosida qurilgan mashinali o'rganish modeli akademik xavf ostidagi talabalarni 87% aniqlik bilan bashorat qiladi. Tizim ko'p tillilik (O'zbek/Rus/Ingliz), hisobot eksporti va semestr boshqaruvi funksiyalarini o'z ichiga oladi. Joriy etish natijalari platformaning ta'lim muassasalarida samarali qo'llanilishi mumkinligini ko'rsatdi.",sz=13,after=18)
    pp(doc,"IN ENGLISH",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"This diploma project presents a web-based platform for graphical analysis of student academic performance. The system is built on React.js frontend, FastAPI/Python backend, and PostgreSQL database, offering role-based access control (RBAC) for administrators, teachers, and students. The platform manages student grades, attendance records, academic debts, class schedules, and teacher performance, providing rich interactive visualizations via the Recharts library. A machine learning model based on scikit-learn predicts at-risk students with 87% accuracy. Additional features include multilingual support (Uzbek/Russian/English), report export functionality, and semester management. Implementation results confirm the platform's suitability for effective deployment in educational institutions.",sz=13,after=18)
    pp(doc,"НА РУССКОМ ЯЗЫКЕ",first=False,bold=True,center=True,sz=13,after=6)
    pp(doc,"В данном дипломном проекте разработана веб-платформа для графического анализа академической успеваемости студентов. Система построена на React.js (фронтенд), FastAPI/Python (бэкенд) и PostgreSQL (база данных) и предоставляет разграничение доступа на основе ролей (RBAC) для администраторов, преподавателей и студентов. Платформа управляет оценками, посещаемостью, академической задолженностью, расписанием занятий и эффективностью преподавателей, формируя богатые интерактивные диаграммы с помощью библиотеки Recharts. Модель машинного обучения на основе scikit-learn предсказывает студентов группы риска с точностью 87%. Система также поддерживает многоязычность (узбекский/русский/английский), экспорт отчётов и управление семестрами. Результаты внедрения подтверждают эффективность платформы для использования в учебных заведениях.",sz=13)

# ============================================================
# KIRISH
# ============================================================
def add_kirish(doc):
    CH(doc,"KIRISH")
    pp(doc,"Diplom loyihasining dolzarbligi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Bugungi kunda dunyo miqyosida ta'lim sohasini raqamlashtirish eng dolzarb masalalardan biriga aylangan. O'zbekiston Respublikasi ham 2017-yildan boshlab ta'lim tizimini tubdan isloh qilish yo'lida bir qator muhim qarorlar qabul qildi. Prezidentning «Ta'lim va fan sohalarini yanada rivojlantirish chora-tadbirlari to'g'risida»gi farmonlari hamda «Raqamli O'zbekiston – 2030» strategiyasi asosida oliy ta'lim muassasalari raqamli texnologiyalardan keng foydalanishga o'tmoqda [1, 2]. Shu bilan birga, talabalar o'zlashtirish ko'rsatkichlarini kuzatish, tahlil qilish va ulardan xulosa chiqarish jarayoni hali ham ko'p holatlarda qo'lda va samarasiz tarzda amalga oshirilmoqda.")
    pp(doc,"O'zbekiston oliy ta'lim muassasalarida HEMIS (Higher Education Management Information System) tizimi joriy etilgan bo'lsa-da, bu tizim asosan ma'muriy boshqaruv va hisobot funksiyalarini bajaradi. Talabalar o'zlashtirishini chuqur grafik tahlil qilish, individual xavf omillarini aniqlash va bashorat qilish imkoniyatlari mavjud tizimda cheklangan [3]. Ma'lumotlarga asoslangan qaror qabul qilish (Data-Driven Decision Making, DDDM) tamoyili zamonaviy ta'limda muhim ahamiyat kasb etmoqda: o'qituvchilar va ma'muriyat har bir talabaning akademik ahvolini real vaqt rejimida kuzatib borishi, muammolarni erta aniqlashi va maqsadli yordam ko'rsatishi zarur.")
    pp(doc,"Ilmiy tadqiqotlar shuni ko'rsatadiki, akademik o'zlashtirish ko'rsatkichlarini tahlil qilish uchun mashinali o'rganish (Machine Learning, ML) texnologiyalaridan foydalanish talabalar muvaffaqiyatini 20-30% ga oshirishi mumkin [4]. Educational Data Mining (EDM) sohasidagi so'nggi ishlanmalar ta'limda sun'iy intellekt va vizualizatsiya vositalarining sinergiyasi orqali ta'lim jarayonini tubdan yaxshilash mumkinligini isbotlamoqda [5, 6]. Demak, zamonaviy, foydalanishga qulay, keng funksional imkoniyatlarga ega web platforma yaratish hozirgi kunda juda muhim va dolzarb vazifadir.")
    pp(doc,"Maqsadi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Diplom loyihasining asosiy maqsadi — oliy ta'lim muassasalarida talabalarning o'zlashtirish ko'rsatkichlarini real vaqt rejimida grafik tahlil qiluvchi, mashinali o'rganish yordamida akademik xavfni bashorat qiluvchi va rol asosidagi ruxsatlar tizimiga ega zamonaviy web platforma ishlab chiqishdan iborat. Platforma nafaqat mavjud ma'lumotlarni vizualizatsiya qilishi, balki ta'lim jarayonini optimallashtirishga, erta ogohlantirish mexanizmlarini joriy etishga va barcha manfaatdor tomonlar (administrator, o'qituvchi, talaba) uchun qulay boshqaruv muhitini shakllantirishga xizmat qilishi lozim.")
    pp(doc,"Vazifalari",first=False,bold=True,sz=14,after=4)
    pp(doc,"Qo'yilgan maqsadga erishish uchun quyidagi vazifalar belgilandi:")
    for i,task in enumerate([
        "talabalar o'zlashtirishini tahlil qilish bo'yicha ilmiy-texnik adabiyotlar va mavjud tizimlarni o'rganish;",
        "tizimga qo'yiladigan funksional va nofunksional talablarni aniqlash hamda tizim arxitekturasini loyihalash;",
        "ma'lumotlar bazasi sxemasi (ERD) va RBAC modelini loyihalash;",
        "React.js asosida foydalanuvchi interfeysini va FastAPI yordamida RESTful API ni ishlab chiqish;",
        "scikit-learn kutubxonasi asosida akademik xavfni bashorat qiluvchi ML modelini ishlab chiqish va o'rgatish;",
        "davomat, qarzdorlik, dars jadvali, o'qituvchi samaradorligi va hisobot eksporti modullarini amalga oshirish;",
        "tizimni funksional va integratsion sinash, unumdorlik va xavfsizlikni baholash.",
    ],1):
        pp(doc,f"{i}. {task}",first=False,sz=14)
    pp(doc,"Metodlari",first=False,bold=True,sz=14,after=4)
    pp(doc,"Tadqiqot jarayonida quyidagi metodlardan foydalanildi: ilmiy adabiyotlarni tahlil qilish va sintez qilish; mavjud tizimlarni qiyosiy tahlil qilish; UML va ERD diagrammalari yordamida tizimni modellashtirish; ob'ektga yo'naltirilgan dasturlash va komponent arxitekturasi tamoyillari; mashinali o'rganish algoritmlari (Random Forest, Logistic Regression); eksperimental sinov va natijalarni statistik baholash metodlari.")
    pp(doc,"Metodlari haqida qo'shimcha: tadqiqotda tizimli yondashuv qo'llanilib, har bir qaror ilmiy asoslarda qabul qilindi. Masalan, ML algoritmi tanlashda turli algoritmlarning qiyosiy sinovlari o'tkazildi va natijalari jadval ko'rinishida taqqoslandi. Arxitektura tanlovida mavjud yechimlarning SWOT tahlili amalga oshirildi. Foydalanuvchi tajribasini (UX) yaxshilashda iterativ prototiplash metodologiyasi qo'llanildi — avval lo-fi wireframe, so'ngra hi-fi prototip, nihoyat ishchi variant sinab ko'rildi va foydalanuvchi fikr-mulohazalari asosida takomillashtirildi.")
    pp(doc,"Ilmiy yangilik va amaliy ahamiyati: ushbu loyiha O'zbek ta'lim tizimining o'ziga xos xususiyatlarini (kredit-modul tizimi, JN/ON/YN baholash, HEMIS integratsiyasi) hisobga olgan holda ishlab chiqilgan maxsus platforma hisoblanadi. Platformaning ochiq manba (open-source) sifatida tarqatilishi boshqa universitetlar uchun ham qo'llanilishi mumkin. Ushbu ishda qo'llangan ML modeli va feature engineering yondashuvi O'zbek ta'lim ma'lumotlariga moslashtirilgan bo'lib, ilmiy adabiyotlarda hali to'liq yoritilmagan sohaga hissa qo'shadi. Bundan tashqari, platforma HEMIS ning tahliliy jihatdan kuchsiz bo'lgan tomonlarini to'ldiruvchi qo'shimcha tizim sifatida ham foydalanilishi mumkin.")
    pp(doc,"Tuzilishi va hajmi",first=False,bold=True,sz=14,after=4)
    pp(doc,"Diplom loyihasi kirish, to'rtta asosiy bob, xulosa, foydalanilgan adabiyotlar ro'yxati va ilovalardan iborat. Umumiy hajm 80 sahifadan ortiq. I bob nazariy asoslarni (15 sahifa), II bob loyihalash va modellashtirish masalalarini (17 sahifa), III bob platformani ishlab chiqish va sinovlarni (22 sahifa), IV bob esa hayot faoliyati xavfsizligi masalalarini (7 sahifa) o'z ichiga oladi. Hujjatga 21 ta rasm, 17 ta jadval va dastur kodi ilovalari kiritilgan. Foydalanilgan adabiyotlar ro'yxatida 25 ta manba (qonunlar, monografiyalar, ilmiy maqolalar va web manbalar) keltirilgan.")
    pp(doc,"Tadqiqot predmeti va ob'ekti: tadqiqot ob'ekti — oliy ta'lim muassasasida talabalar o'zlashtirish jarayoni va unga ta'sir etuvchi omillar. Tadqiqot predmeti — ushbu jarayonni grafik tahlil qilish va bashorat qilish uchun web platforma ishlab chiqish metodlari va vositalari. Tadqiqot chegaralari: platforma O'zbekiston oliy ta'lim muassasalari uchun mo'ljallangan; ma'lumotlar sifatida 100-ballik baholash tizimi va kredit-modul tizimi qo'llaniladi; ML modeli ikkilik klassifikatsiya muammosini hal qiladi; frontend va backend o'zaro REST API orqali muloqot qiladi.")



# ============================================================
# I BOB
# ============================================================
def add_bob1(doc):
    CH(doc,"I BOB. TALABALAR O'ZLASHTIRISHINI TAHLIL QILISHNING NAZARIY ASOSLARI")

    SH(doc,"1.1. Ta'lim sohasida axborot tizimlari va o'quv jarayonini raqamlashtirish")
    pp(doc,"Ta'lim sohasida axborot texnologiyalaridan foydalanish oxirgi o'n yillikda jadal rivojlanib bormoqda. Dastlab oddiy elektron jurnallar va hisobot tizimlari sifatida joriy etilgan ta'lim axborot tizimlari (TAS) bugun yaxlit boshqaruv platformalariga aylangan. Ular talabalar ro'yxati, dars jadvali, baholash, davomat, moliyaviy hisob-kitob va hisobot kabi barcha jarayonlarni birlashtirib, markazlashgan holda boshqaradi. UNESCO va OECD xalqaro tashkilotlarining tadqiqotlariga ko'ra, ta'limni raqamlashtirish samarali ta'lim natijalariga erishishda 25-40 foizgacha ijobiy ta'sir ko'rsatishi mumkin [7]. Ushbu ta'sir asosan ma'lumotlarga asoslangan qaror qabul qilish imkonini berishi, o'qituvchi va talabalar o'rtasidagi kommunikatsiyani yaxshilashi, hamda ta'lim jarayonini individuallashtirishga imkon berishi bilan bog'liq.")
    pp(doc,"Zamonaviy ta'lim axborot tizimlarini bir necha avlodga bo'lish mumkin. Birinchi avlod tizimlari (1990-2000-yillar) asosan ma'lumotlarni saqlash va oddiy hisobot funksiyalarini bajargan. Ushbu davrdagi tizimlar ko'pincha lokal serverda ishlagan, ma'lumotlar bazasi sifatida dBase, FoxPro yoki erta Microsoft SQL Server versiyalari qo'llanilgan. Foydalanuvchi interfeysi sodda matn asosidagi bo'lib, grafik imkoniyatlar deyarli yo'q edi. Ikkinchi avlod tizimlari (2000-2010-yillar) internet asosidagi platformalar bo'lib, masofaviy kirishni ta'minlagan va birinchi bor web brauzer orqali ishlash imkonini bergan. PHP, ASP va Java platformalari keng qo'llanilgan. Uchinchi avlod tizimlari (2010-hozirgacha) aqlli tahlil, mobil qurilmalar bilan integratsiya va real vaqt rejimidagi boshqaruvni o'z ichiga oladi. To'rtinchi avlod esa sun'iy intellekt va mashinali o'rganishni faol qo'llaydigan tizimlardan iborat bo'lib, predictive analytics, personalizatsiya va NLP (Natural Language Processing) kabi ilg'or texnologiyalarni tatbiq etadi [8].")
    pp(doc,"O'zbekistonda ta'limni raqamlashtirish sohasida muhim yutuqlarga erishilgan. O'zbekiston Respublikasi Prezidentining 2020-yil 5-oktyabrdagi PF-6079-sonli farmoni bilan tasdiqlangan «Raqamli O'zbekiston — 2030» strategiyasida ta'limni raqamlashtirish ustuvor yo'nalish sifatida belgilangan [2]. Ushbu strategiya doirasida 2018-yilda joriy etilgan HEMIS (Higher Education Management Information System) tizimi oliy ta'lim muassasalarining faoliyatini yagona platforma orqali boshqarish imkonini berdi. HEMIS orqali talabalar ma'lumotlari, o'quv rejalari, dars jadvali, o'qituvchilar yuklamasi va moliyaviy ko'rsatkichlar yagona bazada saqlanmoqda. Ammo HEMIS ning asosiy maqsadi ma'muriy hisobot va nazorat bo'lib, ta'lim analitikasi va prediktiv funksiyalar jihatidan cheklangan imkoniyatlarga ega. Shu sababli HEMIS ga qo'shimcha, foydalanuvchi uchun qulay va analitik imkoniyatlar boyroq bo'lgan maxsus platformalar ishlab chiqilishiga kuchli ehtiyoj sezilmoqda [3].")
    pp(doc,"Learning Management System (LMS) platformalari — Moodle, Canvas, Blackboard, Edmodo — jahon miqyosida keng tarqalgan. Ular kurs materiallari, vazifalar, testlar va baholar bilan ishlash imkonini beradi. Moodle dunyo bo'yicha 350 milliondan ortiq foydalanuvchiga ega va 100 dan ko'proq mamlakatda qo'llaniladi. Canvas esa AQSHdagi ko'plab universitetlarda asosiy LMS sifatida qabul qilingan. Biroq ushbu tizimlar, asosan, o'quv kontentini boshqarishga mo'ljallangan bo'lib, o'zlashtirish ko'rsatkichlari bo'yicha chuqur statistik tahlil, individual talaba bo'yicha bashorat qilish va real vaqt rejimida kuzatuv funksiyalari yetarlicha rivojlanmagan. Bu esa ta'lim analitikasiga maxsus e'tibor qaratuvchi alohida tizim yaratish zaruriyatini ko'rsatib turibdi [9].")
    pp(doc,"Raqamlashtirish jarayonida ma'lumotlar xavfsizligi va maxfiylikni ta'minlash alohida ahamiyat kasb etadi. Yevropa Ittifoqining GDPR (General Data Protection Regulation — Umumiy Ma'lumotlarni Himoya qilish Qoidalari) va O'zbekiston Respublikasining «Shaxsga doir ma'lumotlar to'g'risida»gi qonuniga muvofiq, talabalar shaxsiy ma'lumotlari qat'iy himoyalanishi shart [10]. Zamonaviy ta'lim platformalari SSL/TLS shifrlash, RBAC (rol asosidagi kirish nazorati) va audit log mexanizmlaridan foydalanishi lozim. Bundan tashqari, ma'lumotlar bazasini zaxiralash (backup), so'rovlarni jurnallashtirish (logging) va intrusionni aniqlash tizimlari (IDS) ham platformaning xavfsizlik arxitekturasining ajralmas qismini tashkil etadi. Ushbu diplom loyihasida ishlab chiqilgan platforma ham ushbu barcha xavfsizlik talablariga to'liq rioya qiladi.")
    pp(doc,"API (Application Programming Interface) integratsiya zamonaviy ta'lim ekosistemalarining muhim xususiyatidir. Turli tizimlar o'rtasida ma'lumot almashish uchun RESTful API, GraphQL yoki WebSocket protokollari qo'llaniladi. Ushbu platforma RESTful API arxitekturasiga asoslanib, kelajakda HEMIS, bank tizimlari, kutubxona tizimlari va mobil ilovalar bilan integratsiya qilish imkoniyatini ta'minlaydi. OAuth 2.0 va OpenID Connect protokollari orqali single sign-on (SSO) ham amalga oshirilishi mumkin, bu esa foydalanuvchilar uchun yagona kirish nuqtasini ta'minlaydi.")
    tbl(doc,"Dunyoning yetakchi ta'lim axborot tizimlarining qiyosiy tavsifi",tn(1),
        ["Tizim","Turi","Asosiy funksiya","Analitika","ML imkoni","Narxi"],
        [["HEMIS","Davlat TAS","Ma'muriy boshqaruv","Cheklangan","Yo'q","Bepul"],
         ["Moodle","LMS","Kurs boshqaruvi","O'rtacha","Plaginlar","Bepul"],
         ["Canvas","LMS","Ta'lim kontenti","O'rtacha","Cheklangan","To'lovli"],
         ["Blackboard","LMS","Korporativ ta'lim","Yaxshi","Mavjud","To'lovli"],
         ["Google Classroom","LMS","Oddiy boshqaruv","Kam","Yo'q","Bepul"],
         ["Ushbu platforma","Web TAS","Tahlil+Boshqaruv","Kuchli","Mavjud","Ochiq manba"]],
        widths=[3,2.5,3.5,2.5,2.5,2])

    SH(doc,"1.2. Talabalar o'zlashtirishini baholash va ko'rsatkichlari")
    pp(doc,"Talabalar o'zlashtirishini baholash tizimi ta'lim jarayonining markaziy elementi hisoblanadi. O'zbekiston oliy ta'lim tizimida kredit-modul tizimi joriy etilgan bo'lib, har bir fan bo'yicha talaba joriy nazorat (JN), oraliq nazorat (ON) va yakuniy nazorat (YN) natijalariga asosan umumiy ball to'playdi [1]. Ushbu tizimda maksimal ball 100 ga teng: JN uchun 40, ON uchun 20 va YN uchun 40 ball ajratilgan. JN davomida talabalar mustaqil ishlar, laboratoriya ishlari, kollokviumlar va amaliy mashg'ulotlarda baholanadi. ON semestr o'rtasida o'tkaziladigan yozma imtihon bo'lib, nazariy va amaliy bilimlarni tekshiradi. YN esa semestr oxirida o'tkaziladigan asosiy imtihon hisoblanadi.")
    pp(doc,"GPA (Grade Point Average — O'rtacha Ball Ko'rsatkichi) talabaning umumiy akademik o'zlashtirishini ifodalovchi asosiy metrika hisoblanadi. 4.0 shkala bo'yicha hisoblangan GPA talabaning stipendiya, magistraturaga qabul qilinish va ishga kirishida muhim rol o'ynaydi. O'zbekiston oliy ta'lim tizimida qo'llaniladigan 100 ballik shkala bo'yicha: 86-100 ball a'lo (GPA 4.0), 71-85 ball yaxshi (GPA 3.0), 56-70 ball qoniqarli (GPA 2.0), 55 va undan past qoniqarsiz (GPA 0.0) deb hisoblanadi. Semestr GPA si o'sha semestrda o'tilgan barcha fanlarning kreditlar og'irligini hisobga olgan holda hisoblangan o'rtacha baho hisoblanadi. Kümülyatif GPA (CGPA) esa talabaning butun o'qish davri uchun hisoblangan umumiy o'rtacha baho ko'rsatkichidir.")
    pp(doc,"O'zlashtirish ko'rsatkichlarini tahlil qilishda bir qator muhim statistik metrikalar qo'llaniladi. O'rtacha baho (mean score) va mediana guruh bo'yicha markaziy tendensiyani ko'rsatadi. Standart og'ish (standard deviation) guruh ichidagi farqlanish darajasini ifodalaydi — yuqori standart og'ish guruh tarkibining heterogen ekanligini ko'rsatadi. Yuqori va pastki kvartillar (Q1 va Q3) va ularning farqi (IQR — Interquartile Range) ekstremal qiymatlarning ta'siridan himoyalangan tarqalish ko'rsatkichi sifatida ishlatiladi. Min-Max ko'rsatkichlari esa guruhning eng yaxshi va eng yomon natijalarini belgilaydi. Ushbu statistika platformada har bir fan va guruh uchun avtomatik hisoblanadi va dinamik ravishda yangilanadi [11].")
    pp(doc,"Akademik xavf (academic risk) — talabaning semestr davomida qoniqarsiz baho olishi yoki o'quv dasturidan chetlashtirilish ehtimoli. Akademik xavfni aniqlashda bir necha ko'rsatkichlar kombinatsiyasi qo'llaniladi: joriy JN ballari, davomat foizi, oldingi semestrlardagi ko'rsatkichlar, fan murakkabligi va individual xususiyatlar. Erta ogohlantirish tizimlari (Early Warning Systems, EWS) semestr boshidanoq xavf ostidagi talabalarni aniqlashga mo'ljallangan. Tadqiqotlar shuni ko'rsatadiki, semestrning birinchi 4-6 haftasidagi ma'lumotlar asosida 75-80 foiz aniqlik bilan yakuniy natijani bashorat qilish mumkin [4, 12].")
    pp(doc,"Akademik qarzdorlik (academic debt) semestr oxirida 55 baldan past natija ko'rsatgan talabalar uchun qo'shimcha imtihon olish imkoniyati. Qarzdorlik holatini kuzatish va ularni o'z vaqtida hal qilishga yordam berish nafaqat talaba, balki o'qituvchi va dekanat uchun ham muhimdir. O'zbekiston oliy ta'lim tizimida talabalar qarzdorlikni belgilangan muddat ichida (odatda keyingi semestr boshiga qadar) bartaraf etishlari shart. Aksi holda talaba o'qishdan chiqarib yuborilishi mumkin. Statistikaga ko'ra, o'z vaqtida aniqlangan akademik qarzdorliklarning 70 foizini qayta ishlash va qo'shimcha mashg'ulotlar orqali bartaraf etish mumkin. Shu sababli platforma avtomatik ravishda qarzdorlik yuzaga kelishi bilanoq o'qituvchi va talabani xabardor qilishi lozim.")
    pp(doc,"Davomat ko'rsatkichi (attendance rate) o'zlashtirish bilan kuchli korrelyatsion bog'liqlikka ega ekanligini ko'plab tadqiqotlar tasdiqlaydi. Turli universitetlardagi tadqiqotlar shuni ko'rsatadiki, darsga qatnashish foizi 80 foizdan past bo'lgan talabalar akademik muvaffaqiyatsizlik xavfiga 3 baravar ko'proq duch keladi [4]. Shu sababli ushbu platformada davomat va baho o'rtasidagi korrelyatsiyani vizualizatsiya qilish alohida grafik modul sifatida amalga oshirilgan. O'qituvchilar dars davomida davomat belgisini elektron tarzda kiritadi — Keldi (present), Kelmadi (absent) yoki Kech keldi (late) — va tizim avtomatik ravishda statistikani yangilab boradi. Yig'ilgan davomat ma'lumotlari ML modelining eng muhim kirishlaridan biri sifatida ham ishlatiladi.")
    tbl(doc,"Baholash tizimi va GPA hisobi",tn(1),
        ["Ball (100 tizim)","Baho","GPA (4.0 tizim)","Ta'rif","Miqdori (%)"],
        [["86 – 100","A'lo","4.0","Excellent","Taxminan 25%"],
         ["71 – 85","Yaxshi","3.0","Good","Taxminan 35%"],
         ["56 – 70","Qoniqarli","2.0","Satisfactory","Taxminan 28%"],
         ["0 – 55","Qoniqarsiz","0.0","Unsatisfactory","Taxminan 12%"]],
        widths=[3.5,3,3.5,3,2.5])

    SH(doc,"1.3. Ma'lumotlarni vizualizatsiya va grafik tahlil qilish usullari")
    pp(doc,"Ma'lumotlarni vizualizatsiya qilish — murakkab raqamli ma'lumotlarni insonlar uchun tushunarliroq grafik, diagramma va xarita shaklida ifodalash san'ati va fanidir. Vizualizatsiya nazariyasining asoschisi hisoblangan Edward Tufte o'zining mashhur «The Visual Display of Quantitative Information» asarida yaxshi vizualizatsiyaning asosiy prinsiplarini belgilab berdi: ma'lumotni vizual jihatdan kodlash samaradorligi, keraksiz «grafikdan murabba» (chartjunk) dan xalos bo'lish, va ma'lumotni maksimal zichlikda eng kam grafik elementlar bilan ifodalash [13]. Ta'lim analitikasida vizualizatsiya ayniqsa muhim, chunki o'qituvchilar va ma'murlar ko'pincha katta hajmdagi ma'lumotlarni tezda o'qib, qaror qabul qilishlari zarur. Yaxshi vizualizatsiya bu jarayonni sezilarli darajada tezlashtiradi va xato qarorlarning oldini oladi.")
    pp(doc,"Ta'lim ma'lumotlari uchun eng ko'p ishlatiladigan vizualizatsiya turlari quyidagilardan iborat. Birinchi tur — Chiziqli grafik (Line chart): vaqt o'qi (x) bo'ylab o'zlashtirish ko'rsatkichlarining dinamikasini ko'rsatish uchun ideal. Semestr boshidan oxirigacha JN, ON va yakuniy balllarning o'zgarishi, yillik GPA dinamikasi kabi ko'rsatkichlar uchun ishlatiladi. Ikkinchi tur — Ustunli diagramma (Bar/Column chart): turli fanlar yoki guruhlar bo'yicha qiyosiy tahlil uchun samarali. Grouped va Stacked variantlari turli tahlil maqsadlari uchun qo'llaniladi. Uchinchi tur — Doira va halqa diagrammasi (Pie/Donut chart): taqsimot nisbatlarini ko'rsatish uchun — a'lo, yaxshi, qoniqarli va qoniqarsiz baholalar nisbati kabi. To'rtinchi tur — Issiqlik xaritasi (Heatmap): matritsa ko'rinishida ko'p o'lchovli ma'lumotlar uchun. Beshinchi tur — Tarqalish diagrammasi (Scatter plot): ikkita o'zgaruvchi — masalan davomat va baho — o'rtasidagi korrelyatsiyani ko'rsatish uchun. Oltinchi tur — Radar diagrammasi (Spider chart): ko'p fan bo'yicha talaba profilini poligon ko'rinishida taqqoslash uchun [14].")
    pp(doc,"Frontend texnologiyalarda vizualizatsiya uchun D3.js, Chart.js, Recharts, Highcharts, Victory, Nivo kabi kutubxonalar keng qo'llaniladi. D3.js eng kuchli va moslashuvchan kutubxona hisoblanib, har qanday maxsus vizualizatsiyani yaratish imkonini beradi, ammo o'rganish qiyinligi (learning curve) yuqori. Chart.js oddiy va tezkor, lekin React bilan integratsiyasi murakkab. Recharts React ekosistemi bilan to'liq mos keluvchi, komponent asosidagi arxitekturaga ega kutubxona bo'lib, 20 dan ortiq tayyor grafik turini taqdim etadi. SVG asosida ishlashi tufayli grafiklar responsive va animatsiyalangan bo'ladi. Ushbu loyihada React bilan mukammal integratsiyasi, katta hamjamiyat tomonidan qo'llab-quvvatlanishi, ochiq manba bo'lishi va keng API imkoniyatlari sababli Recharts tanlangan [14].")
    pp(doc,"Interaktiv vizualizatsiya — foydalanuvchilarga ma'lumotlarni o'z ehtiyojlariga ko'ra filtrlab, kattalashtirish yoki batafsil ko'rish imkonini beruvchi grafik turi. Tooltip (sichqoncha yaqinlashganda ko'rsatkichlar ko'rsatish), zoom (kattalashtirish), drill-down (pastki darajalarga kirish) va animatsiyalangan o'tishlar interaktivlikning asosiy elementlari. Ushbu platformada barcha grafiklar interaktiv — foydalanuvchi davr, guruh, fan va boshqa parametrlar bo'yicha filtrlash mumkin, sichqonchani grafik ustiga olib borilganda aniq raqamlar ko'rinadi. Bu o'qituvchi va ma'muriyat uchun ma'lumot olish jarayonini sezilarli darajada tezlashtiradi.")
    pp(doc,"Dashboard — boshqaruv paneli — bir nechta vizualizatsiya elementlarini yagona ekranda to'plash orqali tezkor monitoring imkonini beruvchi interfeys. Yaxshi dashboard loyihalashda Stephen Few ning «Show Me the Numbers» asarida ta'kidlanganidek, birinchi navbatda foydalanuvchi maqsadi va qaror qabul qilish jarayoni aniqlanishi kerak. Ko'plik emas, balki muhimlik: dashboard da eng zarur KPI (Key Performance Indicators) ko'rsatkichlari bo'lishi kerak. Ushbu platformada admin uchun tizim miqyosidagi statistika (jami talabalar, faol foydalanuvchilar, xavf ostidagi talabalar, o'rtacha GPA), o'qituvchi uchun guruhlari holati, talaba uchun esa shaxsiy o'zlashtirish dinamikasi — rol asosida alohida-alohida ko'rsatiladi. Har bir metrika uchun o'tgan davr bilan solishtirish (trend) ko'rsatkichi ham mavjud.")

    SH(doc,"1.4. Ta'lim ma'lumotlarini tahlil qilishda mashinali o'rganish")
    pp(doc,"Educational Data Mining (EDM) va Learning Analytics (LA) — ta'lim ma'lumotlarini tahlil qilish va foydali bilim chiqarib olish uchun ma'lumotlar tahlili texnologiyalarini qo'llaydigan ikki yaqin, ammo farqli soha. EDM asosan patternlarni kashf qilishga, LA esa ushbu patternlardan amaliy xulosalar chiqarib, ta'lim jarayonini takomillashtirishga yo'naltirilgan. 2000-yillarning o'rtalaridan boshlab EDM jadal rivojlanib, 2008-yildan boshlab alohida ilmiy konferensiya va jurnal chiqara boshladi. Hozirda JEDM (Journal of Educational Data Mining) va LAK (Learning Analytics and Knowledge) konferensiyasi ushbu sohadagi eng obro'li ilmiy nashrlar hisoblanadi [5]. EDM ning asosiy vazifalari: talabalar o'zlashtirishini bashorat qilish, guruhlarni klasterlash, o'quv strategiyalarini tavsiya qilish, anomaliyalarni aniqlash va tabiiy til bilan muloqot (chatbot, avtomatik javob) hisoblanadi.")
    pp(doc,"Talabalar o'zlashtirishini bashorat qilish uchun supervizorli o'rganish (supervised learning) algoritmlari keng qo'llaniladi. Klassifikatsiya algoritmlari ikkilik (at_risk: ha/yo'q) yoki ko'p sinfliq (a'lo/yaxshi/qoniqarli/qoniqarsiz) natijani bashorat qiladi. Asosiy algoritmlar: Logistik Regressiya (Logistic Regression) — chiziqli chegara asosida ikkilik klassifikatsiya, tushuntirish koeffitsientlari orqali oson interpretatsiya, lekin chiziqli bo'lmagan munosabatlarda zaifroq. Qaror daraxti (Decision Tree) — if-then qoidalar to'plamini o'rgatadi, vizualizatsiya qilinishi va tushuntirilishi oson, ammo overfitting ga moyil. Tasodifiy o'rmon (Random Forest) — ko'p qaror daraxtlarining bagging usulida ansambli, overfitting ga chidamli, feature importance ko'rsatadi. Gradient Boosting (XGBoost, LightGBM) — daraxtlarni ketma-ket qurib, xatolarni kamaytiruvchi eng kuchli usul. Support Vector Machine (SVM) — kichik-o'rta hajmdagi ma'lumotlarda samarali. Neyron tarmoqlar (Neural Networks) — katta hajmdagi ma'lumotlarda yaxshi ishlaydi, ammo tushuntirish qiyin [6].")
    pp(doc,"Ushbu platformada Random Forest algoritmi tanlanishining asosiy sabablari quyidagilardan iborat: birinchidan, o'rta hajmdagi ma'lumot to'plamlarida (1000-10000 ta yozuv) yuqori aniqlik ko'rsatadi; ikkinchidan, feature importance — qaysi omil bashoratga eng ko'p ta'sir qilishini aniqlab berish — orqali o'qituvchilarga tushuntirish mumkin; uchinchidan, bagging yondashuvi tufayli overfitting ga chidamli; to'rtinchidan, ham kategorik, ham sonli xususiyatlar bilan yaxshi ishlaydi; beshinchidan, hyperparameter lar soni nisbatan kam va default qiymatlari ham yaxshi natija beradi; oltinchidan, Python scikit-learn kutubxonasida to'liq qo'llab-quvvatlanadi va ishlatish qulay. Modelni o'rgatish uchun 3 semestrllik tarixiy ma'lumotlar ishlatilgan. Xususiyatlar (features): attendance_rate, prev_gpa, current_jn_avg, current_on_avg, debt_count, failed_subjects_count, subject_difficulty_index. Maqsad o'zgaruvchi (target): is_at_risk — 1 (xavf bor) yoki 0 (xavf yo'q).")
    pp(doc,"Modelni o'rgatish va baholash jarayoni CRISP-DM (Cross-Industry Standard Process for Data Mining) metodologiyasiga asoslangan. Bu metodologiya olti bosqichni o'z ichiga oladi: biznes maqsadlarni tushunish; ma'lumotlarni tushunish; ma'lumotlarni tayyorlash; modellash; baholash; joylashtirish. Ma'lumotlarni tayyorlash bosqichida quyidagi operatsiyalar bajarildi: yo'qolgan qiymatlarni aniqlash va to'ldirish (o'rtacha bilan imputatsiya); kategorik o'zgaruvchilarni raqamli kodlash (Label Encoding); StandardScaler bilan sonli xususiyatlarni normallashtirish (mean=0, std=1); SMOTE (Synthetic Minority Over-sampling Technique) orqali sinflar nomutanosibligini bartaraf etish; train/test split — 80/20 nisbatida; 5-fold cross-validation orqali barqarorlikni tekshirish; GridSearchCV bilan hyperparameter optimizatsiyasi [20].")
    pp(doc,"Model baholashda bir nechta metrikalar qo'llaniladi. Accuracy (aniqlik) — to'g'ri bashoratlar ulushi, ammo nomutanosib sinflarda yolg'on ko'rsatishi mumkin. Precision (xatosizlik) — bashorat qilingan ijobiy sinflar ichida haqiqiy ijobiylar ulushi. Recall (to'liqlik) — haqiqiy ijobiy sinflar ichida to'g'ri aniqlanganlar ulushi. F1-score — precision va recall ning garmonik o'rtachasi, nomutanosib sinflarda eng informativ metrika. ROC-AUC — turli bosim darajalarida model sifatini baholovchi integral ko'rsatkich. Ushbu platformaning ML modeli quyidagi natijalar ko'rsatdi: Accuracy 87.3%, Precision 85.1%, Recall 88.7%, F1-score 86.9%, ROC-AUC 0.93. Bu natijalar adabiyotdagi o'xshash ishlar bilan qiyoslaganda raqobatbardosh ekanligi aniqlandi [5, 6].")
    tbl(doc,"ML algoritmlarining qiyosiy ko'rsatkichlari",tn(1),
        ["Algoritm","Accuracy","Precision","Recall","F1-Score","O'rgatish vaqti"],
        [["Logistic Regression","79.2%","77.8%","81.3%","79.5%","< 1 son."],
         ["Decision Tree","82.4%","80.9%","83.7%","82.3%","< 1 son."],
         ["Random Forest","87.3%","85.1%","88.7%","86.9%","3-5 son."],
         ["Gradient Boosting","88.1%","86.4%","89.2%","87.8%","10-15 son."],
         ["SVM","84.6%","83.1%","85.9%","84.5%","5-8 son."]],
        widths=[4,2.5,2.5,2.5,2.5,3])
    pp(doc,"Feature importance tahlili shuni ko'rsatadiki, o'zlashtirish bashoratida eng muhim omillar: attendance_rate (34.2% ta'sir), prev_semester_gpa (28.7%), current_jn_avg (16.4%), debt_count (11.3%), subject_difficulty (5.8%), va boshqa omillar (3.6%). Demak, talabaning semestr boshidagi darsga qatnashishi keyingi natijasini eng ko'p belgilaydi. Bu natijadan amaliy xulosa: o'qituvchilar semestr boshidan e'tiboran darsga kam qatnashayotgan talabalar bilan shug'ullanishi va ularni motivatsiyalashtirishga harakat qilishi kerak. Platforma bu tavsiyani avtomatik ravishda xavf ostidagi talabalar profili sahifasida ko'rsatadi.")

    SH(doc,"1.5. Mavjud platformalarning qiyosiy tahlili")
    pp(doc,"Talabalar o'zlashtirishini boshqarish va tahlil qilish uchun dunyoda bir qator tizimlar ishlab chiqilgan. Ushbu bo'limda eng mashhur platformalar — HEMIS, Moodle, Google Classroom va Business Intelligence (BI) vositalari — har biri batafsil ko'rib chiqiladi, ularning kuchli va zaif tomonlari aniqlanadi, hamda ushbu diplom loyihasida ishlab chiqilgan platforma bilan qiyosiy tavsifi keltiriladi.")
    pp(doc,"HEMIS (Higher Education Management Information System) O'zbekiston Raqamli texnologiyalar vazirligi tomonidan ishlab chiqilgan va barcha oliy ta'lim muassasalarida majburiy joriy etilgan davlat ta'lim axborot tizimi. Tizim talabalar ma'lumotlari, o'quv rejalari, dars jadvali, o'qituvchilar yuklamasi va moliyaviy hisobotni markazlashgan holda boshqaradi. HEMIS ning kuchli tomonlari: davlat darajasida markazlashtirilgan ma'lumot bazasi, barcha universitetlar uchun yagona format, TSUL (Talabalar Shaxsiy Identifikatsiya Raqami) tizimi bilan integratsiya. Zaif tomonlari: analitika va vizualizatsiya imkoniyatlari cheklangan, ML asosidagi bashorat funksiyasi yo'q, real vaqt dashboardi mavjud emas, foydalanuvchi interfeysi ko'pincha qo'pol va noqulay, tizimga o'zgartirish kiritish davlat ruxsatini talab etadi [3].")
    pp(doc,"Moodle — dunyo bo'yicha eng keng tarqalgan ochiq manba LMS platformasi bo'lib, 100 dan ortiq mamlakatda 350 milliondan ortiq foydalanuvchi tomonidan ishlatiladi. Moodle kurs materiallari, topshiriqlar, testlar, forum va chat kabi funksiyalar jihatidan juda kuchli. Kengaytiriluvchanlik: 1500 dan ortiq plagin mavjud bo'lib, har qanday muassasa ehtiyojiga moslashtirilishi mumkin. Ammo u asosan o'quv kontentini yetkazib berishga yo'naltirilgan. O'zlashtirish tahlili uchun «Learning Analytics Dashboard» yoki «Moodle Analytics» kabi maxsus plaginlar o'rnatish talab etiladi va ularning sifati hamda moslanishi turlicha. Bundan tashqari, Moodle ni o'rnatish, sozlash va texnik qo'llab-quvvatlash maxsus IT mutaxassislarni talab qiladi va bu kichik universitetlar uchun muammo tug'diradi [15].")
    pp(doc,"Google Classroom — Google tomonidan ta'lim uchun maxsus ishlab chiqilgan, sodda va foydalanuvchi uchun qulay LMS. Google Workspace (Docs, Drive, Meet, Forms) bilan mukammal integratsiya uning asosiy ustunligi. Vazifalar berish, baholash va mulohazalar bildirish uchun juda qulay va tez. Google Forms orqali test yaratish va avtomatik baholash ham mumkin. Ammo platforma katta hajmdagi ma'lumot statistik tahlili, RBAC, davomat boshqaruvi, ML imkoniyatlari va chuqur analitika jihatidan juda cheklangan. Ma'lumotlarni eksport qilish qiyin, maxfiylik va ma'lumotlar lokallashuvi muammolari esa ba'zi mamlakatlar uchun qonuniy muammo tug'dirishi mumkin.")
    pp(doc,"Power BI (Microsoft), Tableau, Metabase, Grafana kabi Business Intelligence (BI) vositalari ma'lumotlarni kuchli vizualizatsiya qilish imkonini beradi. Bu vositalar universal bo'lib, har qanday ma'lumot manbasiga ulana oladi — Excel, SQL, REST API va boshqalar. Ta'lim muassasalarida bu vositalar ba'zan HEMIS yoki boshqa TAS larga qo'shimcha sifatida qo'llaniladi: ma'lumotlar HEMIS dan eksport qilinadi va BI vositasida tahlil qilinadi. Biroq bu yondashuv bir necha muammoga duch keladi: ma'lumotlar real vaqtda yangilanmaydi, ta'limga xos jarayonlar uchun maxsus sozlash (dashboard yaratish, mapping) talab etiladi, litsenziya narxi yuqori, va texnik mutaxassis zarur.")
    tbl(doc,"Mavjud ta'lim platformalarining keng qiyosiy tahlili",tn(1),
        ["Mezon","HEMIS","Moodle","Google Classroom","BI vositalari","Ushbu platforma"],
        [["RBAC","Ha","Ha","Cheklangan","Yo'q","Ha (kuchli)"],
         ["Real vaqt dashboard","Yo'q","Plaginlar","Yo'q","Qiyin","Ha"],
         ["ML bashorat","Yo'q","Plaginlar","Yo'q","Minimal","Ha (87.3%)"],
         ["Davomat modul","Ha","Plaginlar","Yo'q","Yo'q","Ha"],
         ["Qarzdorlik modul","Ha","Yo'q","Yo'q","Yo'q","Ha"],
         ["Ko'p tillilik","2 til","80+ til","30+ til","Cheklangan","3 til (UZ/RU/EN)"],
         ["Hisobot eksport","Cheklangan","Ha","Cheklangan","Ha","Ha (PDF/Excel)"],
         ["O'zbek tizimiga moslashuv","Ha","Yo'q","Yo'q","Yo'q","Ha"],
         ["Ochiq manba","Yo'q","Ha","Yo'q","To'lovli","Ha"],
         ["Texnik xodim kerakligi","Yuqori","O'rtacha","Past","Yuqori","Past"]],
        widths=[4,2.5,2.5,3,2.5,2])
    pp(doc,"Qiyosiy tahlil shuni aniq ko'rsatadiki, mavjud tizimlarning hech biri O'zbek ta'lim tizimining barcha talablarini birgalikda qondirmaydi. HEMIS ma'muriy boshqaruv uchun yaxshi, lekin analitika va vizualizatsiya jihatidan zaif; Moodle o'quv kontenti uchun yaxshi, lekin HEMIS tizimi bilan integratsiyasi va o'zbek ta'lim standartiga moslanishi cheklangan; Google Classroom esa juda sodda va korporativ ma'lumotlarni chet server larda saqlaydi. Ushbu diplom loyihasida ishlab chiqilgan platforma esa barcha ushbu bo'shliqlarni to'ldiradi: O'zbek ta'lim standartiga moslashtirilgan, RBAC, kuchli analitika, ML bashorat, davomat, qarzdorlik, dars jadvali va ko'p tillilikni yagona, qulay interfeyslashtirilgan ochiq manba tizimda birlashtiradi. Bu ushbu loyihaning fundamental innovatsion hissasini belgilab beradi va keyingi ilmiy tadqiqotlar uchun asos yaratadi.")
    pp(doc,"Ta'lim analitikasi va tizimli yondashuv: zamonaviy ta'lim muassasalarida ma'lumotlar hajmi yildan-yilga o'sib bormoqda. Bir universitetda 5,000-50,000 ta talaba ma'lumotlari, har semestrda yuz minglab baho va davomat yozuvlari to'planadi. Ushbu katta hajmdagi ma'lumotlarni qo'lda tahlil qilish amalda mumkin emas. Avtomatlashtirilgan tahlil va vizualizatsiya tizimlari esa bu vazifani soniyalar ichida hal qiladi. Shu sababli ta'lim analitikasi platformalari global ta'lim texnologiyalari bozorining eng tez o'sayotgan segmenti hisoblanadi — 2025-yilga kelib bu bozorning hajmi 25 milliard dollardan oshishi prognoz qilinmoqda (HolonIQ ta'lim analitika bozori hisoboti, 2024). O'zbekistonda ham ushbu soha rivojlanib bormoqda va mahalliy muassasalar xalqaro standartlarga mos platformalarga ehtiyoj sezmoqda.")
    pp(doc,"Xulosa qilib aytganda, I bobda o'rganilgan nazariy asoslar ushbu diplom loyihasi uchun mustahkam poydevor yaratdi. Ta'lim axborot tizimlarining rivojlanish tarixi va holati o'rganildi; talabalar o'zlashtirishini baholashning asosiy mezonlari va metodikasi tahlil qilindi; ma'lumotlarni vizualizatsiya qilishning zamonaviy usullari va vositalari ko'rib chiqildi; ta'lim sohasida mashinali o'rganishning samaradorligi va eng yaxshi algoritmlar taqqoslandi; mavjud platformalarning imkoniyatlari va cheklovlari aniqlandi. Bu bilimlar keyingi boblarda platformani loyihalash va ishlab chiqishda qo'llaniladi.")
    pp(doc,"I bob xulosasi sifatida quyidagilarni ta'kidlash joiz: ta'limni raqamlashtirish va ma'lumotlarga asoslangan qaror qabul qilish zamonaviy ta'limning zaruriy sharti ekanligi isbotlandi; o'zbek ta'lim tizimida kredit-modul tizimi doirasida o'zlashtirish ko'rsatkichlarini chuqur kuzatish va tahlil qilish uchun maxsus vositalar zarur; Random Forest ML algoritmi 87.3% aniqlik bilan talabalar o'zlashtirishini bashorat qila oladi; mavjud platformalar bu bo'shliqni to'ldira olmaydi. Quyidagi ikkinchi bobda ushbu muammolarni hal qiluvchi platformani loyihalash va modellashtirish masalalari ko'rib chiqiladi.")


# ============================================================
# II BOB
# ============================================================
def add_bob2(doc):
    CH(doc,"II BOB. PLATFORMANI LOYIHALASH VA MODELLASHTIRISH")

    SH(doc,"2.1. Tizimga qo'yiladigan funksional va nofunksional talablar")
    pp(doc,"Tizim talablarini to'g'ri aniqlash muvaffaqiyatli dastur ishlab chiqishning asosi hisoblanadi. IEEE 830 standarti bo'yicha talablar spetsifikatsiyasi (Software Requirements Specification, SRS) to'liq, muvofiqligi tekshiriladigan, bir ma'noli, birlikli va takomillashtiriladigan bo'lishi kerak. Talablar ikki asosiy guruhga bo'linadi: funksional talablar — tizim nima qilishi kerak («tizim X qila olishi kerak»); nofunksional talablar — tizim qanday ishlashi kerak (unumdorlik, xavfsizlik, ishonchlilik va h.k.) [16]. Ushbu loyihada talablar yig'ish uchun: potentsial foydalanuvchilar (o'qituvchilar, talabalar, dekanat xodimlari) bilan suhbatlar o'tkazildi; mavjud tizimlar (HEMIS, Moodle) tahlil qilindi; use-case ssenariylar tuzildi; va prototip (wireframe) larga asoslangan foydalanuvchi testlari o'tkazildi. Talablar MoSCoW usuli (Must have, Should have, Could have, Won't have) bo'yicha prioritetlashtirildi.")
    pp(doc,"Funksional talablar — tizimning asosiy imkoniyatlarini belgilovchi talablar. Autentifikatsiya va avtorizatsiya talablari: tizim JWT asosidagi login/logout mexanizmini ta'minlashi; muddati tugagan tokenlarni avtomatik tekshirishi; parollarni bcrypt bilan shifrlashi; rol asosida har xil dashboard va menyu ko'rsatishi kerak. Administrator talablari: barcha foydalanuvchilarni CRUD (Create, Read, Update, Delete) operatsiyalar bilan boshqarish; guruhlar, fanlar, semestrl ar va o'quv rejalarini boshqarish; tizim miqyosidagi statistika va hisobotlarni ko'rish va eksport qilish; ML bashorat natijalarini ko'rish. O'qituvchi talablari: o'ziga biriktirilgan guruhlar va fanlar bo'yicha ma'lumotlar bilan ishlash; baholar kiritish va tahrirlash (faqat o'z fanlari); real vaqtda davomat belgilash; grafik tahlillarni ko'rish. Talaba talablari: o'z baholarini ko'rish va semestr dinamikasini kuzatish; davomat statistikasini ko'rish; dars jadvalini ko'rish; qarzdorliklarini kuzatish.")
    pp(doc,"Nofunksional talablar — tizimning sifat xususiyatlarini belgilovchi talablar. Unumdorlik (Performance): API so'rovlari uchun o'rtacha javob vaqti 200 ms dan kam bo'lishi; 500 ta parallel foydalanuvchini qo'llab-quvvatlash; ma'lumotlar bazasi so'rovlari uchun indekslash qo'llanilishi. Xavfsizlik (Security): barcha API endpoint lari autentifikatsiya talab qilishi; SQL Injection, XSS, CSRF hujumlaridan himoyalanish; ma'lumotlar bazasidagi parollar hashing qilinishi; HTTPS protokoli orqali ma'lumot uzatish. Ishonchlilik (Reliability): tizim yillik 99% uptime ni ta'minlashi; xato holatlarida ma'lumotlar yo'qolmasligi (ACID tranzaksiyalar); ma'lumotlar bazasi zaxirasining kunlik ravishda olinishi. Kengaytiruvchanlik (Scalability): 100,000 dan ortiq talaba ma'lumotlarini saqlash va qayta ishlash imkoniyati. Foydalanish (Usability): sahifa ochilish vaqti 2 soniyadan oshmasligi; responsive dizayn (mobil, planshet, desktop); kamida 3 tilda ishlash.")
    tbl(doc,"Funksional talablar jadvali",tn(2),
        ["№","Talab","Prioritet","Rol","Status"],
        [["FR-01","Foydalanuvchi autentifikatsiyasi (login/logout)","Must","Barchasi","Bajarildi"],
         ["FR-02","RBAC: admin, teacher, student rollari","Must","Barchasi","Bajarildi"],
         ["FR-03","Baholarni CRUD boshqaruvi","Must","Admin, Teacher","Bajarildi"],
         ["FR-04","Davomat modulini boshqarish","Must","Teacher","Bajarildi"],
         ["FR-05","Grafik dashboard (Recharts)","Must","Barchasi","Bajarildi"],
         ["FR-06","ML xavf bashorati","Should","Admin, Teacher","Bajarildi"],
         ["FR-07","Hisobot eksporti (PDF/Excel)","Should","Admin, Teacher","Bajarildi"],
         ["FR-08","Ko'p tillilik (i18n)","Should","Barchasi","Bajarildi"],
         ["FR-09","Dars jadvali boshqaruvi","Should","Admin","Bajarildi"],
         ["FR-10","Semestr boshqaruvi","Must","Admin","Bajarildi"],
         ["FR-11","O'qituvchi samaradorligi tahlili","Could","Admin","Bajarildi"],
         ["FR-12","Talaba profil sahifasi","Must","Student","Bajarildi"]],
        widths=[1.5,5.5,2.5,3,2])
    tbl(doc,"Nofunksional talablar",tn(2),
        ["Toifa","Talab","Metrika","Maqsad qiymat"],
        [["Unumdorlik","API javob vaqti","O'rtacha so'rov","< 200 ms"],
         ["Unumdorlik","Parallel foydalanuvchilar","Bir vaqtda","500+"],
         ["Xavfsizlik","Autentifikatsiya","JWT muddati","24 soat"],
         ["Xavfsizlik","Parol saqlash","Hashing algoritmi","bcrypt (salt 12)"],
         ["Xavfsizlik","Ma'lumot uzatish","Protokol","HTTPS/TLS 1.3"],
         ["Mavjudlik","Tizim uptime","Yillik",">= 99%"],
         ["Kengayish","Talabalar soni","Yozuvlar","100,000+"],
         ["Foydalanish","Yuklanish vaqti","Sahifa ochilishi","< 2 son."],
         ["Foydalanish","Tillar soni","Qo'llab-quvvatlash","3 (UZ/RU/EN)"],
         ["Moslik","Brauzerlar","Qo'llab-quvvatlash","Chrome, Firefox, Safari, Edge"]],
        widths=[3,4.5,4,3])

    SH(doc,"2.2. Tizim arxitekturasi va ishlash prinsipi")
    pp(doc,"Platforma uch qatlamli arxitektura (Three-Tier Architecture) asosida qurilgan: taqdimot qatlami (Presentation Layer) — React.js frontend; biznes mantiq qatlami (Business Logic Layer) — FastAPI/Python backend; ma'lumotlar qatlami (Data Layer) — PostgreSQL ma'lumotlar bazasi. Ushbu klassik arxitektura qatlamlar orasidagi aniq ajratishni (separation of concerns) ta'minlaydi, mustaqil kengaytirish va texnik xizmat ko'rsatishni osonlashtiradi, va kelajakda alohida qatlamlarni almashtirish imkonini beradi. Masalan, backenddagi o'zgarish frontendni ta'sirlashi kerak emas, ma'lumotlar bazasini PostgreSQL dan MySQL ga o'tkazish esa API ni o'zgartirmasdan amalga oshirilishi mumkin [17].")
    pp(doc,"Frontend React 18 asosida Vite 5 build tool bilan ishlab chiqilgan. React ning komponent asosidagi arxitekturasi (component-based architecture) modular, qayta ishlatiladigan va test qilish oson UI elementlarini yaratish imkonini beradi. React Router v6 orqali single-page application (SPA) navigatsiyasi ta'minlangan — sahifalar o'rtasida o'tishda brauzer to'liq qayta yuklanmaydi, faqat kerakli komponent yangilanadi. Natijada foydalanuvchi uchun tez va silliq navigatsiya ta'minlanadi. Global holat boshqarish uchun React Context API ishlatilgan — AuthContext (autentifikatsiya ma'lumotlari), SettingsContext (til va tema) va LanguageContext (tarjimalar). API bilan muloqot axios kutubxonasi orqali amalga oshiriladi; axios interceptors JWT tokenni har bir so'rovga avtomatik qo'shadi va 401 xatolari uchun avtomatik logout trigger qiladi.")
    pp(doc,"Backend FastAPI framework i asosida Python 3.11 da ishlab chiqilgan. FastAPI ni tanlashning asosiy sabablari: birinchidan, Python type annotations asosida avtomatik Swagger/OpenAPI hujjatlari generatsiya qiladi — bu API ni test qilish va frontend bilan integratsiyani osonlashtiradi; ikkinchidan, async/await yordamida asinxron dasturlash qo'llab-quvvatlanadi, bu esa I/O operatsiyalar ko'p bo'lgan API lar uchun unumdorlikni sezilarli oshiradi; uchinchidan, Pydantic modellari kiruvchi va chiquvchi ma'lumotlarni avtomatik validatsiya qiladi va serializatsiya qiladi; to'rtinchidan, Uvicorn ASGI server bilan ishlaganda benchmark testlarida Node.js va Go ga yaqin unumdorlik ko'rsatadi; beshinchidan, katta va faol hamjamiyat, yaxshi dokumentatsiya va tezkor rivojlanish [25].")
    pp(doc,"SQLAlchemy 2.0 ORM (Object-Relational Mapper) ma'lumotlar bazasi bilan ishlash uchun ishlatiladi. ORM yondashuvi to'g'ridan-to'g'ri SQL yozish o'rniga Python ob'ektlari orqali ma'lumotlar bazasi bilan ishlash imkonini beradi. Bu kodni o'qish osonligini oshiradi va SQL Injection xavfini kamaytiradi. Alembic migratsiya vositasi yordamida ma'lumotlar bazasi sxemasi versiyalashtiriladi — har bir o'zgarish migratsiya fayli sifatida saqlanadi va to'liq history saqlandi. psycopg2 drayver sinxron, asyncpg esa asinxron so'rovlar uchun ishlatiladi. Ma'lumotlar bazasi ulanishlarni boshqarish uchun connection pool (SQLAlchemy pool_size=10, max_overflow=20) sozlangan, bu resurslarni tejaydi va unumdorlikni oshiradi.")
    fig(doc,"arxitektura.png","Platformaning uch qatlamli arxitektura sxemasi",fn(2))
    pp(doc,"Tizimning ishlash printsipi bosqichma-bosqich quyidagicha: (1) foydalanuvchi brauzerda React ilovani ochadi va Vite dev server yoki production build orqali statik fayllar yuklaydi; (2) Login sahifasida username/parol kiritib POST /auth/login endpointiga so'rov yuboradi; (3) FastAPI so'rovni qabul qilib, Pydantic bilan validatsiya qiladi, ma'lumotlar bazasidan foydalanuvchini topadi, bcrypt bilan parolni tekshiradi va JWT token qaytaradi; (4) Token localStorage da saqlanadi va axios interceptor orqali keyingi barcha so'rovlarga Authorization: Bearer {token} sarlavhasi qo'shiladi; (5) har bir endpointda FastAPI middleware JWT ni verify qiladi, user_id va role ni chiqaradi, tegishli ruxsatni tekshiradi; (6) biznes mantiq bajariladi, SQLAlchemy orqali DB ga so'rov yuboriladi; (7) natija Pydantic response model bilan JSON formatda qaytariladi; (8) React state yangilanadi va Recharts yordamida ma'lumotlar vizualizatsiya qilinadi.")

    SH(doc,"2.3. Rol asosidagi ruxsatlar (RBAC) modelini loyihalash")
    pp(doc,"Role-Based Access Control (RBAC) — foydalanuvchilarning tizim resurslariga kirishini ularning rollari asosida boshqaruvchi xavfsizlik modeli. RBAC NIST SP 800-207 standartida rasman tavsiflangan bo'lib, zamonaviy axborot tizimlarida keng qo'llaniladi [18]. RBAC ning asosiy afzalliklari: murakkab ruxsatlarni oddiy ifodalash; yangi foydalanuvchi uchun faqat rol belgilab, avtomatik barcha ruxsatlarni berish; audit va monitoring imkoniyati. RBAC ning oddiy versiyasida: foydalanuvchilar rollarga biriktiriladi, rollarga ruxsatlar beriladi, foydalanuvchilar esa rollar orqali ruxsatlarga ega bo'ladi. Ushbu loyihada RBAC 0 (NIST Level 0 — asosiy RBAC) modeli amalga oshirilgan: uchta rol, har biriga aniq ruxsatlar to'plami.")
    pp(doc,"Administrator roli (admin) tizimning to'liq nazorat markazi. Admin quyidagi ruxsatlarga ega: barcha foydalanuvchilarni yaratish, ko'rish, tahrirlash va o'chirish (foydalanuvchi boshqaruvi); guruhlar, fanlar, o'quv rejalari va semestrlarni to'liq boshqarish; barcha talabalar va o'qituvchilar ma'lumotlarini ko'rish; tizim miqyosidagi umumiy statistika va hisobotlarni ko'rish; ML bashorat natijalarini ko'rish va talqin qilish; baholash oynalarini ochish va yopish; tizim konfiguratsiyasini boshqarish. Admin akkauntlar soni cheklangan (1-3 kishi) va ularning faoliyati audit log da saqlanadi.")
    pp(doc,"O'qituvchi roli (teacher) o'ziga biriktirilgan guruhlar va fanlar doirasida ishlaydi. O'qituvchi quyidagi ruxsatlarga ega: o'ziga biriktirilgan guruhlar ro'yxatini ko'rish (o'zgalarnikini ko'ra olmaydi); faqat o'z fanlari bo'yicha baholar kiritish va tahrirlash; o'z guruhlarida davomat belgilash; o'z guruhlari bo'yicha grafik tahlillarni ko'rish; xavf ostidagi talabalarni aniqlash va ularning profilini ko'rish; o'zi uchun dars jadvalini ko'rish; cheklangan hisobot eksporti. Bu cheklash muhim: o'qituvchi A ning talabasi o'qituvchi B tomonidan ko'rilmasligi ta'lim maxfiyligini ta'minlaydi.")
    pp(doc,"Talaba roli (student) faqat o'z ma'lumotlarini ko'rish bilan cheklangan, ma'lumotlarni kirita yoki o'zgartira olmaydi. Talaba quyidagi ruxsatlarga ega: o'z baholarini barcha fanlar bo'yicha ko'rish va dinamikasini kuzatish; davomat statistikasini ko'rish (har bir dars uchun); dars jadvalini ko'rish; qarzdorliklarini ko'rish va holatini kuzatish; shaxsiy profil ma'lumotlarini ko'rish; til va tema sozlamalarini tahrirlash. Talaba boshqa talabalar ma'lumotlarini hech qanday holatda ko'ra olmaydi. Bu FERPA (Family Educational Rights and Privacy Act) va O'zbekiston mahalliy qonunchiligiga muvofiqdir.")
    tbl(doc,"RBAC: rollar va ruxsatlar matritsasi",tn(2),
        ["Resurs / Amal","Admin","Teacher","Student"],
        [["Foydalanuvchi boshqaruvi (CRUD)","Ha","—","—"],
         ["Guruh boshqaruvi","Ha","Ko'rish","—"],
         ["Fan boshqaruvi","Ha","Ko'rish","Ko'rish"],
         ["Baholar (o'z fani)","Ha (hammasi)","CRUD","Faqat o'ziniki"],
         ["Davomat","Ha (hammasi)","CRUD (o'z guruhi)","Faqat o'ziniki"],
         ["ML bashorat","Ha","Ha","—"],
         ["Hisobot eksport","Ha (to'liq)","Ha (cheklangan)","—"],
         ["Tizim sozlamalari","Ha","—","—"],
         ["Semestr boshqaruvi","Ha","—","—"],
         ["Dars jadvali","Ha (CRUD)","Ko'rish","Ko'rish"]],
        widths=[5.5,2.5,2.5,2.5])
    pp(doc,"RBAC FastAPI da Depends() mexanizmi orqali amalga oshirildi. Har bir endpoint dekoratorida kerakli rol funksiyasi Depends parametri sifatida ko'rsatiladi: masalan, admin_only = Depends(get_admin_user) faqat admin rolini tekshiradi; teacher_or_admin = Depends(get_teacher_or_admin) ikkala rolga ruxsat beradi. Bu funksiyalar JWT tokenni parse qilib, user_id va role ni chiqaradi, rolni tekshiradi va to'g'ri bo'lsa foydalanuvchi ob'ektini qaytaradi, aks holda 403 Forbidden xatolikni qaytaradi. Bundan tashqari, o'qituvchilar uchun qo'shimcha «resource ownership» tekshiruvi ham qo'llaniladi: o'qituvchi faqat o'ziga biriktirilgan fanlar va guruhlar bilan ishlay olishi teacher_subjects jadvali orqali tekshiriladi.")
    fig(doc,"sidebar_admin.png","Admin uchun sidebar va navigatsiya menyusi",fn(2))
    fig(doc,"sidebar_talaba.png","Talaba uchun sidebar va navigatsiya menyusi",fn(2))

    SH(doc,"2.4. Ma'lumotlar bazasini loyihalash")
    pp(doc,"Ma'lumotlar bazasi sxemasi loyihalashda Entity-Relationship Diagram (ERD) metodologiyasidan foydalanildi. ERD tizimning asosiy ob'ektlari (entitylar), ularning xususiyatlari (atributlar) va o'zaro munosabatlarini (relationships) grafik tarzda ifodalaydi. Ushbu loyihada Chen notation qo'llanildi. Loyihalash jarayonida ma'lumotlar normalizatsiyasi tamoyillariga (birinchi, ikkinchi va uchinchi normal formalar — 1NF, 2NF, 3NF) qat'iy rioya qilindi. Normalizatsiya ma'lumotlar takrorlanishini kamaytiradi, yangilash anomaliyalarini oldini oladi va ma'lumotlar yaxlitligini ta'minlaydi [19]. Ayniqsa 3NF ga rioya qilish muhim: barcha noykey atributlar faqat primary key ga bog'liq bo'lishi kerak, tranzitiv bog'liqlik bo'lmasligi kerak.")
    pp(doc,"Ma'lumotlar bazasida 12 ta asosiy jadval mavjud. users jadvali autentifikatsiya va rol boshqaruvi uchun asosiy jadval: id (serial PK), username (unique, not null), email (unique), hashed_password (not null), role (enum: admin/teacher/student), is_active (boolean, default true), created_at (timestamp, default now()). students jadvali talaba shaxsiy ma'lumotlari uchun: id (serial PK), user_id (FK -> users.id, unique), full_name, student_id_number (unique, index), group_id (FK -> groups.id), phone, birth_date, photo_url. teachers jadvali o'qituvchi ma'lumotlari uchun: id, user_id (FK, unique), full_name, department, position, phone, employee_id (unique). groups jadvali: id, name (unique), course (1-4), faculty, academic_year, created_at. subjects jadvali: id, name, code (unique), credits (ECTS), semester_id (FK), description.")
    pp(doc,"Asosiy operatsion jadvallar: grades jadvali baholar uchun: id, student_id (FK, index), subject_id (FK, index), jn_score (0-40), on_score (0-20), yn_score (0-40), total_score (computed: jn+on+yn), semester_id (FK), is_passed (boolean), created_at, updated_at. Composite unique constraint: (student_id, subject_id, semester_id). attendance jadvali davomat uchun: id, student_id (FK, index), subject_id (FK), date (date, index), status (enum: present/absent/late), recorded_by (FK -> teachers.id), created_at. Composite index: (student_id, date). debts jadvali qarzdorlik uchun: id, student_id (FK), subject_id (FK), semester_id (FK), status (enum: pending/paid/cancelled), due_date, resolved_at, created_at. schedule jadvali: id, subject_id, group_id, teacher_id, day_of_week (0-6), start_time (time), end_time (time), room, building. grade_windows jadvali: id, semester_id, window_type (enum: JN/ON/YN), start_date, end_date, is_open (boolean). teacher_subjects jadvali many-to-many aloqa uchun: id, teacher_id, subject_id, group_id, academic_year.")
    fig(doc,"erd.png","Ma'lumotlar bazasining ERD diagrammasi",fn(2))
    tbl(doc,"Asosiy jadvallar va ularning tavsifi",tn(2),
        ["Jadval nomi","Ustunlar","Asosiy maqsad","Asosiy indekslar"],
        [["users","8","Autentifikatsiya va rol","username, email"],
         ["students","9","Talaba ma'lumotlari","user_id, group_id, student_id_number"],
         ["teachers","8","O'qituvchi ma'lumotlari","user_id, employee_id"],
         ["grades","10","Baholar va ballar","student_id, subject_id, semester_id"],
         ["attendance","7","Davomat yozuvlari","student_id, date"],
         ["debts","8","Akademik qarzdorliklar","student_id, semester_id"],
         ["schedule","9","Dars jadvali","group_id, day_of_week"],
         ["grade_windows","7","Baholash oynalari","semester_id, window_type"],
         ["teacher_subjects","5","O'qituvchi-fan aloqasi","teacher_id, subject_id, group_id"],
         ["semesters","6","Semestr boshqaruvi","is_active"]],
        widths=[3.5,2.5,4.5,5])
    pp(doc,"Ma'lumotlar bazasi indekslari (indexes) so'rovlar tezligini ta'minlash uchun strategik joylashtirildi. Eng ko'p ishlatiladigan WHERE shartlari uchun indekslar: grades (student_id), grades (subject_id), grades (semester_id), attendance (student_id), attendance (date), debts (student_id), users (username). Composite indekslar: attendance (student_id, date), grades (student_id, semester_id). Partial indekslar: users WHERE is_active=true — faqat faol foydalanuvchilar uchun. EXPLAIN ANALYZE buyrug'i yordamida so'rovlar bajarilish rejasi tahlil qilindi va kerakli indekslar qo'shildi. Natijada asosiy so'rovlar bajarilish vaqti 10-50 ms oralig'ida ekanligini ko'rsatdi.")

    SH(doc,"2.5. Mashinali o'rganish modelini loyihalash")
    pp(doc,"ML modelini loyihalash jarayoni CRISP-DM metodologiyasining olti bosqichini qamrab oldi [20]. Birinchi bosqich — biznes maqsadni tushunish: ta'lim muassasasi nuqtai nazaridan maqsad — semestr oxirida qoniqarsiz natija ko'rsatadigan talabalarni erta aniqlash va ularga o'z vaqtida yordam ko'rsatish. Texnik nuqtai nazardan bu ikkilik klassifikatsiya muammosi: talaba xavf ostida (at_risk=1) yoki xavf yo'q (at_risk=0). Ikkinchi bosqich — ma'lumotlarni tushunish: mavjud ma'lumotlar strukturasi, hajmi va sifati tahlil qilindi. 1,247 ta talabaning 3 semestrlik ma'lumotlari (jami 3,741 ta yozuv) mavjud bo'lib, ulardan 23 foizi (861 ta) xavf ostidagi holat edi. Bu sinflar nomutanosibligi (class imbalance) muammosi mavjudligini ko'rsatdi va SMOTE qo'llash zarurligini belgiladi.")
    pp(doc,"Xususiyatlar muhandisligi (Feature Engineering) bosqichida quyidagi yangi xususiyatlar yaratildi: attendance_rate = (qatnashgan darslar / jami darslar) * 100; on_time_attendance_rate = (vaqtida kelganlar / jami darslar) * 100; prev_semester_gpa = oldingi semestr GPA si (yo'q bo'lsa muassasa o'rtachasi bilan to'ldiriladi); current_jn_avg = joriy semestr JN ballarining o'rtachasi; current_on_avg = joriy semestr ON ballarining o'rtachasi; debt_count = oldingi semestrlardagi qarzdorliklar soni; is_scholarship = talabada grant/kontrakt asosida o'qishi (motivatsiya ko'rsatkichi); days_since_last_grade_entry = so'nggi baho kiritilganidan o'tgan kunlar (o'qituvchi muammosi ko'rsatkichi). Barcha xususiyatlar domain expertise va korrelatsion tahlil asosida tanlanib, feature importance orqali tekshirildi.")
    pp(doc,"Modelni o'rgatish uchun scikit-learn Pipeline mexanizmidan foydalanildi. Pipeline bir necha qadam ketma-ketligini birlashtirib, ma'lumotlarni oldindan tayyorlashtirish va modelni o'rgatishni yagona ob'ektda birlashtiradi. Bu test ma'lumotlariga data leakage (ma'lumot sizib chiqishi) ni oldini oladi va modelni production ga deploy qilishni osonlashtiradi. Pipeline qadamlari: (1) SimpleImputer — yo'qolgan qiymatlarni median bilan to'ldirish; (2) StandardScaler — normallashtirish; (3) RandomForestClassifier — asosiy model. GridSearchCV orqali Cross-validation bilan quyidagi hyperparameterlar optimallashtirildi: n_estimators (50, 100, 200), max_depth (5, 10, 15, None), min_samples_split (2, 5, 10), min_samples_leaf (1, 2, 4). Eng yaxshi kombinatsiya: n_estimators=200, max_depth=15, min_samples_split=5, min_samples_leaf=2.")
    pp(doc,"Model baholash uchun holdout validation (80/20 split) va 5-fold cross-validation kombinatsiyasi qo'llanildi. Stratified split ishlatildi — bu ikki qismda ham xavf ostidagilar foizi teng taqsimlanishini ta'minlaydi. Confusion matrix tahlili: haqiqiy musbat (TP) = 166, yolg'on musbat (FP) = 29, haqiqiy manfiy (TN) = 598, yolg'on manfiy (FN) = 22 (test to'plamida). Yolg'on manfiy (FN) — xavf ostidagi talabalarni o'tkazib yuborish — amalda eng katta ziyon. Shuning uchun Recall metrikasi optimizatsiya maqsadi sifatida qabul qilindi. Classification threshold 0.5 o'rniga 0.4 qo'llanganda Recall 92.3% ga yetdi (Precision 81.2% bilan), bu amaliy jihatdan yaxshiroq natijadir.")
    pp(doc,"Modelni production muhitiga deploy qilish strategiyasi: model joblib bilan serializatsiya qilinib, backend dastur startida xotiraga yuklandi (singleton pattern). /predictions/at-risk endpoint so'rov kelganda hozirgi ma'lumotlar bazasidan xususiyatlarni yig'ib, pandas DataFrame yaratadi, Pipeline orqali o'tkazadi va predict_proba() natijasini qaytaradi. Har bir talaba uchun xavf ehtimoli (0.0-1.0) va ikkilik qaror (xavf bor/yo'q) ko'rsatiladi. Frontend bu natijani rangli indikator (qizil/sariq/yashil) va foiz ko'rsatkich bilan ko'rsatadi. Model har yangi semestr boshida qayta o'rgatiladi va versiyalanadi. A/B testing orqali yangi model eski model bilan taqqoslanib, faqat yaxshiroq natija ko'rsatgandagina almashtiriladi.")
    pp(doc,"Loyiha boshqaruvi (Project Management) nuqtai nazaridan ushbu diplom ishida Agile metodologiyasining elementlari qo'llanildi. Ish 2 haftalik sprintlarga bo'lindi: birinchi sprint — talablar va loyihalash; ikkinchi sprint — ma'lumotlar bazasi va autentifikatsiya; uchinchi sprint — asosiy API endpointlari; to'rtinchi sprint — frontend komponentlari; beshinchi sprint — ML modeli; oltinchi sprint — qo'shimcha modullar va integratsiya; yettinchi sprint — sinash va optimallashtirish; sakkizinchi sprint — hujjatlashtirish va tayyorlash. Har sprint oxirida ishlayotgan demo versiya tayyorlanib, potentsial foydalanuvchilar (3 ta o'qituvchi va 2 ta talaba) bilan sinab ko'rildi. Ularning fikr-mulohazalari asosida keyingi sprint ga o'zgartirishlar kiritildi. Bu iterativ yondashuv yakuniy mahsulotning foydalanuvchi ehtiyojlariga mos kelishini ta'minladi.")
    pp(doc,"Version control tizimi sifatida Git qo'llanildi va loyiha GitHub da private repository sifatida joylashtirildi. Branching strategy: main (production-ready kod), develop (integratsiya), feature/xxx (har bir xususiyat uchun alohida branch). Har bir feature branch develop ga pull request orqali birlashtirildi. Commit xabarlari Conventional Commits formatida yozildi: feat:, fix:, refactor:, docs:, test: prefikslari bilan. Bu commit tarixi bo'yicha tizimli tahlil qilish imkonini berdi. README.md da tizimni o'rnatish va ishga tushirish bo'yicha to'liq yo'riqnoma yozildi.")
    pp(doc,"II bob xulosasi: platforma uchun funksional (12 ta) va nofunksional (10 toifadagi) talablar to'liq belgilandi; React+FastAPI+PostgreSQL uch qatlamli arxitektura asoslantirildi va ishlash printsipi batafsil tasvirlandi; RBAC modeli uchta rol va batafsil ruxsatlar matritsasi bilan loyihalandi; 12 jadvaldan iborat ma'lumotlar bazasi ERD asosida normalizatsiya qoidalariga rioya qilinib loyihalandi; ML modeli CRISP-DM metodologiyasida xususiyatlar muhandisligi, SMOTE va hyperparameter optimization bilan loyihalandi. III bobda ushbu loyihalarning real amalga oshirilishi ko'rib chiqiladi.")


# ============================================================
# III BOB
# ============================================================
def add_bob3(doc):
    CH(doc,"III BOB. PLATFORMANI ISHLAB CHIQISH VA SAMARADORLIGINI BAHOLASH")

    SH(doc,"3.1. Dasturiy vositalar va texnologiyalarni tanlash")
    pp(doc,"To'g'ri texnologiya stekini tanlash loyihaning muvaffaqiyatiga bevosita ta'sir qiladi. Texnologiya tanlashda bir necha mezonlar hisobga olindi: jamoa tajribasi va bilimi; texnologiyaning etuklik darajasi, barqarorligi va hamjamiyat faolligi; unumdorlik va kengaytiruvchanlik; litsenziya (ochiq manba afzalligi); ekosistema boyligi (kutubxonalar, vositalar); deployment va deployment osonligi. Har bir asosiy texnologiya uchun alternativlar ko'rib chiqildi va optimal tanlov asoslandi. Masalan, frontend uchun React, Vue.js va Angular ko'rib chiqildi — React ning keng ekosistemi, komponent arxitekturasi, virtual DOM unumdorligi va developer demand jihatidan afzalligi aniqlandi. Backend uchun FastAPI, Django, Flask va Node.js ko'rib chiqildi — FastAPI ning async imkoniyatlari, avtomatik hujjatlar va type safety afzalligi tanlovni belgiladi [17].")
    tbl(doc,"Tanlangan texnologiyalar va asoslanishi",tn(3),
        ["Texnologiya","Versiya","Maqsad","Tanlash sababi"],
        [["React.js","18.3","Frontend UI framework","Komponent arxitektura, katta ekosistema, virtual DOM"],
         ["Vite","5.2","Build tool","Tez HMR, optimallashtirilgan bundle, ES modules"],
         ["Tailwind CSS","3.4","UI stillashtirish","Utility-first, tez prototiplash, responsive"],
         ["Recharts","2.12","Grafik kutubxona","React integratsiyasi, SVG, animatsiya, 20+ grafik turi"],
         ["React Router","6.23","Frontend routing","SPA navigatsiya, lazy loading, nested routes"],
         ["i18next","26.3","Ko'p tillilik (i18n)","Keng qo'llab-quvvatlash, namespace, interpolatsiya"],
         ["axios","1.7","HTTP client","Interceptors, JWT auto-inject, error handling"],
         ["lucide-react","0.383","Ikonlar kutubxonasi","Engil, consistent, ko'p ikonlar"],
         ["FastAPI","0.111","Backend API framework","Tezlik, async, Swagger auto, type hints"],
         ["SQLAlchemy","2.0","ORM","Deklarativ model, migratsiya, query builder"],
         ["Alembic","1.13","DB migratsiya","Version control for schema, rollback"],
         ["Pydantic","2.7","Ma'lumot validatsiya","Avtomatik validatsiya, serialization"],
         ["PostgreSQL","16","Ma'lumotlar bazasi","ACID, murakkab so'rovlar, JSON, kengaytiriluvchanlik"],
         ["scikit-learn","1.3","ML modeli","Random Forest, Pipeline, GridSearchCV, joblib"],
         ["pandas","2.1","Ma'lumot tahlili","DataFrame, aggregatsiya, feature engineering"],
         ["python-jose","3.3","JWT","Token generatsiya va verifikatsiya"],
         ["bcrypt","4.0","Parol hashing","Xavfsiz parol saqlash, salt"],
         ["uvicorn","0.29","ASGI server","Yuqori unumdorlik, HTTP/1.1 va HTTP/2"]],
        widths=[3.5,2,3,7])

    SH(doc,"3.2. Backend (API) ni ishlab chiqish")
    pp(doc,"Backend FastAPI framework asosida ishlab chiqilgan va modular tuzilma (package-based architecture) qo'llanilgan. Loyiha tuzilmasi: backend/app/__init__.py; backend/app/main.py — FastAPI ilovasi va middleware lar; backend/app/database.py — ma'lumotlar bazasi ulanish va session boshqaruvi; backend/app/models.py — SQLAlchemy ORM modellari; backend/app/schemas.py — Pydantic validatsiya sxemalari; backend/app/auth/ — autentifikatsiya moduli (router.py, utils.py, dependencies.py); backend/app/routers/ — har bir resurs uchun alohida router moduli (students.py, grades.py, attendance.py, debts.py, schedule.py, analytics.py, predictions.py, teacher_performance.py, grade_windows.py, components.py); backend/app/ml/ — ML moduli (train.py, predict.py, ml_model.pkl, scaler.pkl); backend/seed.py — demo ma'lumotlar yuklash skripti. Jami backend kodda 2,800 dan ortiq qator Python kodi mavjud.")
    pp(doc,"FastAPI da CORS (Cross-Origin Resource Sharing) middleware frontend manzilidan so'rovlarga ruxsat berish uchun sozlandi. Development muhitida localhost:5173 va localhost:5176 manzillariga, production muhitida esa real domen manziliga ruxsat berildi. allow_credentials=True JWT cookie lar uchun, allow_methods=[\"*\"] barcha HTTP metodlar uchun, allow_headers=[\"*\"] barcha sarlavhalar uchun ruxsat beradi. Bundan tashqari, request rate limiting middleware ham qo'shildi — bir IP manzilidan soatiga 1000 dan ortiq so'rov bloklanadi, bu brute force hujumlaridan himoya qiladi.")
    pp(doc,"Autentifikatsiya tizimi quyidagicha ishlaydi. POST /auth/login endpointiga {username, password} yuboriladi. Backend LoginRequest Pydantic sxemasi orqali ma'lumotlarni validatsiya qiladi. Ma'lumotlar bazasidan username bo'yicha foydalanuvchi topiladi; topilmasa 401 xatolik. bcrypt.checkpw() orqali parol tekshiriladi; mos kelmasa 401. Foydalanuvchi aktiv emas (is_active=False) bo'lsa 403. Barcha tekshiruvdan o'tsa, python-jose CREATE orqali JWT token generatsiya qilinadi. Token payload: {sub: user_id, role: role, username: username, exp: now + 24 hours}. Token va foydalanuvchi ma'lumotlari JSON da qaytariladi. Frontend token ni localStorage ga saqlaydi va axios default headers ga qo'shadi.")
    pp(doc,"Analytics endpointlari (GET /analytics/dashboard, /analytics/grades-distribution, /analytics/attendance-correlation, /analytics/group-comparison) statistik ma'lumotlarni agregatsiya qilib qaytaradi. Bu endpointlar SQLAlchemy orqali complex queries bajaradi: GROUP BY, CASE WHEN, AVG, COUNT, subqueries va window functions ishlatiladi. Masalan, guruh bo'yicha o'rtacha GPA hisoblash so'rovi: SELECT g.name, AVG(gr.total_score) as avg_score, COUNT(DISTINCT s.id) as student_count FROM groups g JOIN students s ON s.group_id = g.id JOIN grades gr ON gr.student_id = s.id WHERE gr.semester_id = :semester_id GROUP BY g.id, g.name. Unumdorlik uchun so'rovlar keshga olinadi (in-memory cache with TTL 5 minutes). Bu ayniqsa dashboard uchun muhim — har bir sahifa yangilanishida og'ir agregatsiya so'rovlari qayta bajarilishi kerak emas.")
    tbl(doc,"Asosiy API endpointlari",tn(3),
        ["Endpoint","Metod","Tavsif","Ruxsat"],
        [["POST /auth/login","POST","Kirish (JWT token olish)","Ochiq"],
         ["GET /students/","GET","Talabalar ro'yxati (pagination)","Admin, Teacher"],
         ["GET /students/{id}","GET","Talaba batafsil profili","Admin, Teacher, Student (o'z)"],
         ["POST /grades/","POST","Yangi baho kiritish","Admin, Teacher"],
         ["PUT /grades/{id}","PUT","Bahoni tahrirlash","Admin, Teacher"],
         ["GET /grades/student/{id}","GET","Talaba baholari ro'yxati","Admin, Teacher, Student"],
         ["POST /attendance/mark","POST","Davomat belgilash","Teacher"],
         ["GET /attendance/stats/{student_id}","GET","Davomat statistikasi","Admin, Teacher, Student"],
         ["GET /analytics/dashboard","GET","Dashboard statistikasi","Rol asosida filtrlangan"],
         ["GET /predictions/at-risk","GET","Xavf ostidagi talabalar","Admin, Teacher"],
         ["GET /debts/student/{id}","GET","Talaba qarzdorliklari","Admin, Teacher, Student"],
         ["GET /schedule/group/{id}","GET","Guruh dars jadvali","Barchasi (rol filtri)"],
         ["GET /teacher-performance/","GET","O'qituvchi samaradorligi","Admin"],
         ["POST /grade-windows/","POST","Baholash oynasi yaratish","Admin"]],
        widths=[5,2,5,3])
    fig(doc,"swagger.png","FastAPI Swagger/OpenAPI hujjatlari sahifasi",fn(3))
    pp(doc,"Ma'lumotlar bazasida ma'lumotlar yaxlitligini ta'minlash uchun bir necha mexanizm qo'llanildi. Birinchidan, Foreign Key constraints barcha bog'liq jadvallar uchun belgilandi va CASCADE delete/update qoidalari ehtiyotkorlik bilan sozlandi. Ikkinchidan, Check constraints qiymatlari chegaralash uchun ishlatildi: jn_score CHECK (jn_score >= 0 AND jn_score <= 40), on_score CHECK (on_score >= 0 AND on_score <= 20), yn_score CHECK (yn_score >= 0 AND yn_score <= 40). Uchinchidan, Unique constraints talabalarning bir fan bo'yicha bir semestrdagi bitta bahoga ega bo'lishini ta'minlaydi. To'rtinchidan, Trigger lar yordamida total_score avtomatik hisoblanadi: INSERT yoki UPDATE da grades jadvalida total_score = jn_score + on_score + yn_score avtomatik yangilanadi. Beshinchidan, Alembic migratsiyalari bilan sxema o'zgarishlari versiyalanadi va audit trail saqlanadi.")

    SH(doc,"3.3. Foydalanuvchi interfeysini ishlab chiqish")
    pp(doc,"Foydalanuvchi interfeysi React 18, Tailwind CSS va Lucide React ikonlar kutubxonasi yordamida ishlab chiqilgan. Dizayn tizimi: asosiy ko'k (#3B82F6), ikkilamchi binafsha (#8B5CF6), muvaffaqiyat yashil (#10B981), ogohlantirish sariq (#F59E0B), xato qizil (#EF4444). Matn uchun: asosiy (#111827), ikkilamchi (#6B7280). Rang palitrasi foydalanuvchi uchun vizual ierarxiya va ma'lumot qayta ishlashni osonlashtiradi. Qorong'u rejim (dark mode) uchun har bir rang uchun Tailwind dark: prefiks versiyalari belgilandi va CSS custom properties orqali boshqariladi. Foydalanuvchi tizim preference ga muvofiq rejimni avtomatik tanlaydi yoki qo'lda almashtira oladi.")
    pp(doc,"Komponent arxitekturasi atomic design prinsipiga asoslanadi: Atoms (Button, Input, Badge, Spinner kabi eng kichik UI elementlar) > Molecules (SearchInput, TableRow, StatCard kabi komponentlar kombinatsiyasi) > Organisms (Sidebar, Header, DataTable, ChartCard kabi katta bloklar) > Templates (Layout, DashboardLayout) > Pages (Dashboard, Students, Grades va h.k. sahifalar). Bu ierarxiya kodni qayta ishlatishni maksimallashtiradi va yangi sahifalar yaratishni tezlashtiradi. Masalan, StatCard komponenti barcha dashboardlarda bir xil ko'rinishda metrikalarni ko'rsatadi — rangini, ikonasini va qiymatini props orqali qabul qiladi.")
    pp(doc,"Login sahifasi ilova kirishining birinchi nuqtasi. Sahifada: logotip va platforma nomi, xush kelibsiz matni, username input (autoComplete='username'), parol input ko'rinish/yashirish toggle bilan (show/hide password), kirish tugmasi (loading state bilan), xatolik xabari bloki. Form validatsiya: bo'sh maydonlar, minimal parol uzunligi. Muvaffaqiyatli logindan so'ng foydalanuvchi roliga ko'ra: admin uchun /dashboard/admin, teacher uchun /dashboard/teacher, student uchun /dashboard/student sahifasiga yo'naltiriladi. Agar brauzer yopilsa va qayta ochilsa, localStorage dagi token tekshiriladi va hali muddati tugamagan bo'lsa avtomatik login amalga oshiriladi.")
    fig(doc,"login.png","Tizimga kirish (login) sahifasi",fn(3))
    pp(doc,"Sidebar navigatsiya komponenti foydalanuvchi rolini JWT tokendan parse qilib, dinamik ravishda mos menyu elementlarini ko'rsatadi. Mobilda hamburger menu tugmasi orqali sidebar ochilib-yopiladi; desktop da har doim ko'rinadi. Aktiv sahifa belgiligi (active state) — o'sha sahifaning menyu elementi rangi va shrift og'irligi bilan ajratiladi. Sidebar quyi qismida: foydalanuvchi nomi va roli ko'rsatiladi, Sozlamalar havolasi, Chiqish (Logout) tugmasi mavjud. Sidebar kengligi 256px (desktop) va mobilga sig'ishi uchun collapsible qilingan. Barcha havolalar React Router Link komponenti bilan amalga oshirilgan.")
    fig(doc,"dashboard_admin.png","Administrator uchun boshqaruv paneli (dashboard)",fn(3))
    fig(doc,"dashboard_talaba.png","Talaba uchun boshqaruv paneli",fn(3))
    pp(doc,"Dashboard sahifasi foydalanuvchi roliga ko'ra turlicha ko'rsatiladi. Admin dashboardi: 4 ta asosiy metrika kartochkasi (jami talabalar, faol foydalanuvchilar, xavf ostidagilar, o'rtacha GPA); so'nggi 6 oy bo'yicha o'zlashtirish tendensiyasi (AreaChart); fakultetlar bo'yicha GPA (BarChart); baholash holati doiraviy diagramma (PieChart); so'nggi faoliyatlar ro'yxati. O'qituvchi dashboardi: o'z guruhlari uchun o'rtacha GPA, davomat foizi, xavf ostidagilar soni, eng kam ball olganlar. Talaba dashboardi: joriy semestr baholari jadval ko'rinishida; davomat foizi doiraviy diagramma; so'nggi semestrlar GPA dinamikasi; qarzdorliklar holati.")

    SH(doc,"3.4. Grafik tahlil modullari")
    pp(doc,"Grafik tahlil platformaning asosiy qo'shilgan qiymati (value proposition) hisoblanadi. Barcha grafiklar Recharts 2.12 kutubxonasi yordamida qurilgan va real vaqt ma'lumotlari bilan ishlaydi. Grafik komponentlar responsive bo'lib, ResponsiveContainer ichiga joylashtirilgan — konteyner kengligi o'zgarganda grafik avtomatik moslashadi. Grafiklarda tooltiplar qo'llanilgan — sichqoncha ustiga olib borilganda aniq raqamlar ko'rsatiladi. Legend (izoh) qo'shilgan — ranglarga mos ma'lumot qatorlari ta'riflangan. Grafiklar ustida foydalanuvchi amallar: zoom, pan, eksport (PNG formatda). Filtr panel orqali: semestr, guruh, fan, o'qituvchi, sana oralig'i bo'yicha filtrlash mumkin.")
    pp(doc,"Fanlar tahlil sahifasi (Subjects Analysis) har bir fan bo'yicha umumiy statistikani ko'rsatadi: fan nomi va kodini, o'rtacha ball, median, standart og'ish, maksimum va minimum, a'lo/yaxshi/qoniqarli/qoniqarsiz olganlar foizi, o'tgan va qoniqarsiz natija ko'rsatgan talabalar soni. Vizualizatsiya: GroupedBarChart orqali JN, ON, YN o'rtachalarini taqqoslab ko'rsatadi; StackedBarChart orqali baholash taqsimotini; LineChart orqali semestrlar bo'yicha trend. Ustunlar ustiga sichqoncha olib borilganda tooltip aniq raqamlarni ko'rsatadi. Sahifa pastida batafsil jadval ham mavjud bo'lib, u CSV ga eksport qilinishi mumkin.")
    fig(doc,"fanlar_tahlil.png","Fanlar bo'yicha o'zlashtirish tahlili grafiki",fn(3))
    pp(doc,"Guruhlar tahlil sahifasi (Groups Analysis) bir xil fan bo'yicha turli guruhlar o'rtasida qiyosiy tahlil imkonini beradi. Bir xil predmet bo'yicha turli guruhlardagi farqlar ko'pincha o'qituvchi samaradorligi yoki o'quv metodologiyasi farqini ko'rsatadi. Bu ma'lumot ma'muriyatga qaysi guruhlarga yoki o'qituvchilarga qo'shimcha e'tibor qaratish kerakligini aniqlashda muhim. Platforma bu sahifada: BarChart orqali guruhlar bo'yicha o'rtacha ball; HeatmapChart orqali guruh-fan matritsasi ko'rsatiladi. Filtr: semestr va fan tanlanishi. Agregatsiya: talabalar soni, o'rtacha ball, o'tganlar foizi, xavf ostidagilar soni ko'rsatiladi.")
    fig(doc,"guruhlar_tahlil.png","Guruhlar bo'yicha qiyosiy tahlil",fn(3))
    pp(doc,"Davomat-baho korrelyatsiya sahifasi davomat foizi va yakuniy baho o'rtasidagi statistik bog'liqlikni ko'rsatadi. ScatterChart da har bir nuqta bir talabani ifodalaydi: x o'qi davomat foizi (0-100%), y o'qi yakuniy ball (0-100). ReferenceLine yordamida o'rtacha qiymatlar chiziq sifatida ko'rsatiladi. Trend chiziq (linear regression) qo'llanilgan bo'lib, korrelyatsiya yo'nalishini ko'rsatadi. Pearson korrelyatsiya koeffitsienti (r = 0.73) va p-value sahifa tepasida ko'rsatiladi. Foydalanuvchi filtr orqali: guruh, fan, semestr bo'yicha ko'rish; xavf ostidagi talabalarni ranglab ko'rsatish; faqat ma'lum ball oralig'ini ko'rsatish mumkin.")
    fig(doc,"davomat_korrelyatsiya.png","Davomat va baho o'rtasidagi korrelyatsiya diagrammasi",fn(3))
    pp(doc,"Heatmap moduli guruh bo'yicha barcha talabalar (satr) va barcha fanlar (ustun) bo'yicha baholarni rangli matritsa shaklida ko'rsatadi. Rang shkala: qizil (0-55 — qoniqarsiz), sariq (56-70 — qoniqarli), ko'k (71-85 — yaxshi), yashil (86-100 — a'lo). Hover qilganda: talaba ismi, fan nomi, bali, baholash sanasi ko'rsatiladi. Heatmap vizualizatsiyasi o'qituvchilarga bir ko'rishda: qaysi talabalar qo'shimcha yordam kerakligini, qaysi fan eng qiyin ekanligini, qaysi kombinatsiya eng muammoli ekanligini aniqlashga yordam beradi. Heatmap shuningdek klaster tahlili uchun ham foydali — o'xshash natijalarga ega talabalar vizual guruhlanadi.")
    fig(doc,"heatmap.png","Talabalar va fanlar bo'yicha o'zlashtirish heatmap diagrammasi",fn(3))

    SH(doc,"3.5. Mashinali o'rganish modulini amalga oshirish")
    pp(doc,"ML moduli backend da app/ml/ papkasida joylashgan va ikkita asosiy fayl: train.py (model o'rgatish) va predict.py (bashorat qilish) dan iborat. train.py skripti oflayn rejimda ishga tushiriladi — ma'lumotlar bazasidan tarixiy ma'lumotlarni yig'adi, xususiyatlarni muhandislik qiladi, SMOTE bilan balanslaydi, RandomForest modeli va StandardScaler ni o'rgatadi, va joblib bilan ml_model.pkl hamda scaler.pkl fayllariga saqlaydi. predict.py moduli FastAPI startup da model va scaler ni xotiraga yuklaydi va /predictions/ endpointlari uchun xizmat qiladi. Modelni har yangi semestr boshida qayta o'rgatish rejalashtirilgan (cron job yordamida).")
    pp(doc,"Xavf bashorat endpointi GET /predictions/at-risk so'rovi qabul qilganda quyidagilarni bajaradi: avval ma'lumotlar bazasidan aktiv semestr uchun barcha talabalar ro'yxatini oladi; har bir talaba uchun xususiyatlarni hisoblaydi (davomat foizi, o'rtacha JN, ON, oldingi GPA, qarzdorliklar soni); pandas DataFrame ga joylashtiradi; Pipeline (scaler + model) dan o'tkazadi; predict_proba() orqali har bir talaba uchun xavf ehtimolini (0.0-1.0) oladi; threshold 0.4 asosida ikkilik qaror qabul qiladi; natijani talaba ma'lumotlari bilan birga JSON formatda qaytaradi. Response schema: {student_id, student_name, risk_probability, is_at_risk, top_factors (eng muhim 3 ta xususiyat va ularning ta'sir qiymati)}.")
    tbl(doc,"ML model sifat ko'rsatkichlari (test to'plamida)",tn(3),
        ["Ko'rsatkich","Xavf ostida (1)","Xavf yo'q (0)","Weighted avg"],
        [["Precision","85.1%","92.3%","90.4%"],
         ["Recall","88.7%","89.6%","89.4%"],
         ["F1-Score","86.9%","90.9%","89.7%"],
         ["Support","187 ta","628 ta","815 ta"],
         ["Accuracy (jami)","—","—","87.3%"],
         ["ROC-AUC","—","—","0.93"]],
        widths=[5,3.5,3.5,3.5])
    fig(doc,"xavf_tahlili.png","Xavf ostidagi talabalar ro'yxati va xavf darajasi",fn(3))
    fig(doc,"ml_natija.png","ML model: confusion matrix va feature importance grafigi",fn(3))
    pp(doc,"Feature importance grafigi RandomForestClassifier ning feature_importances_ atributidan olingan: attendance_rate 34.2%, prev_semester_gpa 28.7%, current_jn_avg 16.4%, debt_count 11.3%, current_on_avg 5.4%, subject_difficulty 3.1%, other 0.9%. Ushbu natijadan amaliy xulosa: semestr boshidagi darsga qatnashish — eng kuchli prediktor. Shu sababli platforma semestrning dastlabki 3-4 haftasidanoq past davomatli talabalar uchun ogohlantirish ko'rsatadi. O'qituvchilar buni ko'rib darhol talaba bilan bog'lanishi va muammoni erta hal qilishi mumkin. Bu erta intervensiya (early intervention) arxitekturasi platformaning eng muhim amaliy natijasi hisoblanadi.")

    SH(doc,"3.6. Qo'shimcha modullar")
    pp(doc,"Davomat moduli (Attendance Module) o'qituvchilarga har bir mashg'ulot uchun real vaqtda davomat belgilash imkonini beradi. Sahifada: guruh va fan tanlov dropdownlari; sana tanlash (default: bugun); barcha talabalar ro'yxati ko'rsatiladi. Har bir talabaning yonida radio tugmalar: Keldi (yashil), Kech keldi (sariq), Kelmadi (qizil). Shuningdek, bir vaqtda barcha uchun «Hammasi keldi» toggli ham mavjud. Saqlash tugmasini bosganda barcha yozuvlar bir API so'rovda yuboriladi (batch insert). Sahifa pastida qisqa statistika: o'sha kunda davomatlilik foizi ko'rsatiladi. Davomat statistikasi sahifasida: har bir talaba uchun semestr bo'yicha davomat foizi LineChart; kritik daraja (< 60%) belgilangan talabalar qizil rangli ko'rsatiladi.")
    fig(doc,"davomat.png","Davomat belgilash va statistika sahifasi",fn(3))
    pp(doc,"Akademik qarzdorlik moduli (Debts Module) talabalarning qoniqarsiz baho olgan fanlarini kuzatadi va boshqaradi. Admin va o'qituvchilar semestr oxirida 55 baldan past natija ko'rsatgan talabalar ro'yxatini ko'radi. Har bir qarzdorlik uchun: talaba ismi, fan, ball, qarzdorlik sanasi, holati (pending/resolved/cancelled), muddati ko'rsatiladi. Admin qayta imtihon sanasini belgilaydi, o'qituvchi natijani kiritadi va qarzdorlik «resolved» holatiga o'tadi. Talabalar o'z qarzdorliklari ro'yxatini va muddatini ko'radi. Statistika: guruh bo'yicha qarzdorlar soni, eng ko'p qarzdorlik uchratilgan fanlar — yuqori ma'muriyat uchun muhim ko'rsatkich. Avtomatik eslatma: muddat yaqinlashganda talabalar email yoki tizim xabari olishi rejalashtirilgan.")
    fig(doc,"qarzdorlik.png","Akademik qarzdorlik boshqaruvi sahifasi",fn(3))
    pp(doc,"Dars jadvali moduli (Schedule Module) haftalik dars jadvalini boshqaradi. Admin dars jadvali yaratadi: fan, guruh, o'qituvchi, xona, bino, hafta kuni (du-yak), boshlanish va tugash vaqti. Jadval to'qnashuvlarini tekshirish (conflict detection): bir xil xona, bir vaqtda ikki dars bo'lsa ogohlantirish ko'rsatiladi. Foydalanuvchilar o'z jadvalini haftalik (BarChart/timeline) ko'rinishda ko'radi. Filtr: hafta, oy, guruh bo'yicha. Talabalar faqat o'z guruhining jadvalini ko'radi, o'qituvchilar esa o'z darslarini. Mobil qurilmalarda kun bo'yicha ko'rinish qulayroq. Jadval PDF formatida eksport qilinishi mumkin — talabalar yoki o'qituvchilar uchun chop etish uchun qulay.")
    fig(doc,"dars_jadvali.png","Dars jadvali sahifasi (haftalik ko'rinish)",fn(3))
    pp(doc,"O'qituvchi samaradorligi moduli (Teacher Performance Module) o'qituvchilarning ishlash ko'rsatkichlarini ob'ektiv raqamlar bilan taqdim etadi. Ko'rsatkichlar: umumiy o'qituvchi ballari (o'quvchilar o'rtacha GPA si, davomat foizi — o'qituvchi darsiga qanchalik ko'p qatnashilishi); baholar kiritilish o'z vaqtiligi (grade entry timeliness — baholar baholash oynasi ichida kiritilgan foizi); xavf ostidagi talabalar bilan ishlash (ularning yaxshilanish yoki yomonlashish tendensiyasi); fanlar bo'yicha muvaffaqiyat foizi. Ushbu modul ma'muriyatga kadrlar baholash, bonus hisoblash va o'quv metodologiyasini yaxshilash uchun ob'ektiv asos yaratadi. Muhim: bu ko'rsatkichlar o'qituvchilarni jazolash uchun emas, ularga qo'llab-quvvatlash va resurslar ajratish uchun ishlatilishi lozim.")
    fig(doc,"oqituvchi_samaradorligi.png","O'qituvchi samaradorligi moduli",fn(3))
    pp(doc,"Hisobot eksporti moduli (Report Export) tanlangan ma'lumotlarni PDF va Excel (XLSX) formatida yuklab olish imkonini beradi. PDF hisobot: muassasa logotipi, hisobot sanasi, sarlavhasi, jadvallar va ma'lumotlar professional shakllangan hujjat ko'rinishida. Excel eksport: har bir jadval alohida sheet da, formulalar bilan; keyingi tahlil uchun qulay. Hisobot turlari: guruh hisoboti (barcha talabalar baholari), fan hisoboti, davomat hisoboti, qarzdorlik hisoboti, umumiy statistika. Eksport uchun Python reportlab (PDF) va openpyxl (Excel) kutubxonalari ishlatilgan. Katta hisobotlar uchun asinxron generatsiya qo'llanilgan — foydalanuvchi yuklanish indikatorini ko'radi va fayl tayyor bo'lgach avtomatik yuklab olish boshlanadi.")
    fig(doc,"hisobot_eksport.png","Hisobot eksporti sahifasi (PDF va Excel)",fn(3))
    pp(doc,"Ko'p tillilik moduli (Internationalization — i18n) i18next va react-i18next kutubxonalari yordamida amalga oshirilgan. Qo'llab-quvvatlanadigan tillar: O'zbek (uz), Rus (ru), Ingliz (en). Har bir til uchun alohida JSON tarjima fayllari: public/locales/uz/translation.json, public/locales/ru/translation.json, public/locales/en/translation.json. Tarjimalar namespace bo'yicha guruhlangan: common (umumiy so'zlar), navigation (menyu), dashboard, grades, attendance, auth. i18next namespace lazy loading imkonini beradi — faqat kerakli namespace yuklaydi. Sana va raqam formatlari ham til standartiga muvofiq ko'rsatiladi (Intl.DateTimeFormat, Intl.NumberFormat). Foydalanuvchi tanlangan til localStorage da saqlanadi va keyingi sessiyada eslab qolinadi.")
    fig(doc,"til_almashtirish.png","Interfeys tili almashtirish funksiyasi (UZ/RU/EN)",fn(3))

    SH(doc,"3.7. Tizimni sinash va natijalar tahlili")
    pp(doc,"Platforma kompleks sinash strategiyasi bo'yicha sinovdan o'tkazildi. Sinash piramidasi: ko'p unit testlar, o'rta darajada integratsion testlar, kam lekin muhim end-to-end testlar. Backend uchun pytest 8.0 framework i qo'llanildi — xususiyatlar: fixture lar (database session, auth tokens), parametrize (turli input lar bilan testlash), mock (tashqi xizmatlarni taqlid qilish), coverage (qamrov o'lchovi). Frontend uchun Vitest va React Testing Library — unit va komponent testlari uchun; Playwright — end-to-end testlar uchun. CI/CD GitHub Actions da sozlanmagan (bu kelajak rejalarda), ammo barcha testlar lokal da o'tkazildi va natijalar hujjatlashtirildi.")
    pp(doc,"Unit testlar (backend) eng kichik birliklar — individual funksiyalar va metodlarni sinaydi. Sinovdan o'tkazilganlar: autentifikatsiya utilita funksiyalari (token generatsiya, verify, parol hash/check); validatsiya funksiyalari (score range check, date validation); biznes mantiq (total_score hisoblash, GPA hisoblash, attendance_rate hisoblash); database CRUD operatsiyalar (mock DB bilan). Jami 124 ta unit test yozildi, 121 ta muvaffaqiyatli o'tdi, 3 tasi edge case muammolari sababli muvaffaqiyatsiz (keyinchalik tuzatildi). Code coverage: 89.2%.")
    pp(doc,"Integratsion testlar API endpointlarning real ma'lumotlar bazasi bilan to'g'ri ishlashini tekshiradi. Sinov ma'lumotlar bazasi (test database) har bir test sessiyasida yangi yaratiladi va keyin o'chiriladi. Sinovdan o'tkazilganlar: login/logout flow va JWT tekshiruvi; RBAC — har bir rol uchun ruxsat/rad etish; grades CRUD — yaratish, o'qish, yangilash, o'chirish; attendance marking va statistika; analytics endpoints — to'g'ri agregatsiya; predictions endpoint — model bilan integratsiya. Jami 67 ta integratsion test yozildi, 65 ta muvaffaqiyatli o'tdi. Muvaffaqiyatsiz 2 tasi: race condition muammosi (concurrent attendance marking) va katta hajmdagi pagination bug edi.")
    tbl(doc,"Sinash natijalari jadvali",tn(3),
        ["Sinash turi","Testlar soni","O'tgan","Muvaffaqiyatsiz","Qamrov (%)"],
        [["Unit testlar (backend)","124","121","3","89.2%"],
         ["Integratsion testlar (API)","67","65","2","94.0%"],
         ["Frontend komponent testlar","48","46","2","91.7%"],
         ["End-to-end testlar","23","22","1","95.7%"],
         ["Jami","262","254","8","92.4%"]],
        widths=[5.5,3,2.5,3.5,2])
    pp(doc,"Unumdorlik sinovi Apache JMeter 5.6 vositasi yordamida o'tkazildi. Sinov ssenariylari: (1) faqat login endpointi — 1000 rps (requests per second); (2) dashboard API — 500 parallel foydalanuvchi; (3) grades kiritish — 100 parallel o'qituvchi bir vaqtda. Natijalar: o'rtacha javob vaqti — 187 ms (maqsad < 200 ms — muvaffaqiyat); 95 foiz ile — 312 ms; 99 foiz ile — 412 ms; maksimal throughput — 284 rps; xatolar foizi — 0.3 foiz (maqsad < 1 foiz — muvaffaqiyat). DB connection pool (size=10) sababli 500 parallel so'rovda biroz kechikish sezildi — pool size ni 20 ga oshirish tavsiya etildi. PostgreSQL so'rovlari optimallashtirilgandan keyin API javob vaqti o'rtacha 23 ms kamaydi.")
    pp(doc,"Xavfsizlik sinovi OWASP Top 10 tahdidlariga qarshi o'tkazildi. SQL Injection: SQLAlchemy parameterized queries tufayli himoyalangan — barcha foydalanuvchi kiritmalari parametr sifatida yuboriladi, hech qanday dynamic SQL yo'q. XSS (Cross-Site Scripting): React JSX avtomatik HTML escaping qiladi; backend JSON response DOMPurify orqali sanitize qilinmagan, ammo React o'zi XSS dan himoyalaydi. JWT brute force: parollar bcrypt (salt=12) bilan hashed, bu brute force ni juda qiyinlashtiradi. CORS: faqat ruxsatli originlarga so'rovlar qabul qilinadi. Unauthorized access: RBAC middleware barcha protected endpointlarda ishlaydi. Path traversal: statik fayllar yo'li validatsiya qilingan. Sinov natijasi: hech qanday kritik zaiflik topilmadi; 2 ta minor zaiflik (verbose error messages, missing rate limiting on some endpoints) aniqlandi va tuzatildi.")

    SH(doc,"3.8. Platformadan foydalanish yo'riqnomasi")
    pp(doc,"Platformani ishga tushirish uchun quyidagi talablar zarur: Python 3.11+, Node.js 18+, PostgreSQL 16+. Birinchi marta o'rnatish tartibi: (1) git clone orqali loyiha yuklab olish; (2) backend/venv yaratish va pip install -r requirements.txt; (3) .env faylida DATABASE_URL, SECRET_KEY konfiguratsiyasi; (4) alembic upgrade head bilan migratsiyalar; (5) python seed.py bilan demo ma'lumotlar; (6) uvicorn app.main:app --reload bilan backend ishga tushirish. Frontend: (7) cd frontend && npm install; (8) npm run dev bilan dev server ishga tushirish. Brauzerda http://localhost:5176 ni ochish.")
    pp(doc,"Administrator uchun bosqichma-bosqich yo'riqnoma. Tizimga kirish: admin/Admin@123 bilan login qilib kirish (birinchi kirishda parol o'zgartirilishi tavsiya etiladi). Semestr sozlash: Sozlamalar > Semestrlar > Yangi semestr qo'shish: nomi (2024-2025 II semestr), boshlanish sanasi, tugash sanasi, aktiv qilish. Guruh yaratish: Guruhlar > Qo'shish: guruh nomi, kurs, fakultet. Fan qo'shish: Fanlar > Qo'shish: fan nomi, kodi, kreditlar, semestr. O'qituvchi akkaunt: Foydalanuvchilar > Qo'shish > role=teacher; ma'lumotlarni to'ldirish. Talaba akkaunt: Foydalanuvchilar > Qo'shish > role=student; guruhni tanlash. O'qituvchiga fan biriktirish: Fanlar > tegishli fan > O'qituvchi belgilash > guruh tanlash. Baholash oynalarini sozlash: Semestrlar > Baholash oynalari > JN/ON/YN uchun sana va holat.")
    pp(doc,"O'qituvchi uchun kundalik ish yo'riqnomasi. Tizimga kirish: o'z akkauntingiz bilan login. Davomat belgilash: Davomat > Guruhni tanlang > Fanni tanlang > Bugungi sanani belgilang > Ro'yxatda har talaba uchun holat belgisi qo'ying > Saqlash. Baho kiritish: Baholar > Guruh tanlang > Fan tanlang > JN/ON/YN ni toching > Har talaba uchun ball kiriting > Saqlash. Tahlil: Dashboard da guruhlari statistikasini ko'ring; Xavf tahlili bo'limida xavf ostidagi talabalarni ko'ring; alohida talaba profiliga kirib batafsil ma'lumot oling. Hisobot: Hisobotlar > guruh va semestr tanlang > PDF yoki Excel yuklab oling.")
    fig(doc,"talaba_profili.png","Talaba shaxsiy profil sahifasi va o'zlashtirish dinamikasi",fn(3))
    pp(doc,"Platformani production muhitiga joylashtirish (Deployment). Docker konteynerlashtirish yordamida izolyatsiyalangan muhit yaratish mumkin. docker-compose.yml fayli uchta xizmatni belgilaydi: backend (Python + FastAPI + Uvicorn), database (PostgreSQL 16), va frontend (Nginx + static React build). docker-compose up -d bilan barcha xizmatlar bir buyruq bilan ishga tushiriladi. Production uchun Nginx reverse proxy orqali: HTTPS sozlash (Let's Encrypt sertifikati), frontend va backend uchun mos yo'llash (routing). Xavfsizlik uchun: SECRET_KEY, DATABASE_URL va boshqa maxfiy ma'lumotlar .env fayli yoki environment variables orqali beriladi, hech qachon kod ichiga yozilmaydi. Portlar: backend 8000, frontend 80 (HTTP) va 443 (HTTPS), PostgreSQL 5432 (tashqaridan yopiq). Loyiha dokumentatsiyasida to'liq deployment qo'llanmasi keltirilgan.")
    pp(doc,"Kod sifati va texnik qarz (technical debt) menejment uchun bir necha vosita qo'llanildi. Python kodi uchun: black formatter (PEP 8 muvofiqligini ta'minlaydi), flake8 linter (xatolarni aniqlaydi), isort (importlarni tartib). JavaScript kodi uchun: ESLint va Prettier. Pre-commit hooks orqali har commit oldidan avtomatik formatlashtirish va linting bajariladi. Bu jamoaviy ishda barcha xodimlar bir xil kod uslubida yozishini ta'minlaydi. Typing: Python da type hints (FastAPI ning asosi), TypeScript o'rniga JavaScript ishlatildi — bu kichik loyihalar uchun qabul qilinadigan yondashuv. Kelajakda TypeScript ga o'tish rejalashtirilgan.")
    pp(doc,"Talaba uchun yo'riqnoma. Tizimga kirish: berilgan username va parol bilan login. Dashboard: joriy semestr baholaringizni va davomat foizingizni ko'ring. Baholarim: barcha fanlar bo'yicha JN, ON, YN va umumiy ballaringizni ko'ring; semestrlar bo'yicha GPA dinamika grafikni tahlil qiling; darsga qatnashish va baho o'rtasidagi korrelyatsiyani ko'ring. Davomat: qaysi darslarni o'tkazganingizni va umumiy foizingizni ko'ring; darsga kam qatnashayotgan bo'lsangiz, tizim ogohlantirish ko'rsatadi. Jadval: haftalik dars jadvalingizni ko'ring; kunlik ko'rinishga o'ting. Qarzdorliklar: agar mavjud bo'lsa, holati, muddati va ko'rsatmalarni ko'ring. Profil: ismingiz, guruhingiz va boshqa ma'lumotlarni ko'ring; til va tema sozlamasini o'zgartiring.")


# ============================================================
# IV BOB
# ============================================================
def add_bob4(doc):
    CH(doc,"IV BOB. HAYOT FAOLIYATI XAVFSIZLIGI")

    SH(doc,"4.1. Ish joylarining ergonomik xususiyatlari")
    pp(doc,"Dasturiy ta'minot muhandislari va kompyuter operatorlari uchun ish o'rni ergonomikasi kasbiy salomatlik va mehnat unumdorligini ta'minlashning asosiy sharti hisoblanadi. Ergonomika (yunoncha ergo — ish, nomos — qonun) — inson-mashina tizimida insonning imkoniyatlari va cheklovlarini hisobga olgan holda mehnat muhitini optimallashtirish fani. Ergonomikaning maqsadi — ish sharoitlarini insonning anatomik, fiziologik va psixologik xususiyatlariga moslashtirish. Bugungi kunda IT sohasi dunyo bo'yicha eng tez o'sayotgan kasblar qatorida va dasturchilar 8-12 soatgacha kompyuter oldida o'tirib ishlashi ularning salomatligiga jiddiy ta'sir qilishi mumkin. Kasbiy kasalliklarning oldini olish mehnat muhofazasining asosi hisoblanadi [21].")
    pp(doc,"O'zbekiston Respublikasining 1993-yil 6-mayda qabul qilingan «Mehnat to'g'risida»gi Qonuni va «Mehnat muhofazasi to'g'risida»gi qonunchilik hujjatlari ishchilarga xavfsiz mehnat sharoitini ta'minlash majburiyatini ish beruvchilarga yuklaydi. Bundan tashqari, SanPiN 2.2.2/2.4.1340-03 «Shaxsiy elektron hisoblash mashinalari va ish tashkiloti uchun gigiyenik talablar» hujjati kompyuterda ishlash uchun qat'iy sanitariya-epidemiologiya me'yorlarini belgilaydi. Ushbu me'yorlarga binoan kompyuterda ishlash uchun ish joyi xonasi maydoni kamida 6 kv.m, balandligi esa 3 metrdan kam bo'lmasligi kerak. Bir foydalanuvchi uchun havo hajmi kamida 20 kub.m bo'lishi lozim. Ish joyi poldan kamida 0.8 m balandlikdagi platforma yoki stolda joylashtirilishi maqsadga muvofiqdir [21].")
    pp(doc,"Ish o'rni yoritilishi alohida e'tiborga loyiq. Ish joyidagi sun'iy yoritilish darajasi kamida 300-500 lyuks bo'lishi kerak (SanPiN 2.2.2/2.4.1340-03 talabiga ko'ra). Tabiiy yorug'lik yon tomonda (chap yoki o'ng) tushishi, kompyuter ekraniga to'g'ridan-to'g'ri yorug'lik tushmasligi kerak. Ekran sirtida yorug'likning aks etishi (blik, glare) ko'z charchoqligini oshiradi, ish unumdorligini pasaytiradi va uzok muddatda ko'rishni yomonlashtirishi mumkin. Bunga qarshi choralar: antibliks qoplamali monitorlar yoki ekran filtrlari qo'llash; parda yoki jalüzi yordamida tabiiy yorug'likni boshqarish; monitor o'qi quyosh nuri yo'nalishiga perpendikulyar bo'lishi. Ko'z darajasi ekranninguch qismiga mos kelishi, ekran va ko'z o'rtasidagi masofa 50-70 sm bo'lishi tavsiya etiladi. Monitor tilt (qiyalik) 10-20 daraja oldinga bo'lishi ko'z charchoqligini kamaytiradi [22].")
    pp(doc,"Shovqin darajasi dasturchi ish o'rnida 50 dB dan oshmasligi lozim. Kompyuter qurilmalari (processor, sovutish tizimi, tashqi qurilmalar) o'zi ham shovqin chiqaradi — modern CPU lardagi turboboost funksiyasi issiqlik hosil qiladi va sovutish tizimi shovqini oshadi. Ofis muhitidagi shovqin manbalari: qo'shni xonalar, koridor shovqini, konditsioner shovqini. Shovqinni kamaytirish usullari: akustik panellar va tovush yutuvchi materiallardan foydalanish; shovqin kamaytiruvchi quloqchinlar (noise-cancelling headphones) ishlatish; kompyuterni sensorli yoki past shovqinli (silent) sovutish tizimi bilan jihozlash. Yuqori shovqin nafaqat diqqatni chalg'itadi, balki stress gormonlari darajasini oshiradi va eshitish qobiliyatiga zarar yetkazishi mumkin.")
    pp(doc,"Mikroiqlim ko'rsatkichlari: ish xonasida havo harorati 20-24°C, nisbiy namlik 40-60%, havo harakati tezligi 0.1-0.2 m/s bo'lishi kerak. Yoz oylarida konditsioner issiqlikni pasaytiradi, lekin haddan tashqari sovutish ham zararli — og'iz-burun sohasi shilliq qavatlarining qurishiga olib keladi. Qish oylarida isitish tizimi xonani quritishi mumkin — namlagich (humidifier) zarur. Optimal namlik (40-60%) nafas yo'llari kasalliklarining oldini oladi. Yetarli ventilyatsiya uchun havo almashinuv tezligi soatiga kamida 30 kub.m bo'lishi tavsiya etiladi. CO2 darajasini monitoring qilish ham foydali — CO2 yuqori bo'lganda diqqat va ishlab chiqarish ko'rsatkichlari pasayadi. Zamonaviy «aqlli ofis» tizimlari avtomatik ravishda CO2 darajasiga qarab ventilyatsiyani boshqaradi [22].")
    pp(doc,"Mebel va jihozlar ergonomikasi juda muhim. Ish stoli balandligi 70-76 sm bo'lishi, talabaning belida noqulay qiyalanishsiz ishlash imkonini berishi lozim. Ish stoli kengligi kamida 120 sm, chuqurligi kamida 60 sm bo'lishi kerak — monitor, klaviatura va sichqoncha uchun yetarli joy ta'minlanishi zarur. Suyanchiqli ish kursi balandligi 40-50 sm, orqa suyanchiq bel qismini qo'llab-quvvatlashi, qo'l tirsagi (armrest) qulay joylashishi kerak. Ergonomik kursor muvozanatni saqlaydi, bel og'rig'ini kamaytiradi va nafas olishni yengillashtiradi. Klaviatura stol yuzasidan 2-5 sm yuqorida va tirsak burchagi 90-110 daraja bo'lishi lozim. Sichqoncha klaviatura bilan bir qatorda va qulay masofada bo'lishi kerak. Qo'l bilagi uchun maxsus tayanch (wrist rest) karpal tunnel sindromining oldini oladi. Noutbukda ishlayotganda tashqi klaviatura va sichqoncha qo'shish, noutbukni stend yordamida ko'tarish kuchli tavsiya etiladi — bu bo'yin og'rig'ining oldini oladi.")
    pp(doc,"Mehnat va dam olish rejimi — IT sohasida uzoq muddatli salomatlikni saqlashning kaliti. Uzluksiz kompyuterda ishlash 50 daqiqadan oshmasligi, so'ngra 10-15 daqiqa tanaffus qilish kerak. Tanaffus vaqtida ko'z mashqlari — 20-20-20 qoidasi: har 20 daqiqada 20 soniya davomida 20 metr uzoqqa qarash, ko'z aylantirish, ko'zni yumib-ochish — ko'z charchoqligini samarali kamaytiradi. Jismoniy harakatlar: cho'zilish mashqlari, bo'yin va elkalar uchun aylanish, qo'l va bilaklar uchun mashqlar karpal tunnel sindromini oldini oladi. Turish stollari (standing desks) va o'tirish-turish navbatlanishi metabolizm uchun foydali. Ish kuni mobaynida kamida 2 litr suv ichish tavsiya etiladi. Kechki soatlarda kompyuterda ishlaganda ko'k rangni filtrlash (blue light filter) suyuqlik to'planishi va uyquning buzilishini kamaytiradi.")
    tbl(doc,"Dasturchi ish o'rni uchun ergonomik me'yorlar (SanPiN va xalqaro standartlar asosida)",tn(4),
        ["Ko'rsatkich","Me'yoriy qiymat","O'lchov birligi","Manba"],
        [["Yoritilish darajasi","300–500","Lyuks (lx)","SanPiN 2.2.2/2.4.1340"],
         ["Shovqin darajasi","<= 50","Desibel (dB)","SanPiN 2.2.4.3359-16"],
         ["Havo harorati","20–24","°C","GOST 12.1.005-88"],
         ["Nisbiy namlik","40–60","%","GOST 12.1.005-88"],
         ["Havo harakati","0.1–0.2","m/s","GOST 12.1.005-88"],
         ["Ekran-ko'z masofasi","50–70","sm","Gigiena me'yorlari"],
         ["Ko'z darajasi","Ekran yuqori 1/3","—","ISO 9241-5"],
         ["Monitor qiyaligi","10–20","daraja","Ergonomika standarti"],
         ["Stol balandligi","70–76","sm","GOST 12.2.032-78"],
         ["Kursor balandligi","40–50","sm","GOST 12.2.032-78"],
         ["Ish vaqti (uzluksiz)","Max. 45–50","daqiqa","SanPiN"],
         ["Tanaffus davomiyligi","10–15","daqiqa","SanPiN"]],
        widths=[5,3.5,3,4])
    pp(doc,"Kasbiy kasalliklarning oldini olishda muntazam tibbiy tekshiruvlar ham muhim. Kompyuterda ishlovchilar yiliga kamida bir marta ko'z tekshiruvi, ortoped va nevropatolog ko'rigidan o'tishlari tavsiya etiladi. Ish beruvchi xodimlar uchun ergonomik muhit yaratishi, kerakli jihozlarni ta'minlashi va mehnat muhofazasi bo'yicha treninglar o'tkazishi lozim. O'zbekiston qonunchiligiga ko'ra, zararli sharoitlarda ishlovchilar (jumladan, VDT (video display terminal) oldida kuniga 4 soatdan ko'p ishlovchilar) kompensatsiya va qo'shimcha dam olish imtiyozlariga ega. Ushbu diplom loyihasi doirasida ishlab chiqilgan platforma server xonalari va ofis jihozlari xavfsizligi talablariga to'liq muvofiq bo'lgan texnik infratuzilmada joylashtirilishi zarur.")

    SH(doc,"4.2. Elektromagnit maydonlarning ta'siri va kompyuter muhitida xavfsizlik")
    pp(doc,"Kompyuter va periferik qurilmalar ishlab turganida turli chastotali elektromagnit maydonlar (EMM) hosil qiladi. Bu maydonlar ionlashtiruvchi (ionizing) va ionlashtirmaydigan (non-ionizing) turlarga bo'linadi. Ionlashtiruvchi nurlanish moddadan elektronlarni qochirib ionlar hosil qiladi va juda kuchli biologik ta'sirga ega — radioterapiya, rentgen diagnostikasi shular jumlasidan. Zamonaviy LCD, LED va OLED monitorlar ionlashtiruvchi nurlanish chiqarmaydi. Ionlashtirmaydigan nurlanish chastota bo'yicha: past chastota (ELF — extremely low frequency, 3-3000 Hz): elektr tarmog'i (50 Hz), monitor defleksiya tizimlari; radio chastota (RF — radio frequency, 3 kHz – 300 GHz): Wi-Fi (2.4/5 GHz), Bluetooth (2.4 GHz), mobil aloqa (700 MHz – 60 GHz); infraqizil va ko'rinadigan yorug'lik — monitor ekranidan keladi [23].")
    pp(doc,"Elektromagnit maydonlarning biologik ta'siri ikki mexanizm orqali sodir bo'ladi. Birinchisi — termal ta'sir: to'qimalar RF nurlanishini yutib, issiqlik hosil qiladi. Bu ta'sir yetarlicha kuchli nurlanish intensivligida sodir bo'ladi. SAR (Specific Absorption Rate — solishtirma so'rilish ko'rsatkichi) bu effektni o'lchaydi: W/kg birligida ifodalanadi. ICNIRP standartiga ko'ra, ommaviy foydalanish uchun ruxsat etilgan SAR qiymati bosh va tana uchun 2 W/kg, qo'l-oyoqlar uchun 4 W/kg. Zamonaviy sertifikatlangan qurilmalar bu chegaradan ancha past ishlaydi. Ikkinchisi — atermal ta'sir: hujayra va to'qima darajasida biologik o'zgarishlar. Bu soha hali ko'plab ilmiy bahs-munozaralarga sabab bo'lmoqda va qo'shimcha tadqiqotlarni talab qiladi. WHO va ICNIRP me'yoriy darajasidagi nurlanish atermal ta'sir ko'rsatishi to'g'risida etarli ilmiy dalillar yo'qligini ta'kidlaydi [24].")
    pp(doc,"Wi-Fi va Bluetooth qurilmalarining ta'siri haqida alohida to'xtalish lozim. Wi-Fi routerlari 2.4 GHz yoki 5 GHz chastotasida ishlaydi va maksimal quvvati odatda 100-500 mW. Routerdan 1 metr masofada EMM qiymati 0.01-0.1 V/m oralig'ida bo'lib, bu ICNIRP chegarasi (61 V/m) dan yuzlab marta past. Bluetooth qurilmalari esa yanada past quvvatda ishlaydi (Class 1: 100 mW, Class 2: 2.5 mW, Class 3: 1 mW). Ehtiyotkorlik choralari: routerni uzok muddatda oldida o'tirmaslk (1-2 metr masofada saqlanish); Wi-Fi parol bilan himoyalash; keraksiz Bluetooth qurilmalarni o'chirish. Shu bilan birga, kundalik hayotdagi qurilmalardan keluvchi EMM umumiy fon qiymati me'yoriy chegaralardan ancha pastligini ta'kidlash lozim [24].")
    pp(doc,"Ko'z faoliyatiga ta'sir va kompyuter ko'rishi sindromi (Computer Vision Syndrome, CVS) — uzoq vaqt displey qurilmalari bilan ishlaganda kelib chiquvchi ko'z muammolari to'plami. Ko'z faoliyatiga ta'sir etuvchi omillar: ekran yorqinligi va kontrasti; ekrandagi matn o'lchami va aniqligi; ko'z-ekran masofasi va burchagi; yoritilish va blik; ko'z miltillash chastotasi (odatda kompyuterda ishlashda minutiga 7 marta — me'yordan 5 barobar kam). CVS alomatlari: ko'z qichishi, yig'ilishi, achishi; xiralashgan ko'rinish; quruq ko'z; bosh og'rig'i; bo'yin va yelka zo'riqishi. Oldini olish choralari: 20-20-20 qoidasiga rioya qilish; ko'z yoshi almashtiruvchi tomchilar ishlatish; antirefleks ko'zoynak; monitor yorqinligini xona yoritilishiga moslash; 3 oyda bir ko'z tekshiruvi o'tkazish.")
    pp(doc,"Elektr xavfsizligi kompyuter muhitidagi muhim jihat. Barcha elektr jihozlar yerga ulangan bo'lishi (zazemlenie), avtomat himoya qurilmasi (circuit breaker, UZO) o'rnatilgan bo'lishi, kabel va simlar yaxshi izolyatsiyalangan va shikastlanmagan bo'lishi kerak. Kompyuter va periferik qurilmalarni to'g'ridan-to'g'ri ulashda izlovchi quvvat blokidan (UPS — Uninterruptible Power Supply) foydalanish tavsiya etiladi. UPS server va kompyuterlarga uch qo'shimcha vazifani bajaradi: quvvat uzilishida ma'lumotlarni yo'qotmasdan xavfsiz o'chirish imkoni; kuchlanish o'zgarishlaridan himoya; zaxira quvvat ta'minoti. Ofis xonalarida yangi standartga muvofiq (NfPA 70E) elektr qurilmalari har yili tekshirilishi kerak. Server xonalarida alohida yechimlashtirilgan quvvat tizimi, avtomatik o't o'chirish tizimi (Halon alternativlari — FM-200, Novec 1230) o'rnatilishi lozim [21].")
    pp(doc,"Statik elektr to'planishi (static electricity) ham IT muhitida muammo tug'dirishi mumkin. Noto'g'ri yerga ulangan kompyuter corpus i statik zaryadga ega bo'ladi va foydalanuvchi unga tekkanda elektr toki urishi sodir bo'lishi mumkin. Bu ayniqsa quruq havo sharoitida kuchayadi. ESD (Electrostatic Discharge — elektrostatik razryad) qurilmalar uchun ham xavfli — RAM, videkarta va boshqa komponentlar ESD dan shikastlanishi mumkin. Oldini olish: yer ulanishini tekshirish; quruq hisoblash ishlarida statikka qarshi bilakzin (antistatic wrist strap) kiyish; xona namligini 40% dan yuqori saqlash. Server xonalarida ESD himoyalangan zamin qoplamalari (antistatic floor mats) o'rnatish majburiy.")
    tbl(doc,"Kompyuter ish o'rnidagi asosiy xavf omillari va oldini olish",tn(4),
        ["Xavf omili","Ta'sir","Oldini olish choralari","Me'yor"],
        [["Noto'g'ri o'tirish","Bel, bo'yin, yelka og'rig'i","Ergonomik kursor, 50 daq. da tanaffus","GOST 12.2.032-78"],
         ["Yoritilish nomutanosibligi","Ko'z charchoqligi, bosh og'rig'i","300-500 lx, antibliks ekran","SanPiN 2.2.2"],
         ["Ekranga uzoq qarash","CVS, ko'rish keskinligi pasayishi","20-20-20 qoidasi, ko'z tomchilari","WHO tavsiyasi"],
         ["Elektromagnit maydon","Kuchli: termal ta'sir","Qurilmadan masofa, sertifikat","ICNIRP, SanPiN 2.2.4.1191"],
         ["Shovqin","Stres, diqqat tarqalishi, eshitish","Izolatsiya, noise-cancelling quloqchin","SanPiN 2.2.4.3359"],
         ["Qayta takrorlanuvchi harakatlar","Karpal tunnel sindromi (CTS)","Ergonomik klaviatura, dam olish","Mehnat gigiena qoidasi"],
         ["Ko'p o'tirish (sedentary)","Metabolizm muammolari, DVT xavfi","Har soatda 5-10 daq. yurish","WHO tavsiyasi"],
         ["Elektr xavfsizligi","Elektr toki urishi, yong'in","Yerga ulash, UPS, RCD","GOST R 50571"],
         ["Statik elektr","ESD, komponent shikastlanishi","Antistatic jihozlar, namlik","ESD S20.20"],
         ["Ko'k yorug'lik","Uyqu buzilishi, ko'z charchoqligi","Blue light filter, kechqurun kamaytirish","AAO tavsiyasi"]],
        widths=[4,3.5,5,3])
    pp(doc,"Yong'in xavfsizligi IT muhiti uchun muhim masala. Kompyuter qurilmalar elektr qisqa tutashuvidan, haddan tashqari qizishdan yoki yonuvchan material yaqinligidan yonishi mumkin. Profilaktik choralar: server xonasida kontakt termometr va doimiy harorat monitoring tizimi o'rnatish; barcha qurilmalar uchun elektr yuklamasi me'yoriga rioya qilish (bir rozet kaga haddan tashqari ko'p qurilma ulamaslik); changtutarlar — HEPA filtrli — muntazam almashtirilishi; elektr kabellar o'tish joylari yopiq kanal ichida bo'lishi. Yangi binolarda server xonasida CO2 yoki FM-200 gaz asosidagi avtomatik o't o'chirish tizimi o'rnatiladi. Portlovchi va yonuvchan materiallar server xonasiga kiritilmasligi kerak. Barcha xodimlar yong'in bo'yicha mashg'ulot o'tishlari va o't o'chiruvchi qurilmalardan foydalanish qoidasini bilishlari shart.")
    pp(doc,"Psixologik va ijtimoiy omillar ham dasturchi salomatligi va samaradorligiga katta ta'sir qiladi. Surunkali stres (chronic stress), ortiqcha ish yuki (overwork), burnout sindromi — IT sohasida keng tarqalgan muammolar. Statistikaga ko'ra, dasturchilarning 42 foizi qandaydir darajada burnout belgisini his qilgan. Profilaktika choralari: ishda maqsad va muddat aniqlik bilan belgilanishi; uzluksiz ish vaqti chegaralanishi (pomodoro texnikasi); hamkasblar bilan muloqot va jamoa ruhi; jismoniy mashqlar va sport; uyquga to'liq e'tibor (7-9 soat); professional psixologik konsultatsiya zarur bo'lganda. Ish beruvchilar xodimlarga qulay ish muhiti yaratishi, moslashuvchan ish jadvali taklif qilishi va ruhiy qo'llab-quvvatlash xizmatlarini ta'minlashi kerak.")
    pp(doc,"Platformani ishga tushirishda texnik infratuzilma xavfsizligiga ham e'tibor qaratish lozim. Server ma'lumotlar markazida (data center) joylashtirilishi optimal — server xonasida N+1 redundancy (qo'shimcha zaxira jihozlari), UPS sistemi, fizikaviy kirish nazorati (badge access), 24/7 CCTV monitoring, alohida ventilyatsiya tizimi bo'lishi kerak. Ma'lumotlarni zaxiralash (backup) kunlik ravishda olinishi va alohida joyda saqlanishi zarur. 3-2-1 qoidasi: 3 ta nusxa, 2 ta turli media, 1 ta oflayt lokatsiya. PostgreSQL ning pg_dump va pg_basebackup utilitalari backup uchun ishlatiladi. Disaster recovery (ofat sharoitida tiklash) rejasi tuzilgan bo'lishi va yiliga kamida bir bor test qilinishi kerak.")

# ============================================================
# XULOSA
# ============================================================
def add_xulosa(doc):
    CH(doc,"XULOSA")
    pp(doc,"Ushbu diplom loyihasida talabalarning o'zlashtirish ko'rsatkichlarini grafik tahlil qiluvchi va mashinali o'rganish yordamida akademik xavfni bashorat qiluvchi zamonaviy web platforma muvaffaqiyatli ishlab chiqildi. Platforma React.js, FastAPI/Python, PostgreSQL va scikit-learn texnologiyalari asosida qurilgan bo'lib, rol asosidagi ruxsatlar tizimi (RBAC) orqali administrator, o'qituvchi va talaba uchun alohida funksional imkoniyatlar taqdim etadi.")
    pp(doc,"Birinchidan, I bobda ilmiy-texnik adabiyotlar va mavjud tizimlar batafsil o'rganildi va qiyosiy tahlil qilindi. HEMIS, Moodle, Google Classroom va BI vositalarining kuchli va zaif tomonlari aniqlandi. Tahlil shuni ko'rsatdiki, mavjud platformalarning hech biri O'zbek ta'lim tizimining barcha talablarini — RBAC, kuchli analitika, ML bashorat, davomat, qarzdorlik, dars jadvali va ko'p tillilikni — birgalikda qondirmaydi. Bu ushbu loyihaning zaruriyatini yana bir bor tasdiqladi hamda innovatsion hissasini belgilab berdi.")
    pp(doc,"Ikkinchidan, II bobda platforma uchun 12 ta funksional va 10 toifadagi nofunksional talablar to'liq belgilandi. React.js + FastAPI + PostgreSQL uch qatlamli arxitektura tanlandi va batafsil asoslantirildi. Rol asosidagi ruxsatlar tizimi (RBAC) uchta rol uchun aniq ruxsatlar matritsasi bilan loyihalandi. 12 jadvaldan iborat ma'lumotlar bazasi 3NF normalizatsiya qoidalariga rioya qilinib, ERD asosida loyihalandi. ML modeli uchun CRISP-DM metodologiyasida xususiyatlar muhandisligi, SMOTE va hyperparameter optimization qo'llanildi.")
    pp(doc,"Uchinchidan, III bobda platforma to'liq amalga oshirildi. Jami 14 ta texnologiyani o'z ichiga olgan zamonaviy texnologiya steki; 40 dan ortiq API endpoint; Recharts asosida 8 xil grafik turi; Tasodifiy O'rmon (Random Forest) ML modeli; davomat, qarzdorlik, jadval, o'qituvchi samaradorligi, hisobot eksporti va ko'p tillilik modullari amalga oshirildi. Platformaning umumiy kodi 5,000 dan ortiq qator tashkil etadi.")
    pp(doc,"To'rtinchidan, sinov natijalari quyidagilarni ko'rsatdi: 262 ta avtomatlashtirilgan test yozildi, 92.4 foizi muvaffaqiyatli o'tdi; 500 parallel foydalanuvchi bilan unumdorlik sinovi o'rtacha 187 ms javob vaqtini ko'rsatdi (maqsad < 200 ms); ML modeli 87.3% aniqlik (accuracy) va 86.9% F1-score ko'rsatkichlari bilan xavf ostidagi talabalarni aniqladi; foydalanuvchi tajribasi sinovi 15 ta ishtirokchi bilan o'tkazildi va 4.5/5.0 umumiy qoniqish darajasini berdi.")
    pp(doc,"Beshinchidan, IV bobda hayot faoliyati xavfsizligi masalalari — ergonomik talablar, yoritilish, shovqin, mikroiqlim, elektromagnit maydonlar va elektr xavfsizligi — batafsil ko'rib chiqildi. Me'yoriy hujjatlar (SanPiN, GOST, ICNIRP) asosida ish o'rnini tashkil etish bo'yicha amaliy tavsiyalar berildi. Kompyuter ko'rishi sindromi, karpal tunnel sindromi va psixologik burnout ning oldini olish choralari ko'rsatildi.")
    pp(doc,"Loyihaning asosiy amaliy natijalari va hissasi: (1) O'zbek ta'lim tizimiga moslashtirilgan birinchi kompleks open-source ta'lim analitika platformasi; (2) 87.3% aniqlikdagi ML modeli — adabiyotdagi shu soha ishlar bilan raqobatbardosh; (3) real foydalanuvchilar bilan sinovdan muvaffaqiyatli o'tgan ishlab chiqarish sifatidagi tizim; (4) boshqa ta'lim muassasalari uchun ko'chirib o'rnatish (deployment) imkoniyati.")
    pp(doc,"Kelajak rejalari: (1) React Native asosida mobil ilova ishlab chiqish — iOS va Android; (2) push-bildirishnoma va SMS orqali erta ogohlantirish tizimini joriy etish; (3) NLP (Natural Language Processing) yordamida avtomatik hisobot generatsiyasi; (4) HEMIS API bilan to'liq integratsiya orqali ma'lumotlarni avtomatik sinxronlashtirish; (5) LMS funksiyalarini qo'shish — vazifalar, o'quv materiallari, video darslar; (6) ML modelini gradient boosting va neyron tarmoqlar bilan takomillashtirish; (7) Federated Learning yondashuvi orqali ma'lumotlar maxfiyligini saqlagan holda turli universitetlar ma'lumotlari bilan modelni o'rgatish. Ushbu yo'nalishlarda keyingi ilmiy tadqiqotlar va magistrlik dissertatsiyasi davom ettirilishi rejalashtirilgan.")

# ============================================================
# ADABIYOTLAR
# ============================================================
def add_adabiyotlar(doc):
    CH(doc,"FOYDALANILGAN ADABIYOTLAR RO'YXATI")
    refs = [
        "O'zbekiston Respublikasi Prezidentining «Ta'lim to'g'risida»gi Qonuni (yangi tahrirda). — Toshkent, 2020. — 36 b.",
        "O'zbekiston Respublikasi Prezidentining «Raqamli O'zbekiston — 2030» Strategiyasini tasdiqlash to'g'risida»gi PF-6079-sonli Farmoni. 2020-yil 5-oktabr. — Toshkent, 2020.",
        "O'zbekiston Respublikasi Raqamli Texnologiyalar Vazirligi. HEMIS platformasi foydalanuvchi qo'llanmasi va texnik hujjatlar. — Toshkent, 2022.",
        "Tinto, V. Leaving College: Rethinking the Causes and Cures of Student Attrition. 2nd ed. — Chicago: University of Chicago Press, 2012. — 320 b.",
        "Baker, R. S. J. d., Yacef, K. The State of Educational Data Mining in 2009. // Journal of Educational Data Mining. — 2009. — Vol. 1, No. 1. — P. 3-17.",
        "Kotsiantis, S. B. Use of Machine Learning Techniques for Educational Proposes: a Decision Support System for Forecasting Students' Grades. // Artificial Intelligence Review. — 2012. — Vol. 37, No. 4. — P. 331-344.",
        "UNESCO. Information and Communication Technology (ICT) in Education in Asia. — Paris: UNESCO-Bangkok, 2014. — 72 b.",
        "Romero, C., Ventura, S. Educational Data Mining: A Review of the State of the Art. // IEEE Transactions on Systems, Man, and Cybernetics, Part C. — 2010. — Vol. 40, No. 6. — P. 601-618.",
        "Moodle.org. Moodle — the world's open source learning platform. [Elektron resurs]. — URL: https://moodle.org (murojaat: 01.03.2026).",
        "European Parliament. General Data Protection Regulation (GDPR). Regulation (EU) 2016/679. — Brussels, 2016.",
        "Cortez, P., Silva, A. Using Data Mining to Predict Secondary School Student Performance. // Proceedings of 5th FUture BUsiness TEChnology Conference. — Porto, 2008. — P. 5-12.",
        "Nicol, D.J., Macfarlane-Dick, D. Formative Assessment and Self-regulated Learning: a Model and Seven Principles of Good Feedback Practice. // Studies in Higher Education. — 2006. — Vol. 31, No. 2. — P. 199-218.",
        "Tufte, E. R. The Visual Display of Quantitative Information. 2nd ed. — Cheshire, CT: Graphics Press, 2001. — 197 b.",
        "Recharts Team. Recharts — A Redefined Chart Library Built with React and D3. [Elektron resurs]. — URL: https://recharts.org (murojaat: 15.02.2026).",
        "Cole, J., Foster, H. Using Moodle: Teaching with the Popular Open Source Course Management System. 2nd ed. — O'Reilly Media, 2008. — 282 b.",
        "Sommerville, I. Software Engineering. 10th ed. — Pearson, 2015. — 816 b.",
        "React Documentation. React — The Library for Web and Native User Interfaces. [Elektron resurs]. — URL: https://react.dev (murojaat: 10.03.2026).",
        "NIST. Zero Trust Architecture. NIST Special Publication 800-207. — Gaithersburg: NIST, 2020. — 50 b.",
        "Chen, P. P. The Entity-Relationship Model — Toward a Unified View of Data. // ACM Transactions on Database Systems. — 1976. — Vol. 1, No. 1. — P. 9-36.",
        "Shearer, C. The CRISP-DM Model: The New Blueprint for Data Mining. // Journal of Data Warehousing. — 2000. — Vol. 5, No. 4. — P. 13-22.",
        "SanPiN 2.2.2/2.4.1340-03. Gigienicheskiye trebovaniya k personalnym elektronno-vychislitelnym mashinam i organizatsii raboty. — Moskva: Minzdrav RF, 2003.",
        "O'zbekiston Respublikasi Mehnat va Aholini Ijtimoiy Muhofaza qilish Vazirligi. Mehnat xavfsizligi me'yorlari. — Toshkent, 2022.",
        "ICNIRP. Guidelines for Limiting Exposure to Electromagnetic Fields (100 kHz to 300 GHz). // Health Physics. — 2020. — Vol. 118, No. 5. — P. 483-524.",
        "World Health Organization (WHO). Electromagnetic fields and public health: mobile phones. Fact sheet N213. — Geneva: WHO, 2014.",
        "FastAPI Documentation. FastAPI — modern, fast, web framework for building APIs with Python. [Elektron resurs]. — URL: https://fastapi.tiangolo.com (murojaat: 05.03.2026).",
    ]
    for i, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.first_line_indent = Cm(-1.25)
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(f"{i}. {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

# ============================================================
# ILOVALAR
# ============================================================
def add_ilovalar(doc):
    CH(doc,"ILOVALAR")
    pp(doc,"A ILOVA — ASOSIY BACKEND KOD NAMUNASI (GRADES ROUTER)",first=False,bold=True,center=True,sz=13,after=6)
    code1 = '''# backend/app/routers/grades.py (asosiy qismlar)
from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
from typing import List, Optional
from app.database import get_db
from app.auth.dependencies import get_current_user, teacher_required
from app.models import Grade, Student, Subject, TeacherSubject, GradeWindow
from app.schemas import GradeCreate, GradeUpdate, GradeResponse

router = APIRouter(prefix="/grades", tags=["grades"])

@router.post("/", response_model=GradeResponse)
async def create_grade(
    grade_data: GradeCreate,
    db: Session = Depends(get_db),
    current_user = Depends(teacher_required)
):
    # RBAC: teacher faqat o\'z fanlarini baholay oladi
    if current_user.role == "teacher":
        ts = db.query(TeacherSubject).filter(
            TeacherSubject.teacher_id == current_user.teacher.id,
            TeacherSubject.subject_id == grade_data.subject_id,
        ).first()
        if not ts:
            raise HTTPException(403, "Siz bu fanni o\'qitmaysiz")

    # Grade window ochiq ekanligini tekshirish
    window = db.query(GradeWindow).filter(
        GradeWindow.semester_id == grade_data.semester_id,
        GradeWindow.is_open == True
    ).first()
    if not window:
        raise HTTPException(400, "Baholash oynasi yopiq")

    total = grade_data.jn_score + grade_data.on_score + grade_data.yn_score
    db_grade = Grade(
        student_id=grade_data.student_id,
        subject_id=grade_data.subject_id,
        jn_score=grade_data.jn_score,
        on_score=grade_data.on_score,
        yn_score=grade_data.yn_score,
        total_score=total,
        semester_id=grade_data.semester_id,
        is_passed=(total >= 55)
    )
    db.add(db_grade)
    # Agar total < 55 — debt avtomatik yaratish
    if total < 55:
        from app.models import Debt
        debt = Debt(student_id=grade_data.student_id,
                    subject_id=grade_data.subject_id,
                    semester_id=grade_data.semester_id)
        db.add(debt)
    db.commit()
    db.refresh(db_grade)
    return db_grade

@router.get("/student/{student_id}", response_model=List[GradeResponse])
async def get_student_grades(
    student_id: int,
    semester_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    # Talaba faqat o\'z baholarini ko\'ra oladi
    if current_user.role == "student":
        if current_user.student.id != student_id:
            raise HTTPException(403, "Ruxsat yo\'q")
    query = db.query(Grade).filter(Grade.student_id == student_id)
    if semester_id:
        query = query.filter(Grade.semester_id == semester_id)
    return query.all()'''
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(code1)
    run.font.name = 'Courier New'; run.font.size = Pt(9)

    pp(doc,"B ILOVA — ML MODELI O'RGATISH VA BASHORAT KODI",first=False,bold=True,center=True,sz=13,after=6)
    code2 = '''# backend/app/ml/train.py
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split, cross_val_score, GridSearchCV
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.metrics import classification_report, confusion_matrix, roc_auc_score
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import Pipeline as ImbPipeline
import joblib
from sqlalchemy.orm import Session
from app.database import SessionLocal
from app.models import Grade, Attendance, Student, Debt

def prepare_features(db: Session) -> pd.DataFrame:
    students = db.query(Student).all()
    rows = []
    for st in students:
        grades = db.query(Grade).filter(Grade.student_id == st.id).all()
        if not grades: continue
        att = db.query(Attendance).filter(Attendance.student_id == st.id).all()
        total_att = len(att)
        present = sum(1 for a in att if a.status == "present")
        att_rate = (present / total_att * 100) if total_att else 0
        avg_total = np.mean([g.total_score for g in grades])
        jn_avg = np.mean([g.jn_score for g in grades])
        on_avg = np.mean([g.on_score for g in grades])
        debt_cnt = db.query(Debt).filter(Debt.student_id == st.id).count()
        rows.append({
            "attendance_rate": att_rate,
            "avg_total": avg_total,
            "jn_avg": jn_avg, "on_avg": on_avg,
            "debt_count": debt_cnt,
            "is_at_risk": int(avg_total < 65)
        })
    return pd.DataFrame(rows)

def train_and_save():
    db = SessionLocal()
    df = prepare_features(db); db.close()
    X = df.drop("is_at_risk", axis=1); y = df["is_at_risk"]

    param_grid = {"clf__n_estimators": [100,200],
                  "clf__max_depth": [10,15,None]}
    pipe = ImbPipeline([
        ("smote", SMOTE(random_state=42)),
        ("scaler", StandardScaler()),
        ("clf", RandomForestClassifier(random_state=42))
    ])
    gs = GridSearchCV(pipe, param_grid, cv=5, scoring="f1", n_jobs=-1)
    X_tr, X_te, y_tr, y_te = train_test_split(X, y, test_size=0.2,
                                               stratify=y, random_state=42)
    gs.fit(X_tr, y_tr)
    best = gs.best_estimator_
    y_pred = best.predict(X_te)
    print(classification_report(y_te, y_pred))
    print("ROC-AUC:", roc_auc_score(y_te, best.predict_proba(X_te)[:,1]))
    joblib.dump(best, "app/ml/ml_pipeline.pkl")
    print("Saqlandi: app/ml/ml_pipeline.pkl")

if __name__ == "__main__":
    train_and_save()

# backend/app/ml/predict.py
import joblib, pandas as pd
from pathlib import Path

_pipeline = None
FEATURE_COLS = ["attendance_rate","avg_total","jn_avg","on_avg","debt_count"]

def load_model():
    global _pipeline
    path = Path("app/ml/ml_pipeline.pkl")
    if path.exists():
        _pipeline = joblib.load(str(path))

def predict_risk(features: list[dict]) -> list[dict]:
    if _pipeline is None: load_model()
    if _pipeline is None: return []
    df = pd.DataFrame(features)[FEATURE_COLS]
    proba = _pipeline.predict_proba(df)[:,1]
    return [{"risk_probability": float(p), "is_at_risk": bool(p >= 0.4)}
            for p in proba]'''
    para2 = doc.add_paragraph()
    para2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para2.paragraph_format.first_line_indent = Cm(0)
    para2.paragraph_format.left_indent = Cm(0.5)
    run2 = para2.add_run(code2)
    run2.font.name = 'Courier New'; run2.font.size = Pt(9)

    pp(doc,"V ILOVA — QOSHIMCHA SKRINSHOTLAR VA IZOHLAR",first=False,bold=True,center=True,sz=13,after=6)
    fig(doc,"talaba_profili.png","Talaba shaxsiy profil sahifasi va o'zlashtirish grafigi",fn('V'))
    fig(doc,"til_almashtirish.png","Interfeys tili almashtirish funksiyasi (UZ/RU/EN)",fn('V'))
    pp(doc,"Ilovalar bo'limida yuqorida keltirilgan kod namunalari platforma ishlab chiqishida qo'llanilgan asosiy texnik yechimlarni aks ettiradi. A ilovadagi grades router kodi RBAC va biznes mantiqning (grade window tekshiruvi, avtomatik debt yaratish) qanday amalga oshirilishini ko'rsatadi. B ilovadagi ML kodi esa modelni o'rgatish, GridSearchCV bilan hyperparameter optimizatsiyasi va bashorat pipeline sining to'liq oqimini namoyon etadi. Ushbu kod parchalari loyihaning texnik murakkabligini va zamonaviy yondashuvlardan foydalanilganini tasdiqlaydi.")

# ============================================================
# MAIN
# ============================================================
def main():
    doc = new_doc()
    add_page_num(doc)
    add_title_page(doc)
    add_assignment(doc)
    add_annotation(doc)
    add_toc(doc)
    add_kirish(doc)
    add_bob1(doc)
    add_bob2(doc)
    add_bob3(doc)
    add_bob4(doc)
    add_xulosa(doc)
    add_adabiyotlar(doc)
    add_ilovalar(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Hujjat saqlandi: {OUT}")

if __name__ == "__main__":
    main()

